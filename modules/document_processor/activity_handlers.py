# modules/document_processor/activity_handlers.py
import logging
from docx.enum.text import WD_BREAK
from .formatters import copiar_formatacao_bullets_completa, copiar_formatacao_run
from .table_handlers import (
    tab_mapear_formatacao_paragrafo,
    tab_processar_formatacao_paragrafo,
    tab_copiar_estilo_celula
)

# Configuração de logging
logger = logging.getLogger(__name__)

def processar_atividades_bullet_points(doc, atividades):
    """
    Localiza o marcador [ITEM] no documento e o substitui pelas atividades fornecidas,
    mantendo a formatação original e adicionando apenas uma quebra de página sem texto.
    """
    # Validação inicial dos dados de entrada
    if not atividades:
        logger.warning("Nenhuma atividade fornecida para substituição. O documento não será modificado.")
        return
    
    # Preparação para localizar o marcador [ITEM]
    paragrafo_referencia = None
    indice_elemento = None
    body = doc._body._body
    elementos = list(body)
    
    # Localizar o parágrafo com [ITEM] e o parágrafo de quebra de página
    paragrafo_quebra = None
    
    for i, paragrafo in enumerate(doc.paragraphs):
        if '[ITEM]' in paragrafo.text:
            paragrafo_referencia = paragrafo
            indice_elemento = elementos.index(paragrafo._element)
            logger.info(f"Marcador [ITEM] encontrado no documento, índice {i}")
            
            # Verifica se o próximo parágrafo contém quebra de página
            if i+1 < len(doc.paragraphs):
                # Verificamos se o próximo parágrafo contém uma quebra de página
                for run in doc.paragraphs[i+1].runs:
                    if hasattr(run, '_element') and 'br' in str(run._element.xml):
                        paragrafo_quebra = doc.paragraphs[i+1]
                        logger.info(f"Parágrafo com quebra de página encontrado no índice {i+1}")
                        break
            break
    
    # Verificação de segurança - se não encontrou o marcador
    if paragrafo_referencia is None:
        logger.error("Marcador [ITEM] não encontrado no documento. Nenhuma modificação será realizada.")
        return
    
    # Capturar informações do parágrafo de quebra de página (se existir)
    quebra_elemento = None
    if paragrafo_quebra:
        quebra_elemento = paragrafo_quebra._element
        indice_quebra = elementos.index(quebra_elemento)
        logger.info(f"Elemento de quebra encontrado no índice {indice_quebra}")
    
    try:
        # Lista para armazenar os parágrafos criados
        paragrafos_criados = []
        
        # Processa cada atividade
        for idx, atividade in enumerate(atividades):
            # Criação de um novo parágrafo
            novo_paragrafo = doc.add_paragraph()
            
            # Remove o novo parágrafo do documento (para manipulação)
            body.remove(novo_paragrafo._element)
            
            # Preserva a formatação do parágrafo original
            copiar_formatacao_bullets_completa(novo_paragrafo, paragrafo_referencia)
            
            # Adiciona o texto da atividade
            novo_run = novo_paragrafo.add_run(atividade['nome'])
            if paragrafo_referencia.runs:
                copiar_formatacao_run(novo_run, paragrafo_referencia.runs[0])
            
            # Armazena o novo parágrafo
            paragrafos_criados.append(novo_paragrafo._element)
            logger.info(f"Criado parágrafo para atividade: {atividade['nome']}")
            
            # Se for o último item, adiciona a quebra de página diretamente a ele
            if idx == len(atividades) - 1:
                novo_run.add_break(WD_BREAK.PAGE)
                logger.info("Adicionada quebra de página ao último item da lista")
        
        # Remove o parágrafo de referência ([ITEM]) e quebra de página (se existir)
        body.remove(paragrafo_referencia._element)
        if quebra_elemento is not None and quebra_elemento.getparent() is not None:
            body.remove(quebra_elemento)
        
        # Insere os novos parágrafos de atividades
        for i, paragrafo in enumerate(paragrafos_criados):
            body.insert(indice_elemento + i, paragrafo)
        
        logger.info(f"Adicionados {len(paragrafos_criados)} itens com quebra de página no último item")
        
    except Exception as e:
        logger.error(f"Erro durante a substituição das atividades: {e}")
        logger.exception("Detalhes do erro:")
        raise

def processar_atividades_tabela(doc, atividades):
    """
    Processa uma tabela do Word para inserir atividades mantendo a formatação original.
    """
    try:
        # Vamos localizar a tabela específica que contém os marcadores
        tabela_alvo = None
        linha_item = None
        linha_total = None

        # Iteramos por todas as tabelas do documento
        for table in doc.tables:
            for idx_row, row in enumerate(table.rows):
                texto_linha = ' '.join(cell.text for cell in row.cells)
                if '[ITEM]' in texto_linha:
                    tabela_alvo = table
                    linha_item = idx_row
                elif 'Total' in texto_linha and tabela_alvo == table:
                    linha_total = idx_row
                    break
            if tabela_alvo and linha_total is not None:
                break

        if not tabela_alvo:
            raise ValueError("Não foi possível encontrar a tabela com os marcadores necessários")

        logger.info(f"Tabela encontrada: linha do item={linha_item}, linha do total={linha_total}")

        # Preservamos a linha de totais removendo-a temporariamente
        linha_total_row = tabela_alvo.rows[linha_total]
        linha_total_elemento = linha_total_row._tr
        tabela_alvo._tbl.remove(linha_total_elemento)

        # Processamos a primeira atividade (substituindo os marcadores)
        primeira_linha = tabela_alvo.rows[linha_item]
        
        # Substituímos os marcadores na primeira linha
        for cell in primeira_linha.cells:
            for paragraph in cell.paragraphs:
                texto = paragraph.text
                if '[ITEM]' in texto:
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(paragraph)
                    tab_processar_formatacao_paragrafo(paragraph, atividades[0]['nome'], mapa_formatacao)
                elif '[HORAS_ITEM]' in texto:
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(paragraph)
                    tab_processar_formatacao_paragrafo(paragraph, str(atividades[0]['horas']), mapa_formatacao)

        # Para cada atividade adicional, criamos uma nova linha
        for atividade in atividades[1:]:
            nova_linha = tabela_alvo.add_row()
            
            # Copiamos a formatação da primeira linha para a nova
            for idx, (cell_destino, cell_origem) in enumerate(zip(nova_linha.cells, primeira_linha.cells)):
                # Copiamos o estilo da célula usando a função existente
                tab_copiar_estilo_celula(cell_destino, cell_origem)
                
                # Substituímos o conteúdo mantendo a formatação
                if idx == 0:  # Coluna do nome da atividade 
                    p_destino = cell_destino.paragraphs[0]
                    p_origem = cell_origem.paragraphs[0]
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(p_origem)
                    tab_processar_formatacao_paragrafo(p_destino, atividade['nome'], mapa_formatacao)
                elif idx == 1:  # Coluna das horas
                    p_destino = cell_destino.paragraphs[0]
                    p_origem = cell_origem.paragraphs[0]
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(p_origem)
                    tab_processar_formatacao_paragrafo(p_destino, str(atividade['horas']), mapa_formatacao)

        # Recolocamos a linha de totais no final
        tabela_alvo._tbl.append(linha_total_elemento)

        logger.info("Processamento das atividades na tabela concluído com sucesso")

    except Exception as e:
        logger.error(f"Erro ao processar atividades na tabela: {e}")
        logger.exception("Detalhes do erro:")
        raise