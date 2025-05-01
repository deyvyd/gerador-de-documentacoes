# modules/document_processor/table_handlers.py
import logging
from copy import deepcopy
from .formatters import copiar_formatacao_run

# Configuração de logging
logger = logging.getLogger(__name__)

def tab_mapear_formatacao_paragrafo(paragraph):
    """
    Cria um mapa da formatação de cada parte do texto em um parágrafo.
    """
    mapa_formatacao = []
    posicao_atual = 0
    
    # Analisamos cada run (segmento de texto) no parágrafo
    for run in paragraph.runs:
        # Só mapeamos se o run contiver texto
        comprimento = len(run.text)
        if comprimento > 0:
            mapa_formatacao.append({
                'inicio': posicao_atual,
                'fim': posicao_atual + comprimento,
                'run': run  # Guardamos o run original para copiar sua formatação
            })
            posicao_atual += comprimento
            
    return mapa_formatacao

def tab_processar_formatacao_paragrafo(paragraph, texto_novo, mapa_formatacao):
    """
    Processa e aplica formatação completa a um parágrafo, incluindo alinhamento e estilos.
    """
    try:
        # Parte 1: Obtenção do alinhamento
        alinhamento_real = None
        
        # 1.1 Verifica alinhamento direto
        if paragraph.alignment is not None:
            logger.debug(f"Alinhamento encontrado diretamente: {paragraph.alignment}")
            alinhamento_real = paragraph.alignment
            
        # 1.2 Verifica alinhamento nas propriedades XML
        if alinhamento_real is None and hasattr(paragraph._p, 'pPr'):
            pPr = paragraph._p.pPr
            if pPr is not None:
                jc_element = pPr.xpath('.//w:jc')
                if jc_element:
                    val = jc_element[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    logger.debug(f"Alinhamento encontrado no XML: {val}")
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    alignment_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    alinhamento_real = alignment_map.get(val)
        
        # 1.3 Verifica alinhamento na célula da tabela
        if alinhamento_real is None:
            cell = paragraph._p.getparent()
            while cell is not None:
                if cell.tag.endswith('tc'):
                    tcPr = cell.xpath('.//w:tcPr')
                    if tcPr:
                        align_elements = tcPr[0].xpath('.//w:vAlign')
                        if align_elements:
                            val = align_elements[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            logger.debug(f"Alinhamento encontrado na célula: {val}")
                            from docx.enum.text import WD_ALIGN_PARAGRAPH
                            cell_alignment_map = {
                                'center': WD_ALIGN_PARAGRAPH.CENTER,
                                'left': WD_ALIGN_PARAGRAPH.LEFT,
                                'right': WD_ALIGN_PARAGRAPH.RIGHT
                            }
                            alinhamento_real = cell_alignment_map.get(val)
                    break
                cell = cell.getparent()

        # Parte 2: Preservação e aplicação da formatação
        # 2.1 Salva propriedades do parágrafo
        estilo_paragrafo = None
        if hasattr(paragraph._p, 'pPr') and paragraph._p.pPr is not None:
            estilo_paragrafo = deepcopy(paragraph._p.pPr)
        
        # 2.2 Remove runs existentes
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
        
        # 2.3 Aplica o novo texto com formatação
        if not mapa_formatacao:
            run = paragraph.add_run(texto_novo)
        else:
            texto_original_len = mapa_formatacao[-1]['fim']
            proporcao = len(texto_novo) / texto_original_len if texto_original_len > 0 else 1
            
            for formato in mapa_formatacao:
                novo_inicio = int(formato['inicio'] * proporcao)
                novo_fim = int(formato['fim'] * proporcao)
                novo_fim = min(novo_fim, len(texto_novo))
                
                if novo_inicio < novo_fim:
                    texto_parte = texto_novo[novo_inicio:novo_fim]
                    run_novo = paragraph.add_run(texto_parte)
                    try:
                        copiar_formatacao_run(run_novo, formato['run'])
                    except Exception as e:
                        logger.warning(f"Erro ao copiar formatação do run: {e}")
                        configurar_fonte_padrao(run_novo)
        
        # 2.4 Restaura alinhamento e propriedades
        if alinhamento_real is not None:
            paragraph.alignment = alinhamento_real
            logger.debug(f"Alinhamento restaurado: {alinhamento_real}")
        
        if estilo_paragrafo is not None:
            if hasattr(paragraph._p, 'pPr'):
                for child in list(paragraph._p.pPr):
                    paragraph._p.pPr.remove(child)
                    
            for prop in estilo_paragrafo.xpath('*'):
                paragraph._p.get_or_add_pPr().append(deepcopy(prop))

    except Exception as e:
        logger.error(f"Erro ao processar formatação: {e}")
        # Tratamento de erro: mantém pelo menos o texto e alinhamento básico
        run = paragraph.add_run(texto_novo)
        if alinhamento_real is not None:
            paragraph.alignment = alinhamento_real

def tab_copiar_estilo_celula(celula_destino, celula_origem):
    """
    Copia a formatação completa de uma célula para outra, mantendo a formatação
    específica de cada parte do texto exatamente na posição correta.
    """
    try:
        # Primeiro capturamos o alinhamento original
        alinhamento_origem = None
        if celula_origem.paragraphs:
            alinhamento_origem = celula_origem.paragraphs[0].alignment
            logger.debug(f"Alinhamento original capturado: {alinhamento_origem}")

        # Copia propriedades da célula
        if hasattr(celula_origem._tc, 'tcPr'):
            tcPr_origem = celula_origem._tc.tcPr
            tcPr_destino = celula_destino._tc.get_or_add_tcPr()
            
            # Limpa propriedades existentes
            for child in list(tcPr_destino):
                tcPr_destino.remove(child)
            
            # Copia propriedades da célula
            for prop in tcPr_origem.xpath('*'):
                new_prop = deepcopy(prop)
                tcPr_destino.append(new_prop)

        # Trata os parágrafos
        if celula_origem.paragraphs and celula_destino.paragraphs:
            p_origem = celula_origem.paragraphs[0]
            p_destino = celula_destino.paragraphs[0]
            
            # Aplica o alinhamento capturado anteriormente
            if alinhamento_origem is not None:
                p_destino.alignment = alinhamento_origem
                logger.debug(f"Alinhamento aplicado: {alinhamento_origem}")
            
            # Copia propriedades específicas do parágrafo
            if hasattr(p_origem._p, 'pPr') and p_origem._p.pPr is not None:
                pPr_origem = p_origem._p.pPr
                pPr_destino = p_destino._p.get_or_add_pPr()
                
                # Limpa propriedades existentes mantendo o alinhamento
                alinhamento_temp = None
                if hasattr(pPr_destino, 'jc'):
                    alinhamento_temp = deepcopy(pPr_destino.jc)
                
                for child in list(pPr_destino):
                    pPr_destino.remove(child)
                
                # Restaura o alinhamento se existia
                if alinhamento_temp is not None:
                    pPr_destino.append(alinhamento_temp)
                
                # Copia outras propriedades do parágrafo
                for prop in pPr_origem.xpath('*'):
                    if prop.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc':
                        new_prop = deepcopy(prop)
                        pPr_destino.append(new_prop)

            # Obtém o texto completo da célula destino
            texto_destino = celula_destino.text.strip()
            palavras_destino = texto_destino.split()
            
            # Remove runs existentes mantendo o alinhamento
            alinhamento_atual = p_destino.alignment
            for run in p_destino.runs:
                run._element.getparent().remove(run._element)
            
            # Restaura o alinhamento
            p_destino.alignment = alinhamento_atual
            
            # Processa o texto preservando a formatação
            if texto_destino:
                run_destino = p_destino.add_run(texto_destino)
                if p_origem.runs:
                    copiar_formatacao_run(run_destino, p_origem.runs[0])

            logger.info("Formatação da célula copiada com sucesso, mantendo posicionamento correto")
            logger.debug(f"Alinhamento final: {p_destino.alignment}")

    except Exception as e:
        logger.error(f"Erro ao copiar estilo da célula: {e}")
        raise