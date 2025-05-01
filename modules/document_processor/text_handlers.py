# modules/document_processor/text_handlers.py
import logging
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
from .hyperlink_handlers import criar_hyperlink
from .formatters import copiar_formatacao_run

# Configuração de logging
logger = logging.getLogger(__name__)

def substituir_texto_preservando_formatacao(paragraph, substituicoes):
    """
    Substitui texto em um parágrafo preservando formatações e tratando links especialmente.
    Mantém quebras de linha simples sem adicionar linhas vazias extras.
    """
    try:
        texto_original = ''.join(run.text for run in paragraph.runs)
        
        # Tratamento especial para o link do board
        if '[LINK_BOARD]' in texto_original:
            url = substituicoes.get('[LINK_BOARD]', '')
            
            # Se o URL estiver vazio, retornamos sem fazer nada, pois o parágrafo
            # já terá sido removido pela função remover_paragrafo_link_board
            if not url or url.strip() == '':
                return False
                
            run_modelo = paragraph.runs[0] if paragraph.runs else None
            
            # Textos padrão para o link
            texto_antes = "Para maior detalhamento do fluxo das atividades basta acessar o "
            texto_depois = " com o seu login e senha da plataforma do GitLab. Caso seu usuário não tenha acesso a esse board, por favor solicitar acesso a LogAp."
            
            # Limpa o parágrafo mantendo alinhamento
            alinhamento_atual = paragraph.alignment
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
            paragraph.alignment = alinhamento_atual
            
            # Adiciona texto antes do link
            run_antes = paragraph.add_run(texto_antes)
            if run_modelo:
                copiar_formatacao_run(run_antes, run_modelo)
            
            # Cria e adiciona o hyperlink
            hyperlink = criar_hyperlink(paragraph, url, url, run_modelo)
            paragraph._p.append(hyperlink)
            
            # Adiciona texto depois do link
            run_depois = paragraph.add_run(texto_depois)
            if run_modelo:
                copiar_formatacao_run(run_depois, run_modelo)
            
            return True
            
        # Se não há marcadores para substituir, retorna
        if not any(marcador in texto_original for marcador in substituicoes.keys()):
            return False
            
        # Detecta se estamos processando o marcador de descrição
        processando_descricao = '[DESCRICAO]' in texto_original
        
        # Mapeia cada run com sua posição exata e formatação
        mapa_formatacao = []
        posicao = 0
        
        for run in paragraph.runs:
            if run.text:
                mapa_formatacao.append({
                    'inicio': posicao,
                    'fim': posicao + len(run.text),
                    'texto': run.text,
                    'run': run
                })
                posicao += len(run.text)
        
        # Faz todas as substituições necessárias
        novo_texto = texto_original
        substituicoes_realizadas = []
        
        for marcador, substituicao in substituicoes.items():
            if marcador in novo_texto and marcador != '[LINK_BOARD]':
                # Para a descrição, garantimos que as quebras de linha sejam simples
                if marcador == '[DESCRICAO]' and isinstance(substituicao, str):
                    # Normaliza todos os tipos de quebras de linha para \n
                    substituicao = substituicao.replace('\r\n', '\n').replace('\r', '\n')
                    # Evita quebras de linha duplas que causariam linhas vazias
                    substituicao = substituicao.replace('\n\n', '\n')
                
                # Guarda as posições antes da substituição
                pos_inicio = novo_texto.find(marcador)
                while pos_inicio != -1:
                    substituicoes_realizadas.append({
                        'inicio_original': pos_inicio,
                        'fim_original': pos_inicio + len(marcador),
                        'texto_novo': substituicao,
                        'tamanho_original': len(marcador),
                        'tamanho_novo': len(substituicao)
                    })
                    pos_inicio = novo_texto.find(marcador, pos_inicio + 1)
                
                novo_texto = novo_texto.replace(marcador, substituicao)
        
        # Remove runs existentes mantendo alinhamento
        alinhamento_atual = paragraph.alignment
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
        paragraph.alignment = alinhamento_atual
        
        # Para o texto da descrição, dividimos por quebras de linha e adicionamos cada parte como um run separado
        if processando_descricao and '\n' in novo_texto:
            linhas = novo_texto.split('\n')
            for i, linha in enumerate(linhas):
                run = paragraph.add_run(linha)
                # Aplicamos a formatação do primeiro run do parágrafo original
                if mapa_formatacao and mapa_formatacao[0]['run']:
                    copiar_formatacao_run(run, mapa_formatacao[0]['run'])
                
                # Adicionamos quebra de linha após cada linha, exceto a última
                if i < len(linhas) - 1:
                    run.add_break()  # Adiciona quebra de linha simples
        else:
            # Reconstrói o texto mantendo as formatações exatas
            posicao_atual = 0
            while posicao_atual < len(novo_texto):
                # Encontra qual run original corresponde a esta posição
                run_correspondente = None
                ajuste_posicao = 0
                
                # Ajusta a posição considerando as substituições realizadas
                posicao_ajustada = posicao_atual
                for subst in substituicoes_realizadas:
                    if posicao_atual >= subst['inicio_original']:
                        ajuste_posicao += subst['tamanho_novo'] - subst['tamanho_original']
                        posicao_ajustada = posicao_atual - ajuste_posicao
                
                # Encontra o run original correspondente à posição ajustada
                for formato in mapa_formatacao:
                    if formato['inicio'] <= posicao_ajustada < formato['fim']:
                        run_correspondente = formato['run']
                        break
                
                # Define quanto texto processar neste run
                if run_correspondente:
                    # Encontra o próximo ponto de quebra
                    proximo_ponto = len(novo_texto)
                    for formato in mapa_formatacao:
                        if formato['inicio'] > posicao_ajustada:
                            proximo_ponto = min(proximo_ponto, 
                                            formato['inicio'] + ajuste_posicao)
                    
                    texto_parte = novo_texto[posicao_atual:proximo_ponto]
                    
                    if texto_parte:
                        novo_run = paragraph.add_run(texto_parte)
                        copiar_formatacao_run(novo_run, run_correspondente)
                    
                    posicao_atual = proximo_ponto
                else:
                    # Se não encontrou formatação, usa a do último run
                    texto_parte = novo_texto[posicao_atual:]
                    novo_run = paragraph.add_run(texto_parte)
                    if mapa_formatacao:
                        copiar_formatacao_run(novo_run, mapa_formatacao[-1]['run'])
                    break
        
        return True
        
    except Exception as e:
        logger.error(f"Erro ao substituir texto preservando formatação: {e}")
        try:
            paragraph.text = texto_original
        except:
            pass
        return False

def remover_paragrafo_link_board(doc):
    """
    Remove parágrafos que contêm o marcador [LINK_BOARD] quando o link não foi fornecido
    e adiciona uma única quebra de página explícita no lugar.
    """
    # Lista de parágrafos que contêm o marcador e devem ser removidos
    paragrafos_para_remover = []
    indices_paragrafos = []
    
    # Primeiro, identificamos os parágrafos que contêm o marcador
    for i, paragrafo in enumerate(doc.paragraphs):
        if '[LINK_BOARD]' in paragrafo.text:
            paragrafos_para_remover.append(paragrafo._element)
            indices_paragrafos.append(i)
            logger.info(f"Marcado parágrafo com [LINK_BOARD] para remoção: índice {i}")
    
    # Remove os parágrafos identificados e adiciona quebra de página
    for elemento, indice in zip(paragrafos_para_remover, indices_paragrafos):
        if elemento.getparent() is not None:
            # Obtém o elemento pai e o índice do elemento no pai
            pai = elemento.getparent()
            indice_no_pai = pai.index(elemento)
            
            # Remove o elemento
            pai.remove(elemento)
            logger.info("Parágrafo com [LINK_BOARD] removido com sucesso")
            
            try:
                # Cria um novo parágrafo com quebra de página
                novo_paragrafo = doc.add_paragraph()
                novo_elemento = novo_paragrafo._element
                
                # Remove o parágrafo da posição atual
                if novo_elemento.getparent() is not None:
                    novo_elemento.getparent().remove(novo_elemento)
                
                # Adiciona o novo parágrafo na posição do removido
                pai.insert(indice_no_pai, novo_elemento)
                
                # Adiciona APENAS UMA quebra de página ao novo parágrafo
                run = novo_paragrafo.add_run()
                run.add_break(WD_BREAK.PAGE)
                
                logger.info("Uma única quebra de página adicionada no lugar do parágrafo removido")
            except Exception as e:
                logger.error(f"Erro ao adicionar quebra de página: {str(e)}")
    
    # Também verifica nas tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                paragrafos_para_remover = []
                for i, paragrafo in enumerate(celula.paragraphs):
                    if '[LINK_BOARD]' in paragrafo.text:
                        paragrafos_para_remover.append(paragrafo._element)
                        logger.info(f"Marcado parágrafo com [LINK_BOARD] em tabela para remoção")
                
                for elemento in paragrafos_para_remover:
                    if elemento.getparent() is not None:
                        # Remove o elemento
                        pai = elemento.getparent()
                        indice_no_pai = pai.index(elemento)
                        pai.remove(elemento)
                        logger.info("Parágrafo com [LINK_BOARD] em tabela removido com sucesso")
                        
                        try:
                            # Para tabelas, não adicionamos quebra de página, apenas um parágrafo vazio
                            # para manter o espaçamento
                            novo_paragrafo = celula.add_paragraph()
                            novo_elemento = novo_paragrafo._element
                            
                            # Remove o parágrafo da posição atual
                            if novo_elemento.getparent() is not None:
                                novo_elemento.getparent().remove(novo_elemento)
                            
                            # Adiciona o novo parágrafo na posição do removido
                            pai.insert(indice_no_pai, novo_elemento)
                            logger.info("Parágrafo vazio adicionado na tabela para substituir o removido")
                        except Exception as e:
                            logger.error(f"Erro ao adicionar parágrafo vazio na tabela: {str(e)}")