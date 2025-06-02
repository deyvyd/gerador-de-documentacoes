import logging
import os
import base64
import re
import html
from io import BytesIO
from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from copy import deepcopy
from bs4 import BeautifulSoup
from PIL import Image

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    encoding='utf-8'  # Adicionar encoding UTF-8
)
logger = logging.getLogger(__name__)

def processar_requisitos_funcionais(doc, requisitos):
    """
    Processa requisitos funcionais no documento de Estratégia de Solução.
    
    Esta função localiza o parágrafo modelo e a tabela modelo no documento,
    e então cria cópias para cada requisito funcional, substituindo os marcadores
    pelos valores correspondentes.
    
    Args:
        doc: Documento Word (.docx) aberto
        requisitos: Lista de requisitos funcionais para processar
    """
    if not requisitos or len(requisitos) == 0:
        logger.warning("Nenhum requisito funcional para processar")
        return
    
    logger.info(f"Processando {len(requisitos)} requisitos funcionais")
    
    # 1. Localizar parágrafo modelo (RF-[N_RF]:[TITULO_RF])
    paragrafo_modelo = None
    paragrafo_indice = None
    
    for i, p in enumerate(doc.paragraphs):      
        if "RF-[N_RF]" in p.text:
            paragrafo_modelo = p
            paragrafo_indice = i
            logger.info(f"Encontrado parágrafo modelo no índice {i}: {p.text}")
            break

    if not paragrafo_modelo:
        logger.error("Não foi possível encontrar o parágrafo modelo com 'RF-[N_RF]'")
        return
    
    # 2. Localizar tabela modelo que contém marcadores de RF
    tabela_modelo = None
    tabela_indice = None
    
    for i, tabela in enumerate(doc.tables):
        contem_marcador = False
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    if any(marcador in p.text for marcador in ["[LOCAL_RF]", "[USUARIO_RF]", "[TIPO_RF]"]):
                        contem_marcador = True
                        logger.info(f"Encontrado marcador na tabela {i}")
                        break
                if contem_marcador:
                    break
            if contem_marcador:
                break
        
        if contem_marcador:
            tabela_modelo = tabela
            tabela_indice = i
            logger.info(f"Encontrada tabela modelo no índice {i}")
            break
    
    if not tabela_modelo:
        logger.error("Não foi possível encontrar a tabela modelo com marcadores RF")
        return
    
    # 3. Localizar parágrafos de requisitos não-funcionais para preservá-los
    paragrafos_rnf = []
    for i, p in enumerate(doc.paragraphs):
        if "3 - Requisitos Não-Funcionais" in p.text or "RNF-[N_RNF]" in p.text or "[DESCRICAO_RNF]" in p.text:
            paragrafos_rnf.append((i, p))
            logger.info(f"Encontrado parágrafo RNF no índice {i}: '{p.text}'")
    
    # Armazenar o corpo do documento para inserções
    body = doc._body._body
    
    # 4. Detectar onde está a seção de RNF para inserir os RFs antes
    posicao_insercao_final = None
    for i, (idx, p) in enumerate(paragrafos_rnf):
        if "3 - Requisitos Não-Funcionais" in p.text:
            posicao_insercao_final = idx-1
            logger.info(f"Seção de RNF encontrada na posição {idx}")
            break
    
    # Se não encontrou a seção de RNF, usar o final do documento
    if posicao_insercao_final is None:
        posicao_insercao_final = len(doc.paragraphs)
        logger.info(f"Seção de RNF não encontrada, usando o final do documento (índice {posicao_insercao_final})")
    
    # CORREÇÃO: Criar uma lista de elementos a serem inseridos e depois inserir tudo de uma vez
    elementos_a_inserir = []
    
    # Para cada requisito, criar parágrafo de título e tabela
    for idx, requisito in enumerate(requisitos):
        logger.info(f"Processando requisito {idx+1}: {requisito.get('tituloRF', 'Sem título')}")
        
        elementos_requisito = []
        
        # Adicionar quebra de página apenas para requisitos APÓS o primeiro
        if idx > 0:
            logger.info("Adicionando quebra de página entre requisitos")
            quebra_pagina = doc.add_paragraph()
            run = quebra_pagina.add_run()
            run.add_break(WD_BREAK.PAGE)
            elementos_requisito.append(quebra_pagina._p)
        
        # Criar novo parágrafo para título
        novo_titulo = doc.add_paragraph()
        
        # Copiar estilo do parágrafo modelo
        if hasattr(paragrafo_modelo, 'style') and paragrafo_modelo.style:
            novo_titulo.style = paragrafo_modelo.style
            
        # Copiar formatação do parágrafo modelo
        if hasattr(paragrafo_modelo, 'paragraph_format'):
            # Copiar alinhamento
            if hasattr(paragrafo_modelo.paragraph_format, 'alignment'):
                novo_titulo.paragraph_format.alignment = paragrafo_modelo.paragraph_format.alignment
            # Copiar recuo
            if hasattr(paragrafo_modelo.paragraph_format, 'left_indent'):
                novo_titulo.paragraph_format.left_indent = paragrafo_modelo.paragraph_format.left_indent
            # Copiar espaçamento
            if hasattr(paragrafo_modelo.paragraph_format, 'space_before'):
                novo_titulo.paragraph_format.space_before = paragrafo_modelo.paragraph_format.space_before
            if hasattr(paragrafo_modelo.paragraph_format, 'space_after'):
                novo_titulo.paragraph_format.space_after = paragrafo_modelo.paragraph_format.space_after
        
        # Substituir os marcadores no texto do título
        num_rf = str(idx + 1).zfill(2)  # RF-01, RF-02, etc.
        titulo = requisito.get('tituloRF', '[Sem título]')
        
        # Criar o texto do título formatado
        run = novo_titulo.add_run(f"RF-{num_rf}: {titulo}")
        
        # Copiar formatação do run original
        if paragrafo_modelo.runs:
            run_modelo = paragrafo_modelo.runs[0]
            if hasattr(run_modelo, 'font'):
                # Copiar fonte
                if hasattr(run_modelo.font, 'name'):
                    run.font.name = run_modelo.font.name
                # Copiar tamanho
                if hasattr(run_modelo.font, 'size'):
                    run.font.size = run_modelo.font.size
                # Copiar negrito
                if hasattr(run_modelo.font, 'bold'):
                    run.font.bold = run_modelo.font.bold
                # Copiar itálico
                if hasattr(run_modelo.font, 'italic'):
                    run.font.italic = run_modelo.font.italic
                # Copiar cor
                if hasattr(run_modelo.font, 'color') and run_modelo.font.color and run_modelo.font.color.rgb:
                    run.font.color.rgb = run_modelo.font.color.rgb
        
        elementos_requisito.append(novo_titulo._p)
        
        # Obter valores (ou padrões) para os campos
        local = requisito.get('local', '')
        usuario = requisito.get('usuario', '')
        perfil = requisito.get('perfil', '')
        tipo = requisito.get('tipo', '')

        # ATUALIZAÇÃO: Todos os campos agora usam processamento HTML
        descricao_html = requisito.get('descricao', '')
        if not descricao_html:
            descricao_html = ""

        # Tratar validações com processamento HTML
        validacoes_html = requisito.get('validacoes', '')
        if not validacoes_html:
            validacoes_html = "Não há validação de campos."

        # Tratar regras com processamento HTML
        regras_html = requisito.get('regras', '')
        if not regras_html:
            regras_html = "Não há alterações de regra de negócio."

        # Tratar banco com processamento HTML
        banco_html = requisito.get('banco', '')
        if not banco_html:
            banco_html = "Não há alterações de banco de dados."
        
        # Criar uma cópia da tabela modelo via XML para preservar formatações
        nova_tabela_xml = deepcopy(tabela_modelo._tbl)
        
        # Configurar substituições para marcadores (apenas valores simples)
        substituicoes = {
            "[LOCAL_RF]": local,
            "[USUARIO_RF]": usuario,
            "[PERFIL_RF]": perfil,
            "[TIPO_RF]": tipo,
        }
        
        # Substituir os marcadores simples na tabela
        substituir_marcadores_tabela(nova_tabela_xml, substituicoes)
        
        elementos_requisito.append(nova_tabela_xml)
        
        # Armazenar os elementos para inserção posterior
        elementos_a_inserir.append((elementos_requisito, descricao_html, validacoes_html, regras_html, banco_html, requisito.get('imagens', [])))
    
    # 5. Remover os elementos originais do modelo antes de inserir os novos
    try:
        # Remover parágrafo modelo original
        if paragrafo_modelo._p.getparent() is not None:
            body.remove(paragrafo_modelo._p)
            logger.info(f"Removido parágrafo modelo original (índice {paragrafo_indice})")
        
        # Remover tabela modelo original
        if tabela_modelo._tbl.getparent() is not None:
            body.remove(tabela_modelo._tbl)
            logger.info("Removida tabela modelo original")
    except Exception as e:
        logger.error(f"Erro ao remover elementos originais: {str(e)}")
    
    # 6. Inserit todos os elementos criados antes da seção de RNF
    for idx, (elementos_requisito, descricao_html, validacoes_html, regras_html, banco_html, imagens) in enumerate(elementos_a_inserir):
        logger.info(f"Inserindo elementos do requisito {idx+1}")
        
        # Inserir todos os elementos do requisito
        for elemento in elementos_requisito:
            body.insert(posicao_insercao_final, elemento)
            posicao_insercao_final += 1
        
        # Obter a referência para a tabela inserida
        tabela_inserida = None
        for i, tabela in enumerate(doc.tables):
            # Verificamos se é a tabela que acabamos de adicionar
            contem_marcador = False
            for linha in tabela.rows:
                for celula in linha.cells:
                    for p in celula.paragraphs:
                        if any(marcador in p.text for marcador in ["[DESCRICAO_RF]", "[REGRAS_CAMPOS_RF]", "[REGRAS_NEGOCIO_RF]", "[BANCO_RF]", "[IMAGEM_RF]"]):
                            contem_marcador = True
                            break
                    if contem_marcador:
                        break
                if contem_marcador:
                    break
            
            if contem_marcador:
                tabela_inserida = tabela
                logger.info(f"Encontrada tabela inserida no índice {i}")
                break
        
        if tabela_inserida:
            # ATUALIZAÇÃO: Usar processamento HTML para todos os campos relevantes
            substituir_html_com_formatacao(tabela_inserida, "[DESCRICAO_RF]", descricao_html)
            substituir_html_com_formatacao(tabela_inserida, "[REGRAS_CAMPOS_RF]", validacoes_html)
            substituir_html_com_formatacao(tabela_inserida, "[REGRAS_NEGOCIO_RF]", regras_html)
            substituir_html_com_formatacao(tabela_inserida, "[BANCO_RF]", banco_html)

            # Processar imagens (se houver)
            if imagens:
                logger.info(f"Processando {len(imagens)} imagens para o requisito {idx+1}")
                
                # Localizar a célula com o marcador [IMAGEM_RF]
                celula_imagem = None
                
                for i, linha in enumerate(tabela_inserida.rows):
                    for j, celula in enumerate(linha.cells):
                        for p in celula.paragraphs:
                            if "[IMAGEM_RF]" in p.text:
                                celula_imagem = celula
                                logger.info(f"Encontrado marcador [IMAGEM_RF] na linha {i}, coluna {j}")
                                break
                        if celula_imagem:
                            break
                    if celula_imagem:
                        break
                
                if celula_imagem:
                    # Limpar o conteúdo da célula (remover o marcador [IMAGEM_RF])
                    for p in celula_imagem.paragraphs:
                        if "[IMAGEM_RF]" in p.text:
                            # Limpar o texto do parágrafo
                            p.clear()
                            # Definir alinhamento central para as imagens
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            # Adicionar as imagens ao parágrafo
                            processar_imagens(p, imagens)
                            logger.info("Imagens adicionadas na célula substituindo [IMAGEM_RF]")
                            break
                else:
                    logger.warning(f"Não foi possível encontrar célula com [IMAGEM_RF] para o requisito {idx+1}")
            else:
                # Se não há imagens, remover o marcador [IMAGEM_RF]
                for linha in tabela_inserida.rows:
                    for celula in linha.cells:
                        for p in celula.paragraphs:
                            if "[IMAGEM_RF]" in p.text:
                                p.text = p.text.replace("[IMAGEM_RF]", "")
        else:
            logger.warning(f"Não foi possível encontrar a tabela inserida para o requisito {idx+1}")
    
    # Não remover os parágrafos de requisitos não-funcionais
    logger.info("Preservando seção de requisitos não-funcionais")
    
    logger.info(f"Processamento de {len(requisitos)} requisitos concluído com sucesso")

    return posicao_insercao_final

# Funções auxiliares permanecem inalteradas
def substituir_marcadores_tabela(tabela_xml, substituicoes):
    """
    Substitui marcadores em uma tabela XML.
    
    Args:
        tabela_xml: Elemento XML da tabela
        substituicoes: Dicionário com os valores de substituição para cada marcador
    """
    # Namespace XML para Word
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    
    # Caso especial para [USUARIO_RF] | [PERFIL_RF]
    for elemento_t in tabela_xml.findall(".//w:t", namespace):
        texto = elemento_t.text
        
        if texto:
            # Caso especial para [USUARIO_RF] | [PERFIL_RF]
            if "[USUARIO_RF]" in texto and "[PERFIL_RF]" in texto:
                texto_substituido = texto.replace("[USUARIO_RF]", substituicoes["[USUARIO_RF]"])
                texto_substituido = texto_substituido.replace("[PERFIL_RF]", substituicoes["[PERFIL_RF]"])
                elemento_t.text = texto_substituido
            else:
                # Substituições normais (exceto os que precisam de formatação especial)
                marcadores_especiais = ["[IMAGEM_RF]", "[DESCRICAO_RF]", "[REGRAS_CAMPOS_RF]", "[REGRAS_NEGOCIO_RF]", "[BANCO_RF]"]
                for marcador, valor in substituicoes.items():
                    if marcador in texto and marcador not in marcadores_especiais:
                        elemento_t.text = texto.replace(marcador, valor)

def limpar_texto_html(html_texto):
    """
    Remove tags HTML e retorna o texto puro.
    Usada como uma função de fallback para substituição simples.
    
    Args:
        html_texto: Texto HTML do Quill Editor
        
    Returns:
        Texto limpo sem tags HTML
    """
    if not html_texto:
        return ""
    
    # Usar BeautifulSoup para limpar o HTML
    soup = BeautifulSoup(html_texto, 'html.parser')
    
    # Obter o texto limpo
    return soup.get_text()

def substituir_html_com_formatacao(tabela, marcador, html_texto):
    """
    Substitui um marcador por texto HTML formatado corretamente.
    
    Args:
        tabela: Tabela onde buscar o marcador
        marcador: Marcador a ser substituído (ex: "[DESCRICAO_RF]")
        html_texto: Texto HTML para substituir o marcador
    """
    if not html_texto:
        return
    
    # Localizar a célula que contém o marcador
    celula_destino = None
    for linha in tabela.rows:
        for celula in linha.cells:
            for p in celula.paragraphs:
                if marcador in p.text:
                    celula_destino = celula
                    break
            if celula_destino:
                break
        if celula_destino:
            break
    
    if not celula_destino:
        logger.warning(f"Marcador {marcador} não encontrado na tabela")
        return
    
    # Parsear o HTML
    soup = BeautifulSoup(html_texto, 'html.parser')
    
    # Limpar o conteúdo atual da célula
    for p in list(celula_destino.paragraphs):
        if marcador in p.text:
            # Mantemos o parágrafo, mas limpamos o texto
            p.clear()
            primeiro_paragrafo = p
            break
    else:
        # Se não encontrou o marcador, cria um novo parágrafo
        primeiro_paragrafo = celula_destino.add_paragraph()
    
    # Processar o conteúdo HTML e aplicar ao documento
    paragrafo_atual = primeiro_paragrafo
    
    # Manter controle do primeiro parágrafo
    primeiro_processado = True
    
    # Função auxiliar para processar conteúdo com formatação
    def processar_conteudo_formatado(elemento, paragrafo):
        """Processa elementos mantendo formatação como negrito, itálico, sublinhado e cores"""
        # Criar o run primeiro
        run = paragrafo.add_run(elemento.get_text())
        
        # Negrito
        if elemento.name == 'strong' or elemento.name == 'b':
            run.bold = True
        
        # Itálico
        elif elemento.name == 'em' or elemento.name == 'i':
            run.italic = True
        
        # Sublinhado
        elif elemento.name == 'u':
            run.underline = True
        
        # Código
        elif elemento.name == 'code':
            # Configurar fonte Courier New, tamanho 10
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            
            # Cor do texto azul
            run.font.color.rgb = RGBColor(3, 105, 161)  # #0369A1
            
            # Cor de fundo azul claro
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D2F0FD')  # #D2F0FD
            run._r.get_or_add_rPr().append(shading_elm)
        
        # Span com estilos
        elif elemento.name == 'span':
            style = elemento.get('style', '')
            logger.info(f"Processando span com estilo: {style}")
            
            # Verificar cor do texto
            cor_match = re.search(r'color:\s*rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', style)
            if cor_match:
                r, g, b = map(int, cor_match.groups())
                run.font.color.rgb = RGBColor(r, g, b)
                logger.info(f"Cor do texto definida: RGB({r}, {g}, {b})")
            
            # Verificar classe de alinhamento
            align_class = elemento.get('class', '')
            if align_class:
                if 'ql-align-center' in align_class:
                    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    logger.info("Alinhamento centralizado aplicado ao parágrafo")
                elif 'ql-align-right' in align_class:
                    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    logger.info("Alinhamento à direita aplicado ao parágrafo")
                elif 'ql-align-justify' in align_class:
                    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    logger.info("Alinhamento justificado aplicado ao parágrafo")
                elif 'ql-align-left' in align_class:
                    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    logger.info("Alinhamento à esquerda aplicado ao parágrafo")
        
        return run
    
    # Função auxiliar para processar listas
    def processar_lista(lista_elemento, usar_numeros=True):
        """Processa listas ordenadas ou não ordenadas"""
        nonlocal primeiro_processado, paragrafo_atual
        
        contador = 1
        for li in lista_elemento.find_all('li', recursive=False):
            if not primeiro_processado:
                paragrafo_atual = celula_destino.add_paragraph()
            else:
                primeiro_processado = False
            
            # Verificar se o item da lista tem data-list="bullet"
            if li.get('data-list') == 'bullet' or lista_elemento.get('data-list') == 'bullet':
                # Lista não ordenada
                paragrafo_atual.add_run("• ")
            elif usar_numeros:
                # Lista ordenada
                paragrafo_atual.add_run(f"{contador}. ")
                contador += 1
            else:
                # Lista não ordenada (fallback)
                paragrafo_atual.add_run("• ")
            
            # Processar o conteúdo do item de lista mantendo formatação
            for item in li.contents:
                if item.name in ['strong', 'b', 'em', 'i', 'u', 'code', 'span']:  # Adicionado 'u', 'code' e 'span'
                    processar_conteudo_formatado(item, paragrafo_atual)
                elif item.name in ['ul', 'ol']:
                    # Lista aninhada - processar recursivamente
                    processar_lista(item, item.name == 'ol' and item.get('data-list') != 'bullet')
                elif item.string:
                    paragrafo_atual.add_run(item.string)
            
            # Aplicar recuo
            paragrafo_atual.paragraph_format.left_indent = Pt(20)
    
    # Processar cada elemento do HTML
    for elemento in soup:
        # Ignorar elementos vazios ou apenas com espaço
        if elemento.name is None and (not elemento.string or not elemento.string.strip()):
            continue
            
        # Processar parágrafos
        if elemento.name == 'p':
            # Se não for o primeiro parágrafo, criar um novo
            if not primeiro_processado:
                paragrafo_atual = celula_destino.add_paragraph()
            else:
                primeiro_processado = False
            
            # Verificar classe de alinhamento no parágrafo
            align_class = elemento.get('class', '')
            if align_class:
                if 'ql-align-center' in align_class:
                    paragrafo_atual.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    logger.info("Alinhamento centralizado aplicado ao parágrafo")
                elif 'ql-align-right' in align_class:
                    paragrafo_atual.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    logger.info("Alinhamento à direita aplicado ao parágrafo")
                elif 'ql-align-justify' in align_class:
                    paragrafo_atual.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    logger.info("Alinhamento justificado aplicado ao parágrafo")
                elif 'ql-align-left' in align_class:
                    paragrafo_atual.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    logger.info("Alinhamento à esquerda aplicado ao parágrafo")
            
            # Processar o conteúdo do parágrafo
            for conteudo in elemento.contents:
                if conteudo.name in ['strong', 'b', 'em', 'i', 'u', 'code', 'span']:  # Adicionado 'u', 'code' e 'span'
                    processar_conteudo_formatado(conteudo, paragrafo_atual)
                elif conteudo.name == 'ul' or conteudo.name == 'ol':
                    # Lista dentro de parágrafo
                    usar_numeros = conteudo.name == 'ol' and conteudo.get('data-list') != 'bullet'
                    processar_lista(conteudo, usar_numeros)
                elif conteudo.string:
                    paragrafo_atual.add_run(conteudo.string)
        
        # Se for lista
        elif elemento.name == 'ul' or elemento.name == 'ol':
            # Verificar se é uma lista não ordenada do Quill
            usar_numeros = elemento.name == 'ol' and elemento.get('data-list') != 'bullet'
            processar_lista(elemento, usar_numeros)
        
        # Se for texto plano fora de tags
        elif elemento.string and elemento.string.strip():
            if not primeiro_processado:
                paragrafo_atual = celula_destino.add_paragraph()
            else:
                primeiro_processado = False
            paragrafo_atual.add_run(elemento.string.strip())
        
        # Se for quebra de linha
        elif elemento.name == 'br':
            if not primeiro_processado:
                paragrafo_atual = celula_destino.add_paragraph()
            else:
                primeiro_processado = False
                
def substituir_texto_com_quebras(tabela, marcador, texto):
    """
    Substitui um marcador por texto preservando quebras de linha.
    
    Args:
        tabela: Tabela onde buscar o marcador
        marcador: Marcador a ser substituído (ex: "[REGRAS_NEGOCIO_RF]")
        texto: Texto com quebras de linha para substituir o marcador
    """
    if not texto:
        return
    
    # Localizar a célula que contém o marcador
    celula_destino = None
    paragrafo_original = None
    
    for linha in tabela.rows:
        for celula in linha.cells:
            for p in celula.paragraphs:
                if marcador in p.text:
                    celula_destino = celula
                    paragrafo_original = p
                    break
            if celula_destino:
                break
        if celula_destino:
            break
    
    if not celula_destino:
        logger.warning(f"Marcador {marcador} não encontrado na tabela")
        return
    
    # Dividir o texto pelas quebras de linha
    linhas = texto.split('\n')
    
    # Limpar o parágrafo original
    paragrafo_original.clear()
    
    # Adicionar a primeira linha ao parágrafo original
    paragrafo_original.add_run(linhas[0])
    
    # Para cada linha adicional, criar um novo parágrafo
    for i in range(1, len(linhas)):
        if linhas[i].strip():  # Ignorar linhas vazias
            p = celula_destino.add_paragraph()
            p.add_run(linhas[i])
            
            # Se a linha começa com um marcador de lista (•, -, *), aplicar recuo
            if linhas[i].strip().startswith(('•', '-', '*')):
                p.paragraph_format.left_indent = Pt(20)

def processar_imagens(paragrafo, imagens):
    """
    Processa imagens em base64 e as adiciona ao parágrafo.
    
    Args:
        paragrafo: Parágrafo do documento Word
        imagens: Lista de imagens em formato base64
    """
    if not imagens or len(imagens) == 0:
        logger.warning("Nenhuma imagem para processar")
        return
    
    logger.info(f"Processando {len(imagens)} imagens")
    
    # Processar todas as imagens da lista
    for i, imagem in enumerate(imagens):
        # Verifica se é uma imagem em base64
        if isinstance(imagem, str) and imagem.startswith('data:image'):
            try:
                # Extrair a parte base64
                base64_data = imagem.split(',')[1]
                imagem_bytes = BytesIO(base64.b64decode(base64_data))
                
                # Adicionar uma quebra de linha entre imagens se não for a primeira
                if i > 0:
                    paragrafo.add_run().add_break()
                
                # Adicionar ao documento
                run = paragrafo.add_run()
                # Obter dimensões originais da imagem
                from PIL import Image
                import io

                # Resetar o ponteiro do BytesIO para o início
                imagem_bytes.seek(0)
                with Image.open(imagem_bytes) as img:
                    largura_original_px, altura_original_px = img.size

                # Resetar novamente para o python-docx
                imagem_bytes.seek(0)

                # Converter pixels para polegadas (assumindo 96 DPI)
                DPI = 96
                largura_original_inches = largura_original_px / DPI

                # Definir largura máxima (6 polegadas)
                largura_maxima_inches = 6

                # Se a largura original for menor que o máximo, usar o tamanho original
                # Caso contrário, redimensionar mantendo a proporção
                if largura_original_inches <= largura_maxima_inches:
                    # Usar tamanho original
                    run.add_picture(imagem_bytes)
                    logger.info(f"Imagem {i+1} adicionada com tamanho original: {largura_original_px}x{altura_original_px}px ({largura_original_inches:.2f}\")")
                else:
                    # Redimensionar para a largura máxima
                    run.add_picture(imagem_bytes, width=Inches(largura_maxima_inches))
                    logger.info(f"Imagem {i+1} redimensionada para largura máxima: {largura_maxima_inches}\" (original: {largura_original_inches:.2f}\")")
                
                logger.info(f"Imagem {i+1} processada com sucesso")
                
            except Exception as e:
                logger.error(f"Erro ao processar imagem base64 {i+1}: {str(e)}")
                erro_run = paragrafo.add_run(f"[Erro ao processar imagem {i+1}]")
                # Destacar o erro em vermelho
                erro_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            logger.error(f"Formato de imagem não suportado para imagem {i+1}")
            erro_run = paragrafo.add_run(f"[Formato de imagem {i+1} não suportado]")
            erro_run.font.color.rgb = RGBColor(255, 0, 0)