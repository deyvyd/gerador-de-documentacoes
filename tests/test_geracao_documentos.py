"""Testes da rota /gerar_documentos (formulários técnico e desenvolvimento)."""
import io
import json
import zipfile

import pypdf
from docx import Document

from conftest import extrair_zip


# ---------------------------------------------------------------- JSON

def test_tec_apenas_json(client, form_tec):
    form_tec.update({"apenas_json": "true", "gerar_json": "true"})
    r = client.post("/gerar_documentos", data=form_tec)
    assert r.status_code == 200
    # apenas_json responde o arquivo JSON diretamente, sem zip
    assert r.content_type.startswith("application/json")

    dados = json.loads(r.data)
    assert dados["tipo"] == "tecnica"
    assert dados["info"]["numeroSS"] == "10"
    # Datas convertidas para formato brasileiro
    assert dados["info"]["dataInicio"] == "01/07/2026"
    assert dados["info"]["dataFim"] == "10/07/2026"
    assert len(dados["atividades"]) == 2


def test_dev_apenas_json(client, form_dev):
    form_dev.update({"apenas_json": "true", "gerar_json": "true"})
    r = client.post("/gerar_documentos", data=form_dev)
    assert r.status_code == 200
    assert r.content_type.startswith("application/json")
    dados = json.loads(r.data)
    assert dados["tipo"] == "desenvolvimento"


# ---------------------------------------------------------------- DOCX

def test_tec_gera_docx_validos(client, form_tec):
    form_tec.update({"gerar_docx": "true"})
    r = client.post("/gerar_documentos", data=form_tec)
    assert r.status_code == 200
    z = extrair_zip(r)
    docxs = [n for n in z.namelist() if n.endswith(".docx")]
    # Os 3 modelos técnicos: Estimativa, Estratégia, Relatório
    assert len(docxs) == 3

    for nome in docxs:
        doc = Document(io.BytesIO(z.read(nome)))  # falha se docx corrompido
        texto = "\n".join(p.text for p in doc.paragraphs)
        assert texto.strip(), f"{nome} sem conteúdo"


def test_tec_docx_contem_dados_do_formulario(client, form_tec):
    form_tec.update({"gerar_docx": "true"})
    r = client.post("/gerar_documentos", data=form_tec)
    z = extrair_zip(r)
    estrategia = next(n for n in z.namelist() if "Estrat" in n and n.endswith(".docx"))
    doc = Document(io.BytesIO(z.read(estrategia)))
    texto = "\n".join(p.text for p in doc.paragraphs)
    texto += "\n".join(
        c.text for t in doc.tables for row in t.rows for c in row.cells
    )
    assert "SS de teste" in texto
    assert "Análise" in texto or "Desenvolvimento" in texto


def test_tec_docx_sumario_com_links_internos(client, form_tec):
    """Sumário do docx deve ter hyperlinks de âncora (w:anchor)."""
    form_tec.update({"gerar_docx": "true"})
    r = client.post("/gerar_documentos", data=form_tec)
    z = extrair_zip(r)
    estrategia = next(n for n in z.namelist() if "Estrat" in n and n.endswith(".docx"))
    with zipfile.ZipFile(io.BytesIO(z.read(estrategia))) as dz:
        xml = dz.read("word/document.xml").decode("utf-8")
    assert xml.count("w:anchor=") > 0, "sumário sem hyperlinks internos"
    assert xml.count("bookmarkStart") > 0, "documento sem bookmarks"


# ---------------------------------------------------------------- Nome do ZIP

def test_nome_do_zip_segue_padrao(client, form_tec):
    form_tec.update({"gerar_docx": "true"})
    r = client.post("/gerar_documentos", data=form_tec)
    cd = r.headers.get("Content-Disposition", "")
    assert "SS 010-2026.zip" in cd


# ---------------------------------------------------------------- Erros

def test_form_sem_campos_obrigatorios_retorna_json_de_erro(client):
    r = client.post("/gerar_documentos", data={"tipo": "tecnica"})
    assert r.status_code == 500
    assert r.content_type.startswith("application/json"), (
        "erro deve ser JSON, nunca HTML (frontend faz response.json())"
    )
    assert "error" in r.get_json()


def test_atividades_json_invalido_retorna_json_de_erro(client, form_tec):
    form_tec["atividades"] = "{isso não é json"
    r = client.post("/gerar_documentos", data=form_tec)
    assert r.status_code == 500
    assert r.content_type.startswith("application/json")
