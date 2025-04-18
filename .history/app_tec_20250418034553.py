# app_tec.py
import json
import logging
import os
import re
import shutil
import tempfile
import zipfile
from copy import deepcopy
from datetime import date, datetime
from urllib.parse import quote

from docx import Document
from flask import request, send_file, jsonify

# Configuração do logger
logger = logging.getLogger(__name__)

class RelatorioAcompanhamentoProjeto:
    def __init__(self, numero_ss, ano_ss, iniciais_autor, titulo_ss, descricao,
                 data_inicio, data_fim, total_horas, link_board, atividades=None):
        self.numero_ss = numero_ss
        self.ano_ss = ano_ss
        self.iniciais_autor = iniciais_autor
        self.titulo_ss = titulo_ss
        self.descricao = descricao
        self.data_inicio = datetime.strptime(data_inicio, '%Y-%m-%d').date()
        self.data_fim = datetime.strptime(data_fim, '%Y-%m-%d').date()
        self.total_horas = int(total_horas)
        self.link_board = link_board 
        self.atividades = atividades or []
        
        # Calcular valores derivados
        self.dias_uteis = self._calcular_dias_uteis()
        self.n_pf = self._calcular_pontos_funcao()
        
    def _calcular_dias_uteis(self):
        dias = (self.data_fim - self.data_inicio).days + 1
        dias_uteis = sum(1 for day in range(dias) 
                        if (self.data_inicio + date.resolution * day).weekday() < 5)
        return dias_uteis
    
    def _calcular_pontos_funcao(self):
        """
        Calcula os pontos de função baseado no total de horas
        Fórmula: (total_horas / 10) * (250 / 100)
        """
        try:
            return round((self.total_horas / 10) * (250 / 100), 2)
        except Exception as e:
            logger.error(f"Erro ao calcular pontos de função: {e}")
            return 0
    
    def get_substituicoes(self):
        """Retorna o dicionário de substituições para o documento"""
        data_atual = datetime.now().strftime('%d/%m/%Y')
        
        return {
            '[NNN]': str(self.numero_ss).zfill(3),
            '[AAAA]': str(self.ano_ss),
            '[INICIAIS_AUTOR]': str(self.iniciais_autor),
            '[TITULO]': str(self.titulo_ss),
            '[DESCRICAO]': str(self.descricao),
            '[DATA_ATUAL]': data_atual,
            '[DATA_INICIO]': self.data_inicio.strftime('%d/%m/%Y'),
            '[DATA_FIM]': self.data_fim.strftime('%d/%m/%Y'),
            '[TOTAL_HORAS]': str(self.total_horas),
            '[DIAS_UTEIS]': str(self.dias_uteis),
            '[N_PF]': str(self.n_pf),
            '[LINK_BOARD]': str(self.link_board)
        }

def gerar_nome_arquivo(tipo_documento, numero_ss, ano_ss):
    """Gera o nome do arquivo baseado no tipo de documento"""
    base_codigo = "CTNI-0534-2021.00"
    
    if tipo_documento == 'relatorio':
        prefixo = "AP"
        sufixo = "Relatório de Acompanhamento de Projeto"
    elif tipo_documento == 'estimativa':
        prefixo = "ES"
        sufixo = "Estimativa de Esforço e Cronograma"
    elif tipo_documento == 'estrategia':
        prefixo = "ES"
        sufixo = "Estratégia de solução"
    else:
        prefixo = "DOC"
        sufixo = "Documento"

    nome_final_arquivo = f"{prefixo}-{base_codigo}-{numero_ss}-{ano_ss} - {sufixo}.docx"
    logger.info(f"Nome final arquivo: {nome_final_arquivo}")
    
    return nome_final_arquivo

def processar_atividades_bullet_points(doc, atividades):
    """
    Localiza o marcador [ITEM] no documento e o substitui pelas atividades fornecidas,
    mantendo a formatação original e adicionando apenas uma quebra de página sem texto.
    """
    # Validação inicial dos dados de entrada
    if not atividades:
        logger.warning("Nenhuma atividade fornecida para substituição. O documento não será modificado.")
        return
    
    # Preparação para localizar o marcador [ITEM]
    paragrafo_referencia = None
    indice_elemento = None
    body = doc._body._body
    elementos = list(body)
    
    # Localizar o parágrafo com [ITEM] e o parágrafo de quebra de página
    paragrafo_quebra = None
    
    for i, paragrafo in enumerate(doc.paragraphs):
        if '[ITEM]' in paragrafo.text:
            paragrafo_referencia = paragrafo
            indice_elemento = elementos.index(paragrafo._element)
            logger.info(f"Marcador [ITEM] encontrado no documento, índice {i}")
            
            # Verifica se o próximo parágrafo contém quebra de página
            if i+1 < len(doc.paragraphs):
                # Verificamos se o próximo parágrafo contém uma quebra de página
                for run in doc.paragraphs[i+1].runs:
                    if hasattr(run, '_element') and 'br' in str(run._element.xml):
                        paragrafo_quebra = doc.paragraphs[i+1]
                        logger.info(f"Parágrafo com quebra de página encontrado no índice {i+1}")
                        break
            break
    
    # Verificação de segurança - se não encontrou o marcador
    if paragrafo_referencia is None:
        logger.error("Marcador [ITEM] não encontrado no documento. Nenhuma modificação será realizada.")
        return
    
    # Capturar informações do parágrafo de quebra de página (se existir)
    quebra_elemento = None
    if paragrafo_quebra:
        quebra_elemento = paragrafo_quebra._element
        indice_quebra = elementos.index(quebra_elemento)
        logger.info(f"Elemento de quebra encontrado no índice {indice_quebra}")
    
    try:
        # Lista para armazenar os parágrafos criados
        paragrafos_criados = []
        
        # Processa cada atividade
        for idx, atividade in enumerate(atividades):
            # Criação de um novo parágrafo
            novo_paragrafo = doc.add_paragraph()
            
            # Remove o novo parágrafo do documento (para manipulação)
            body.remove(novo_paragrafo._element)
            
            # Preserva a formatação do parágrafo original
            copiar_formatacao_bullets_completa(novo_paragrafo, paragrafo_referencia)
            
            # Adiciona o texto da atividade
            novo_run = novo_paragrafo.add_run(atividade['nome'])
            if paragrafo_referencia.runs:
                copiar_formatacao_run(novo_run, paragrafo_referencia.runs[0])
            
            # Armazena o novo parágrafo
            paragrafos_criados.append(novo_paragrafo._element)
            logger.info(f"Criado parágrafo para atividade: {atividade['nome']}")
            
            # Se for o último item, adiciona a quebra de página diretamente a ele
            if idx == len(atividades) - 1:
                from docx.enum.text import WD_BREAK
                novo_run.add_break(WD_BREAK.PAGE)
                logger.info("Adicionada quebra de página ao último item da lista")
        
        # Remove o parágrafo de referência ([ITEM]) e quebra de página (se existir)
        body.remove(paragrafo_referencia._element)
        if quebra_elemento is not None and quebra_elemento.getparent() is not None:
            body.remove(quebra_elemento)
        
        # Insere os novos parágrafos de atividades
        for i, paragrafo in enumerate(paragrafos_criados):
            body.insert(indice_elemento + i, paragrafo)
        
        logger.info(f"Adicionados {len(paragrafos_criados)} itens com quebra de página no último item")
        
    except Exception as e:
        logger.error(f"Erro durante a substituição das atividades: {e}")
        logger.exception("Detalhes do erro:")
        raise

def processar_atividades_tabela(doc, atividades):
    """
    Processa uma tabela do Word para inserir atividades mantendo a formatação original.
    
    Esta função realiza as seguintes operações:
    1. Localiza a tabela que contém os marcadores [ITEM] e [HORAS_ITEM]
    2. Preserva o cabeçalho da tabela
    3. Substitui os marcadores pelos dados da primeira atividade
    4. Adiciona linhas adicionais para as demais atividades
    5. Mantém a linha de totais no final com sua formatação original
    """
    try:
        # Vamos localizar a tabela específica que contém os marcadores
        tabela_alvo = None
        linha_item = None
        linha_total = None

        # Iteramos por todas as tabelas do documento
        for table in doc.tables:
            for idx_row, row in enumerate(table.rows):
                texto_linha = ' '.join(cell.text for cell in row.cells)
                if '[ITEM]' in texto_linha:
                    tabela_alvo = table
                    linha_item = idx_row
                elif 'Total' in texto_linha and tabela_alvo == table:
                    linha_total = idx_row
                    break
            if tabela_alvo and linha_total is not None:
                break

        if not tabela_alvo:
            raise ValueError("Não foi possível encontrar a tabela com os marcadores necessários")

        logger.info(f"Tabela encontrada: linha do item={linha_item}, linha do total={linha_total}")

        # Preservamos a linha de totais removendo-a temporariamente
        linha_total_row = tabela_alvo.rows[linha_total]
        linha_total_elemento = linha_total_row._tr
        tabela_alvo._tbl.remove(linha_total_elemento)

        # Processamos a primeira atividade (substituindo os marcadores)
        primeira_linha = tabela_alvo.rows[linha_item]
        
        # Substituímos os marcadores na primeira linha
        for cell in primeira_linha.cells:
            for paragraph in cell.paragraphs:
                texto = paragraph.text
                if '[ITEM]' in texto:
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(paragraph)
                    tab_processar_formatacao_paragrafo(paragraph, atividades[0]['nome'], mapa_formatacao)
                elif '[HORAS_ITEM]' in texto:
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(paragraph)
                    tab_processar_formatacao_paragrafo(paragraph, str(atividades[0]['horas']), mapa_formatacao)

        # Para cada atividade adicional, criamos uma nova linha
        for atividade in atividades[1:]:
            nova_linha = tabela_alvo.add_row()
            
            # Copiamos a formatação da primeira linha para a nova
            for idx, (cell_destino, cell_origem) in enumerate(zip(nova_linha.cells, primeira_linha.cells)):
                # Copiamos o estilo da célula usando a função existente
                tab_copiar_estilo_celula(cell_destino, cell_origem)
                
                # Substituímos o conteúdo mantendo a formatação
                if idx == 0:  # Coluna do nome da atividade 
                    p_destino = cell_destino.paragraphs[0]
                    p_origem = cell_origem.paragraphs[0]
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(p_origem)
                    tab_processar_formatacao_paragrafo(p_destino, atividade['nome'], mapa_formatacao)
                elif idx == 1:  # Coluna das horas
                    p_destino = cell_destino.paragraphs[0]
                    p_origem = cell_origem.paragraphs[0]
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(p_origem)
                    tab_processar_formatacao_paragrafo(p_destino, str(atividade['horas']), mapa_formatacao)

        # Recolocamos a linha de totais no final
        tabela_alvo._tbl.append(linha_total_elemento)

        logger.info("Processamento das atividades na tabela concluído com sucesso")

    except Exception as e:
        logger.error(f"Erro ao processar atividades na tabela: {e}")
        logger.exception("Detalhes do erro:")
        raise

def gerar_relatorio_tecnico():
    """
    Função para gerar relatórios técnicos.
    Esta função foi extraída do app.py original para ser chamada pelo novo ponto de entrada.
    """
    logger.info("Iniciando geração de relatórios técnicos")
    
    # Listas para controle de arquivos temporários e arquivos de saída
    temp_files = []    # Arquivos temporários que serão deletados ao final
    output_files = []  # Arquivos que serão incluídos no ZIP final
    zip_path = None    # Caminho do arquivo ZIP que será gerado
    
    try:
        # Obtém os dados do formulário
        dados = request.form
        
        # Obtém o tipo de formulário (sempre tecnica neste caso)
        tipo_formulario = 'tecnica'
        
        # Processa as opções de formato selecionadas pelo usuário
        gerar_json = request.form.get('gerar_json', 'true').lower() == 'true'  # Por padrão, sempre gera JSON
        gerar_docx = request.form.get('gerar_docx', 'false').lower() == 'true'
        gerar_pdf = request.form.get('gerar_pdf', 'false').lower() == 'true'
        apenas_json = request.form.get('apenas_json', 'false').lower() == 'true'

        # Dados específicos do formulário técnico
        atividades = json.loads(dados.get('atividades', '[]'))
        total_horas = dados.get('totalHoras', 0)

        if apenas_json:
            try:
                # Criamos um dicionário com os dados disponíveis
                dados_json = {
                    'tipo': tipo_formulario,  # Adiciona o tipo
                    'info': {
                        'numeroSS': dados.get('numeroSS', ''),
                        'anoSS': dados.get('anoSS', ''), 
                        'tituloSS': dados.get('tituloSS', ''),
                        'descricao': dados.get('descricao', ''),
                        'dataInicio': dados.get('dataInicio', ''),
                        'dataFim': dados.get('dataFim', ''),
                        'linkBoard': dados.get('linkBoard', ''),
                        'iniciaisAutor': dados.get('iniciaisAutor', ''),
                        'totalHoras': total_horas
                    },
                    'atividades': atividades
                }
                
                # Caminho e nome do arquivo JSON
                json_filename = f"SS {dados.get('numeroSS', '').zfill(3)}-{dados.get('anoSS', '')}.json"
                json_path = os.path.join(tempfile.gettempdir(), json_filename)
                
                # Salvar arquivo JSON
                with open(json_path, 'w', encoding='utf-8') as json_file:
                    json.dump(dados_json, json_file, ensure_ascii=False, indent=2)
                
                # Retornar resposta JSON com confirmação
                return send_file(
                    json_path,
                    as_attachment=True,
                    download_name=json_filename,
                    mimetype='application/json'
                )
                
            except Exception as e:
                logger.error(f"Erro ao gerar arquivo JSON: {str(e)}")
                return jsonify({"error": f"Erro ao gerar JSON: {str(e)}"}), 500

        json_path = None
        try:
            # Criamos um dicionário com todos os dados
            dados_json = {
                'tipo': tipo_formulario,  # Adiciona o tipo
                'info': {
                    'numeroSS': dados['numeroSS'],
                    'anoSS': dados['anoSS'], 
                    'tituloSS': dados['tituloSS'],
                    'descricao': dados['descricao'],
                    'dataInicio': dados['dataInicio'],
                    'dataFim': dados['dataFim'],
                    'totalHoras': dados.get('totalHoras', 0),
                    'linkBoard': dados['linkBoard'],
                    'iniciaisAutor': dados['iniciaisAutor']
                },
                'atividades': atividades
            }
            
            # Caminho do arquivo JSON
            json_path = os.path.join(tempfile.gettempdir(), f"SS {dados['numeroSS'].zfill(3)}-{dados['anoSS']}.json")
            
            # Salvar arquivo JSON
            with open(json_path, 'w', encoding='utf-8') as json_file:
                json.dump(dados_json, json_file, ensure_ascii=False, indent=2)
            
            # Adicionar à lista de arquivos temporários e de saída
            temp_files.append(json_path)
            output_files.append({
                'path': json_path,
                'filename': os.path.basename(json_path)
            })
            
        except Exception as e:
            logger.error(f"Erro ao gerar arquivo JSON: {str(e)}")

        # Cria instância do relatório com os dados fornecidos
        relatorio = RelatorioAcompanhamentoProjeto(
            numero_ss=dados['numeroSS'],
            ano_ss=dados['anoSS'],
            iniciais_autor=dados['iniciaisAutor'],
            titulo_ss=dados['tituloSS'],
            descricao=dados['descricao'],
            data_inicio=dados['dataInicio'],
            data_fim=dados['dataFim'],
            total_horas=dados.get('totalHoras', 0),
            link_board=dados['linkBoard'],
            atividades=atividades
        )

        # Define o caminho dos modelos de documento
        modelos_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'modelos', 'tecnica')
        modelos = [
            os.path.join(modelos_path, "ModeloTec - Estimativa de Esforço e Cronograma.docx"),
            os.path.join(modelos_path, "ModeloTec - Estratégia de solução.docx"),
            os.path.join(modelos_path, "ModeloTec - Relatório de Acompanhamento de Projeto.docx")
        ]
        
        # Verifica se todos os modelos existem
        for modelo in modelos:
            if not os.path.exists(modelo):
                return {"error": f"Modelo não encontrado: {os.path.basename(modelo)}"}, 404

        # Obtém as substituições que serão feitas nos documentos
        substituicoes = relatorio.get_substituicoes()
        
        # Importa as funções necessárias do app.py
        from app import (
            identificar_tipo_documento, 
            processar_documento, 
            atualizar_sumario_com_python_docx, 
            obter_titulos_sumario, 
            obter_paginas_pdf, 
            gerar_pdf_do_docx
        )
        
        # Processa cada modelo de documento
        for modelo_path in modelos:
            # Cria uma cópia temporária do modelo
            temp_output = None
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_output = temp_file.name
                temp_files.append(temp_output)
                shutil.copy2(modelo_path, temp_output)
            
            try:
                # Processa o documento em um contexto controlado
                with open(temp_output, 'rb') as doc_file:
                    doc = Document(doc_file)
                    doc.core_properties.title = os.path.basename(modelo_path)
                    
                    # Identifica o tipo de documento e gera o nome do arquivo
                    tipo_documento = identificar_tipo_documento(doc)
                    nome_arquivo = gerar_nome_arquivo(
                        tipo_documento,
                        dados['numeroSS'].zfill(3),
                        dados['anoSS']
                    )

                    # Define o título do documento sem a extensão
                    titulo_documento = os.path.splitext(nome_arquivo)[0]
                    logger.debug(f"Título definido no DOCX: {titulo_documento}")
                    
                    # Define as propriedades do documento DOCX
                    doc.core_properties.title = titulo_documento

                    # Processa o documento substituindo os marcadores
                    processar_documento(doc, tipo_documento, substituicoes, atividades)
                    
                    # Define o caminho do arquivo temporário final
                    temp_final = os.path.join(tempfile.gettempdir(), nome_arquivo)

                    # Salvamos o documento com as substituições
                    doc.save(temp_final)
                    pdf_pre_sumario = gerar_pdf_do_docx(temp_final)
                    titulos_sumario = obter_titulos_sumario(temp_final)
                    titulos_atualizados = obter_paginas_pdf(pdf_pre_sumario, titulos_sumario)

                    temp_files.append(temp_final)
                    
                    # Atualizamos o sumário - SEMPRE, independente se vai gerar PDF ou não
                    atualizar_sumario_com_python_docx(temp_final, titulo_documento, titulos_atualizados)
                        
                    # Adiciona o DOCX aos arquivos de saída se solicitado
                    if gerar_docx:
                        output_files.append({
                            'path': temp_final,
                            'filename': nome_arquivo
                        })
                    
                    # Gera PDF se solicitado
                    if gerar_pdf:
                        pdf_path = gerar_pdf_do_docx(temp_final)
                        if pdf_path:
                            temp_files.append(pdf_path)
                            output_files.append({
                                'path': pdf_path,
                                'filename': os.path.splitext(nome_arquivo)[0] + '.pdf'
                            })
                        else:
                            # Se falhou ao gerar PDF mas o usuário quer só PDF, adiciona mensagem de erro
                            if not gerar_docx:
                                return {"error": "Falha ao gerar PDF"}, 500
                
            except Exception as e:
                logger.error(f"Erro ao processar documento {modelo_path}: {str(e)}")
                raise
        
        # Verifica se temos arquivos para incluir no ZIP
        if not output_files:
            return {"error": "Nenhum arquivo foi gerado com sucesso"}, 500
        
        # Cria nome do arquivo ZIP
        zip_filename = f"SS {dados['numeroSS'].zfill(3)}-{dados['anoSS']}.zip"
        zip_path = os.path.join(tempfile.gettempdir(), zip_filename)
        
        # Cria o arquivo ZIP com os documentos gerados
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for output in output_files:
                if os.path.exists(output['path']):
                    zipf.write(output['path'], output['filename'])

        # Configura e retorna a resposta com o arquivo ZIP
        response = send_file(
            zip_path,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )

        # Garante que o Content-Disposition esteja correto para caracteres especiais
        encoded_filename = quote(zip_filename)
        response.headers["Content-Disposition"] = (
            f"attachment; filename=\"{zip_filename}\"; "
            f"filename*=UTF-8''{encoded_filename}"
        )

        # Função para limpeza dos arquivos temporários após envio da resposta
        @response.call_on_close
        def cleanup():
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.unlink(temp_file)
                except Exception as e:
                    logger.error(f"Erro ao limpar arquivo temporário: {e}")
            
            try:
                if os.path.exists(zip_path):
                    os.unlink(zip_path)
            except Exception as e:
                logger.error(f"Erro ao limpar arquivo ZIP: {e}")

        return response

    except Exception as e:
        logger.exception("Erro durante a geração dos relatórios")
        return {"error": str(e)}, 500
        
    finally:
        # Limpeza final de segurança dos arquivos temporários
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except:
                pass
            
        try:
            if zip_path and os.path.exists(zip_path):
                os.unlink(zip_path)
        except:
            pass

def init_app_tec(app):
    """
    Função para inicializar as rotas específicas de documentos técnicos.
    """
    @app.route('/gerar_relatorio_tecnico', methods=['POST'])
    def gerar_relatorio_tecnico_route():
        return gerar_relatorio_tecnico()