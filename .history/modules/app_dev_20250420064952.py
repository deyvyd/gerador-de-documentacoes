# modules/app_dev.py - Funcionalidades para relatórios de desenvolvimento
import json
import logging
import os
import shutil
import tempfile
import zipfile
from urllib.parse import quote
from datetime import datetime

from flask import request, send_file, jsonify
from docx import Document

# Importações dos módulos compartilhados
from models.report_models import DocumentacaoTecnica
from modules.utils import gerar_nome_arquivo, obter_paginas_pdf
from modules.document_processor import (
    identificar_tipo_documento, 
    processar_documento, 
    processar_atividades_bullet_points,
    processar_atividades_tabela,
    obter_titulos_sumario,
    atualizar_sumario_com_python_docx,
    gerar_pdf_do_docx
)

# Configuração do logger
logger = logging.getLogger(__name__)

# Implementação para o arquivo app-dev.py
def gerar_documentos_dev():
    """
    Função para gerar relatórios de desenvolvimento.
    """
    logger.info("Iniciando geração de relatórios de desenvolvimento")

    # Verificar valor do tipo recebido
    tipo_formulario = request.form.get('tipo', 'não especificado')
    logger.info(f"Tipo de formulário recebido: {tipo_formulario}")
    
    # Listas para controle de arquivos temporários e arquivos de saída
    temp_files = []    # Arquivos temporários que serão deletados ao final
    output_files = []  # Arquivos que serão incluídos no ZIP final
    zip_path = None    # Caminho do arquivo ZIP que será gerado
    
    try:
        # Obtém os dados do formulário
        dados = request.form
        
        # Obtém o tipo de formulário (sempre desenvolvimento neste caso)
        tipo_formulario = 'desenvolvimento'
        
        # Processa as opções de formato selecionadas pelo usuário
        gerar_json = request.form.get('gerar_json', 'true').lower() == 'true'
        gerar_docx = request.form.get('gerar_docx', 'false').lower() == 'true'
        gerar_pdf = request.form.get('gerar_pdf', 'false').lower() == 'true'
        apenas_json = request.form.get('apenas_json', 'false').lower() == 'true'

        # Dados específicos do formulário de desenvolvimento
        requisitos = json.loads(dados.get('requisitos', '[]'))
        requisitos_nao_funcionais = json.loads(dados.get('requisitosNaoFuncionais', '[]'))
        total_pontos_funcao = dados.get('totalPontosFuncao', 0)

        # Se apenas JSON foi solicitado, retorna diretamente sem processar os documentos
        if apenas_json:
            try:
                # Criamos um dicionário com os dados disponíveis

                # Se o total de pontos de função for maior que zero, incluir campos adicionais
                if float(total_pontos_funcao) > 0:
                    dados_json = {
                        'tipo': tipo_formulario,
                        'info': {
                            'numeroSS': dados['numeroSS'],
                            'anoSS': dados['anoSS'], 
                            'tituloSS': dados['tituloSS'],
                            'descricao': dados['descricao'],
                            'dataInicio': dados['dataInicio'],
                            'dataFim': dados['dataFim'],
                            'linkBoard': dados['linkBoard'],
                            'iniciaisAutorCriacao': dados['iniciaisAutor'],
                            'iniciaisAutorModificacao': dados['iniciaisAutor'],
                            'dataCriacao': datetime.now().strftime('%d/%m/%Y'),
                            'dataMoficacao': datetime.now().strftime('%d/%m/%Y'),
                            'totalPontosFuncao': total_pontos_funcao
                        },
                        'requisitos': requisitos,
                        'requisitosNaoFuncionais': requisitos_nao_funcionais,
                        'totalPontosFuncao': total_pontos_funcao
                    }
                else:
                    dados_json = {
                        'tipo': tipo_formulario,
                        'info': {
                            'numeroSS': dados['numeroSS'],
                            'anoSS': dados['anoSS'], 
                            'tituloSS': dados['tituloSS'],
                            'descricao': dados['descricao'],
                            'dataInicio': dados['dataInicio'],
                            'dataFim': dados['dataFim'],
                            'linkBoard': dados['linkBoard'],
                            'iniciaisAutorCriacao': dados['iniciaisAutor'],
                            'dataCriacao': datetime.now().strftime('%d/%m/%Y')
                        },
                        'requisitos': requisitos,
                        'requisitosNaoFuncionais': requisitos_nao_funcionais,
                        'totalPontosFuncao': total_pontos_funcao
                    }
                
                # Caminho e nome do arquivo JSON
                json_filename = f"SS {dados['numeroSS'].zfill(3)}-{dados['anoSS']}.json"
                json_path = os.path.join(tempfile.gettempdir(), json_filename)
                
                # Salvar arquivo JSON
                with open(json_path, 'w', encoding='utf-8') as json_file:
                    json.dump(dados_json, json_file, ensure_ascii=False, indent=2)
                
                # Retornar resposta JSON com confirmação
                return send_file(
                    json_path,
                    as_attachment=True,
                    download_name=json_filename,
                    mimetype='application/json'
                )
                
            except Exception as e:
                logger.error(f"Erro ao gerar arquivo JSON: {str(e)}")
                return jsonify({"error": f"Erro ao gerar JSON: {str(e)}"}), 500

        json_path = None
        try:
            # Criamos um dicionário com todos os dados

            # Se o total de pontos de função for maior que zero, incluir campos adicionais
            if float(total_pontos_funcao) > 0:
                dados_json = {
                    'tipo': tipo_formulario,
                    'info': {
                        'numeroSS': dados['numeroSS'],
                        'anoSS': dados['anoSS'], 
                        'tituloSS': dados['tituloSS'],
                        'descricao': dados['descricao'],
                        'dataInicio': dados['dataInicio'],
                        'dataFim': dados['dataFim'],
                        'linkBoard': dados['linkBoard'],
                        'iniciaisAutor': dados['iniciaisAutor'],
                        'dataCriacao': datetime.now().strftime('%d/%m/%Y'),
                        'dataMoficacao': datetime.now().strftime('%d/%m/%Y'),
                        'iniciaisAutor2': dados['iniciaisAutor'],
                        'totalPontosFuncao': total_pontos_funcao
                    },
                    'requisitos': requisitos,
                    'requisitosNaoFuncionais': requisitos_nao_funcionais,
                    'totalPontosFuncao': total_pontos_funcao
                }
            else:
                dados_json = {
                    'tipo': tipo_formulario,
                    'info': {
                        'numeroSS': dados['numeroSS'],
                        'anoSS': dados['anoSS'], 
                        'tituloSS': dados['tituloSS'],
                        'descricao': dados['descricao'],
                        'dataInicio': dados['dataInicio'],
                        'dataFim': dados['dataFim'],
                        'linkBoard': dados['linkBoard'],
                        'iniciaisAutor': dados['iniciaisAutor'],
                        'dataCriacao': datetime.now().strftime('%d/%m/%Y')
                    },
                    'requisitos': requisitos,
                    'requisitosNaoFuncionais': requisitos_nao_funcionais,
                    'totalPontosFuncao': total_pontos_funcao
                }
            
            # Caminho do arquivo JSON
            json_path = os.path.join(tempfile.gettempdir(), f"SS {dados['numeroSS'].zfill(3)}-{dados['anoSS']}.json")
            
            # Salvar arquivo JSON
            with open(json_path, 'w', encoding='utf-8') as json_file:
                json.dump(dados_json, json_file, ensure_ascii=False, indent=2)
            
            # Adicionar à lista de arquivos temporários e de saída
            temp_files.append(json_path)
            output_files.append({
                'path': json_path,
                'filename': os.path.basename(json_path)
            })
            
        except Exception as e:
            logger.error(f"Erro ao gerar arquivo JSON: {str(e)}")
        
        class DocumentacaoDesenvolvimento(DocumentacaoTecnica):
            """Extensão da classe DocumentacaoTecnica com campos específicos de desenvolvimento"""
            def __init__(self, numero_ss, ano_ss, iniciais_autor, titulo_ss, descricao, 
                        data_inicio, data_fim, link_board, requisitos=None, 
                        requisitos_nao_funcionais=None, total_pontos_funcao=0,
                        data_modificacao=None, iniciais_autor_modificacao=None):
                
                # Chama o construtor da classe pai
                super().__init__(numero_ss, ano_ss, iniciais_autor, titulo_ss, descricao,
                                data_inicio, data_fim, 0, link_board)  # Total horas 0
                
                # Campos específicos de desenvolvimento
                self.requisitos = requisitos or []
                self.requisitos_nao_funcionais = requisitos_nao_funcionais or []
                self.total_pontos_funcao = float(total_pontos_funcao)

                # Campos adicionais
                self.data_modificacao = data_modificacao
                self.iniciais_autor_modificacao = iniciais_autor_modificacao
                
            def get_substituicoes(self):
                """Retorna o dicionário de substituições para o documento"""
                # Obtém as substituições básicas da classe pai
                substituicoes = super().get_substituicoes()
                
                # Verifica se temos valores para DATA_ATUAL2 e INICIAIS_AUTOR2
                # Podemos ter esses valores por dois motivos:
                # 1. O total de pontos de função é maior que zero (preenchido agora)
                # 2. Esses valores foram carregados do JSON importado
                
                if hasattr(self, 'data_modificacao') and self.data_modificacao:
                    substituicoes['[DATA_ATUAL2]'] = self.data_modificacao
                elif self.total_pontos_funcao > 0:
                    from datetime import datetime
                    substituicoes['[DATA_ATUAL2]'] = datetime.now().strftime('%d/%m/%Y')
                
                if hasattr(self, 'iniciais_autor_modificacao') and self.iniciais_autor_modificacao:
                    substituicoes['[INICIAIS_AUTOR2]'] = self.iniciais_autor_modificacao
                elif self.total_pontos_funcao > 0:
                    substituicoes['[INICIAIS_AUTOR2]'] = substituicoes['[INICIAIS_AUTOR]']
                
                # Adiciona o número de pontos de função formatado
                substituicoes['[N_PF]'] = f"{self.total_pontos_funcao:.2f}".replace('.', ',')
                
                return substituicoes
                
        # Cria a instância de documentação de desenvolvimento
        documentacao = DocumentacaoDesenvolvimento(
            numero_ss=dados['numeroSS'],
            ano_ss=dados['anoSS'],
            iniciais_autor=dados['iniciaisAutor'],
            titulo_ss=dados['tituloSS'],
            descricao=dados['descricao'],
            data_inicio=dados['dataInicio'],
            data_fim=dados['dataFim'],
            link_board=dados['linkBoard'],
            requisitos=requisitos,
            requisitos_nao_funcionais=requisitos_nao_funcionais,
            total_pontos_funcao=total_pontos_funcao,
            data_modificacao=dados.get('dataModificacao'),
            iniciais_autor_modificacao=dados.get('iniciaisAutorModificacao')
        )

        # Define o caminho dos modelos de documento
        app_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        modelos_path = os.path.join(app_root, 'modelos', 'desenvolvimento')
        modelos = [
            os.path.join(modelos_path, "ModeloDev - Estimativa de Esforço e Cronograma.docx"),
            os.path.join(modelos_path, "ModeloDev - Estratégia de solução.docx"),
            os.path.join(modelos_path, "ModeloDev - Relatório de Acompanhamento de Projeto.docx")
        ]
        
        # Verifica se todos os modelos existem
        for modelo in modelos:
            if os.path.exists(modelo):
                logger.info(f"Modelo encontrado: {modelo}")
            else:
                logger.error(f"Modelo não encontrado: {modelo}")
                return {"error": f"Modelo não encontrado: {os.path.basename(modelo)}"}, 404

        # Obtém as substituições que serão feitas nos documentos
        substituicoes = documentacao.get_substituicoes()
        
        # Processa cada modelo de documento
        for modelo_path in modelos:
            # Cria uma cópia temporária do modelo
            temp_output = None
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_output = temp_file.name
                temp_files.append(temp_output)
                shutil.copy2(modelo_path, temp_output)
            
            try:
                # Processa o documento em um contexto controlado
                with open(temp_output, 'rb') as doc_file:
                    doc = Document(doc_file)
                    doc.core_properties.title = os.path.basename(modelo_path)
                    
                    # Identifica o tipo de documento e gera o nome do arquivo
                    tipo_documento = identificar_tipo_documento(doc)
                    nome_arquivo = gerar_nome_arquivo(
                        tipo_documento,
                        dados['numeroSS'].zfill(3),
                        dados['anoSS']
                    )

                    # Define o título do documento sem a extensão
                    titulo_documento = os.path.splitext(nome_arquivo)[0]
                    logger.debug(f"Título definido no DOCX: {titulo_documento}")
                    
                    # Define as propriedades do documento DOCX
                    doc.core_properties.title = titulo_documento

                    # Processa o documento substituindo os marcadores
                    processar_documento(doc, tipo_documento, substituicoes)
                    
                    # Processa os requisitos funcionais
                    processar_requisitos_funcionais(doc, documentacao.requisitos)
                    
                    # Processa os requisitos não funcionais
                    processar_requisitos_nao_funcionais(doc, documentacao.requisitos_nao_funcionais)
                    
                    # Define o caminho do arquivo temporário final
                    temp_final = os.path.join(tempfile.gettempdir(), nome_arquivo)

                    # Salvamos o documento com as substituições
                    doc.save(temp_final)
                    pdf_pre_sumario = gerar_pdf_do_docx(temp_final)
                    titulos_sumario = obter_titulos_sumario(temp_final)
                    titulos_atualizados = obter_paginas_pdf(pdf_pre_sumario, titulos_sumario)

                    temp_files.append(temp_final)
                    
                    # Atualizamos o sumário - SEMPRE, independente se vai gerar PDF ou não
                    atualizar_sumario_com_python_docx(temp_final, titulo_documento, titulos_atualizados)
                        
                    # Adiciona o DOCX aos arquivos de saída se solicitado
                    if gerar_docx:
                        output_files.append({
                            'path': temp_final,
                            'filename': nome_arquivo
                        })
                    
                    # Gera PDF se solicitado
                    if gerar_pdf:
                        pdf_path = gerar_pdf_do_docx(temp_final)
                        if pdf_path:
                            temp_files.append(pdf_path)
                            output_files.append({
                                'path': pdf_path,
                                'filename': os.path.splitext(nome_arquivo)[0] + '.pdf'
                            })
                        else:
                            # Se falhou ao gerar PDF mas o usuário quer só PDF, adiciona mensagem de erro
                            if not gerar_docx:
                                return {"error": "Falha ao gerar PDF"}, 500
                
            except Exception as e:
                logger.error(f"Erro ao processar documento {modelo_path}: {str(e)}")
                raise
        
        # Verifica se temos arquivos para incluir no ZIP
        if not output_files:
            return {"error": "Nenhum arquivo foi gerado com sucesso"}, 500
        
        # Cria nome do arquivo ZIP
        zip_filename = f"SS {dados['numeroSS'].zfill(3)}-{dados['anoSS']}.zip"
        zip_path = os.path.join(tempfile.gettempdir(), zip_filename)
        
        # Cria o arquivo ZIP com os documentos gerados
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for output in output_files:
                if os.path.exists(output['path']):
                    zipf.write(output['path'], output['filename'])

        # Configura e retorna a resposta com o arquivo ZIP
        response = send_file(
            zip_path,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )

        # Garante que o Content-Disposition esteja correto para caracteres especiais
        encoded_filename = quote(zip_filename)
        response.headers["Content-Disposition"] = (
            f"attachment; filename=\"{zip_filename}\"; "
            f"filename*=UTF-8''{encoded_filename}"
        )

        # Função para limpeza dos arquivos temporários após envio da resposta
        @response.call_on_close
        def cleanup():
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.unlink(temp_file)
                except Exception as e:
                    logger.error(f"Erro ao limpar arquivo temporário: {e}")
            
            try:
                if os.path.exists(zip_path):
                    os.unlink(zip_path)
            except Exception as e:
                logger.error(f"Erro ao limpar arquivo ZIP: {e}")

        return response

    except Exception as e:
        logger.exception("Erro durante a geração dos relatórios de desenvolvimento")
        return {"error": str(e)}, 500
        
    finally:
        # Limpeza final de segurança dos arquivos temporários
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except:
                pass
            
        try:
            if zip_path and os.path.exists(zip_path):
                os.unlink(zip_path)
        except:
            pass


def processar_requisitos_funcionais(doc, requisitos):
    """
    Processa as tabelas de requisitos funcionais no documento.
    Substitui os marcadores [N_RF], [TITULO_RF], etc.
    """
    if not requisitos:
        logger.info("Nenhum requisito funcional para processar")
        return
    
    logger.info(f"Processando {len(requisitos)} requisitos funcionais")
    
    # Identificar tabelas de requisitos funcionais
    tabelas_rf = []
    
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Verifica se o parágrafo contém algum marcador de requisito funcional
                    if any(marcador in paragraph.text for marcador in 
                          ['[N_RF]', '[TITULO_RF]', '[TIPO_RF]', '[LOCAL_RF]', 
                           '[USUARIO_RF]', '[PERFIL_RF]', '[DESCRICAO_RF]',
                           '[REGRAS_CAMPOS_RF]', '[REGRAS_NEGOCIO_RF]', '[BANCO_RF]']):
                        if i not in tabelas_rf:
                            tabelas_rf.append(i)
                            logger.info(f"Tabela de RF encontrada: índice {i}")
    
    # Processar cada tabela de requisitos funcionais
    for tabela_idx in tabelas_rf:
        tabela = doc.tables[tabela_idx]
        
        # Encontrar a linha de modelo (com marcadores)
        linha_modelo_idx = None
        for i, row in enumerate(tabela.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '[N_RF]' in paragraph.text:
                        linha_modelo_idx = i
                        logger.info(f"Linha modelo encontrada na posição {i}")
                        break
                if linha_modelo_idx is not None:
                    break
            if linha_modelo_idx is not None:
                break
        
        if linha_modelo_idx is None:
            logger.warning(f"Não foi encontrada linha modelo na tabela {tabela_idx}")
            continue
            
        # Guardar a linha modelo antes de modificá-la
        linha_modelo = tabela.rows[linha_modelo_idx]
        
        # Processar cada requisito funcional
        for idx, requisito in enumerate(requisitos):
            if idx == 0:
                # Para o primeiro requisito, usamos a linha modelo
                linha_atual = linha_modelo
            else:
                # Para os demais, adicionamos uma nova linha duplicando a modelo
                linha_atual = tabela.add_row()
                
                # Copiar o estilo de cada célula
                for cell_idx, cell in enumerate(linha_modelo.cells):
                    # Copiar formato da célula
                    if cell_idx < len(linha_atual.cells):
                        tab_copiar_estilo_celula(linha_atual.cells[cell_idx], cell)
            
            # Substitui os marcadores nas células
            for idx_cell, cell in enumerate(linha_atual.cells):
                for paragraph in cell.paragraphs:
                    texto_original = paragraph.text
                    
                    # Dicionário com as substituições específicas para cada requisito
                    substituicoes_rf = {
                        '[N_RF]': f"{idx + 1:02d}",
                        '[TITULO_RF]': requisito.get('tituloRF', ''),
                        '[TIPO_RF]': requisito.get('tipo', ''),
                        '[LOCAL_RF]': requisito.get('local', ''),
                        '[USUARIO_RF]': requisito.get('usuario', ''),
                        '[PERFIL_RF]': requisito.get('perfil', ''),
                        '[DESCRICAO_RF]': requisito.get('descricao', ''),
                        '[REGRAS_CAMPOS_RF]': requisito.get('validacoes', ''),
                        '[REGRAS_NEGOCIO_RF]': requisito.get('regras', ''),
                        '[BANCO_RF]': requisito.get('banco', '')
                    }
                    
                    # Verifica se há marcadores de RF neste parágrafo
                    if any(marcador in texto_original for marcador in substituicoes_rf.keys()):
                        # Obter o mapeamento de formatação do parágrafo
                        mapa_formatacao = tab_mapear_formatacao_paragrafo(paragraph)
                        
                        # Substituir cada marcador mantendo a formatação
                        texto_substituido = texto_original
                        for marcador, valor in substituicoes_rf.items():
                            if marcador in texto_substituido:
                                texto_substituido = texto_substituido.replace(marcador, valor)
                        
                        # Aplicar o texto substituído com a formatação original
                        tab_processar_formatacao_paragrafo(paragraph, texto_substituido, mapa_formatacao)
                    
                    # Tratamento especial para imagens - processa aqui se o documento tiver suporte
                    if '[IMAGEM_RF]' in texto_original and requisito.get('imagens'):
                        # Esta implementação requer integração com manipulação de imagens
                        # que seria adicionada posteriormente
                        logger.info(f"Imagens encontradas para o requisito {idx+1}, mas processamento não implementado")
                        
                        # Removemos o marcador para não deixá-lo no documento
                        tab_processar_formatacao_paragrafo(paragraph, "", mapa_formatacao)


def processar_requisitos_nao_funcionais(doc, requisitos_nao_funcionais):
    """
    Processa as tabelas de requisitos não funcionais no documento.
    Substitui os marcadores [N_RNF], [TITULO_RNF], [DESCRICAO_RNF].
    """
    if not requisitos_nao_funcionais:
        logger.info("Nenhum requisito não funcional para processar")
        return
    
    logger.info(f"Processando {len(requisitos_nao_funcionais)} requisitos não funcionais")
    
    # Identificar tabelas de requisitos não funcionais
    tabelas_rnf = []
    
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Verifica se o parágrafo contém algum marcador de requisito não funcional
                    if any(marcador in paragraph.text for marcador in 
                          ['[N_RNF]', '[TITULO_RNF]', '[DESCRICAO_RNF]']):
                        if i not in tabelas_rnf:
                            tabelas_rnf.append(i)
                            logger.info(f"Tabela de RNF encontrada: índice {i}")
    
    # Processar cada tabela de requisitos não funcionais
    for tabela_idx in tabelas_rnf:
        tabela = doc.tables[tabela_idx]
        
        # Encontrar a linha de modelo (com marcadores)
        linha_modelo_idx = None
        for i, row in enumerate(tabela.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '[N_RNF]' in paragraph.text:
                        linha_modelo_idx = i
                        logger.info(f"Linha modelo encontrada na posição {i}")
                        break
                if linha_modelo_idx is not None:
                    break
            if linha_modelo_idx is not None:
                break
        
        if linha_modelo_idx is None:
            logger.warning(f"Não foi encontrada linha modelo na tabela {tabela_idx}")
            continue
            
        # Guardar a linha modelo antes de modificá-la
        linha_modelo = tabela.rows[linha_modelo_idx]
        
        # Processar cada requisito não funcional
        for idx, requisito in enumerate(requisitos_nao_funcionais):
            if idx == 0:
                # Para o primeiro requisito, usamos a linha modelo
                linha_atual = linha_modelo
            else:
                # Para os demais, adicionamos uma nova linha duplicando a modelo
                linha_atual = tabela.add_row()
                
                # Copiar o estilo de cada célula
                for cell_idx, cell in enumerate(linha_modelo.cells):
                    # Copiar formato da célula
                    if cell_idx < len(linha_atual.cells):
                        tab_copiar_estilo_celula(linha_atual.cells[cell_idx], cell)
            
            # Substitui os marcadores nas células
            for idx_cell, cell in enumerate(linha_atual.cells):
                for paragraph in cell.paragraphs:
                    texto_original = paragraph.text
                    
                    # Dicionário com as substituições específicas para cada requisito
                    substituicoes_rnf = {
                        '[N_RNF]': f"{idx + 1:02d}",
                        '[TITULO_RNF]': requisito.get('titulo', ''),
                        '[DESCRICAO_RNF]': requisito.get('descricao', '')
                    }
                    
                    # Verifica se há marcadores de RNF neste parágrafo
                    if any(marcador in texto_original for marcador in substituicoes_rnf.keys()):
                        # Obter o mapeamento de formatação do parágrafo
                        mapa_formatacao = tab_mapear_formatacao_paragrafo(paragraph)
                        
                        # Substituir cada marcador mantendo a formatação
                        texto_substituido = texto_original
                        for marcador, valor in substituicoes_rnf.items():
                            if marcador in texto_substituido:
                                texto_substituido = texto_substituido.replace(marcador, valor)
                        
                        # Aplicar o texto substituído com a formatação original
                        tab_processar_formatacao_paragrafo(paragraph, texto_substituido, mapa_formatacao)

def init_app_dev(app):
    """
    Função para inicializar funções específicas de relatórios de desenvolvimento.
    Não registra rotas adicionais.
    """
    # Pode adicionar outras inicializações aqui, mas sem registrar rotas
