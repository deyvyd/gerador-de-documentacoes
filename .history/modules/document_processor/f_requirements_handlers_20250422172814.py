# f_requirements_handlers.py - Solução corrigida para formatação HTML e quebras de linha
import logging
import os
import base64
import re
import html
from io import BytesIO
from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from copy import deepcopy
from bs4 import BeautifulSoup

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    encoding='utf-8'  # Adicionar encoding UTF-8
)
logger = logging.getLogger(__name__)

def processar_requisitos_funcionais(doc, requisitos):
    """
    Processa requisitos funcionais no documento de Estratégia de Solução.
    
    Esta função localiza o parágrafo com o marcador [DESCRICAO] e o substitui,
    depois localiza o modelo de RF (1 - [TITULO]) e cria uma nova página para cada RF.
    
    Args:
        doc: Documento Word (.docx) aberto
        requisitos: Lista de requisitos funcionais para processar
    """
    if not requisitos or len(requisitos) == 0:
        logger.warning("Nenhum requisito funcional para processar")
        return
    
    logger.info(f"Processando {len(requisitos)} requisitos funcionais")
    
    # 1. Localizar título da seção de requisitos e a descrição
    paragrafo_descricao = None
    paragrafo_descricao_indice = None
    
    for i, p in enumerate(doc.paragraphs):
        if "[DESCRICAO]" in p.text:
            paragrafo_descricao = p
            paragrafo_descricao_indice = i
            logger.info(f"Encontrado parágrafo de descrição no índice {i}: {p.text}")
            break
    
    if not paragrafo_descricao:
        logger.error("Não foi possível encontrar o parágrafo com [DESCRICAO]")
        return
    
    # Substituir o marcador [DESCRICAO] pela descrição geral
    paragrafo_descricao.text = paragrafo_descricao.text.replace("[DESCRICAO]", "Abaixo segue a descrição dos requisitos funcionais.")
    logger.info("Substituído marcador [DESCRICAO] pela descrição geral")
    
    # 2. Localizar parágrafo modelo para o título do RF (1 - [TITULO])
    paragrafo_modelo = None
    paragrafo_modelo_indice = None
    
    for i, p in enumerate(doc.paragraphs):
        if "1 - [TITULO]" in p.text or "1 – [TITULO]" in p.text:
            paragrafo_modelo = p
            paragrafo_modelo_indice = i
            logger.info(f"Encontrado parágrafo modelo de título no índice {i}: {p.text}")
            break
    
    if not paragrafo_modelo:
        logger.error("Não foi possível encontrar o parágrafo modelo com '1 - [TITULO]'")
        return
    
    # 3. Localizar parágrafo com a palavra "Descrição" (após o título)
    paragrafo_descricao_titulo = None
    paragrafo_descricao_titulo_indice = None
    
    # Começar a busca a partir do parágrafo do título do RF
    for i in range(paragrafo_modelo_indice + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == "Descrição":
            paragrafo_descricao_titulo = doc.paragraphs[i]
            paragrafo_descricao_titulo_indice = i
            logger.info(f"Encontrado parágrafo 'Descrição' no índice {i}")
            break
    
    # 4. Localizar o parágrafo com [DESCRICAO]
    paragrafo_descricao_conteudo = None
    paragrafo_descricao_conteudo_indice = None
    
    # Começar a busca a partir do parágrafo "Descrição"
    if paragrafo_descricao_titulo_indice:
        for i in range(paragrafo_descricao_titulo_indice + 1, len(doc.paragraphs)):
            if "[DESCRICAO]" in doc.paragraphs[i].text:
                paragrafo_descricao_conteudo = doc.paragraphs[i]
                paragrafo_descricao_conteudo_indice = i
                logger.info(f"Encontrado parágrafo [DESCRICAO] no índice {i}")
                break
    
    # 5. Localizar tabela modelo que contém marcadores de RF
    tabela_modelo = None
    tabela_indice = None
    
    for i, tabela in enumerate(doc.tables):
        contem_marcador = False
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    if any(marcador in p.text for marcador in ["[LOCAL_RF]", "[USUARIO_RF]", "[TIPO_RF]"]):
                        contem_marcador = True
                        logger.info(f"Encontrado marcador na tabela {i}")
                        break
                if contem_marcador:
                    break
            if contem_marcador:
                break
        
        if contem_marcador:
            tabela_modelo = tabela
            tabela_indice = i
            logger.info(f"Encontrada tabela modelo no índice {i}")
            break
    
    if not tabela_modelo:
        logger.error("Não foi possível encontrar a tabela modelo com marcadores RF")
        return
    
    # 6. Armazenar o corpo do documento para inserções
    body = doc._body._body
    
    # 7. Para cada requisito, criar uma cópia do parágrafo e da tabela em uma nova página
    elementos_adicionados = []
    
    # Ponto de inserção após a descrição geral
    posicao_insercao = paragrafo_descricao_indice + 1
    
    # Para cada requisito, criar uma nova página com o título do RF e a tabela
    for idx, requisito in enumerate(requisitos):
        logger.info(f"Processando requisito {idx+1}: {requisito.get('tituloRF', 'Sem título')}")
        
        # Adicionar quebra de página antes de cada requisito
        quebra_pagina = doc.add_paragraph()
        run = quebra_pagina.add_run()
        run.add_break(WD_BREAK.PAGE)
        body.insert(posicao_insercao, quebra_pagina._p)
        posicao_insercao += 1
        elementos_adicionados.append(quebra_pagina._p)
        
        # Criar novo parágrafo para título
        novo_titulo = doc.add_paragraph()
        
        # Copiar estilo e formatação do parágrafo modelo
        if hasattr(paragrafo_modelo, 'style') and paragrafo_modelo.style:
            novo_titulo.style = paragrafo_modelo.style
            
        if hasattr(paragrafo_modelo, 'paragraph_format'):
            if hasattr(paragrafo_modelo.paragraph_format, 'alignment'):
                novo_titulo.paragraph_format.alignment = paragrafo_modelo.paragraph_format.alignment
            if hasattr(paragrafo_modelo.paragraph_format, 'left_indent'):
                novo_titulo.paragraph_format.left_indent = paragrafo_modelo.paragraph_format.left_indent
            if hasattr(paragrafo_modelo.paragraph_format, 'space_before'):
                novo_titulo.paragraph_format.space_before = paragrafo_modelo.paragraph_format.space_before
            if hasattr(paragrafo_modelo.paragraph_format, 'space_after'):
                novo_titulo.paragraph_format.space_after = paragrafo_modelo.paragraph_format.space_after
        
        # Formatar o título do RF no formato "1 - Nome"
        titulo = requisito.get('tituloRF', '[Sem título]')
        run = novo_titulo.add_run(f"{idx + 1} - {titulo}")
        
        # Copiar formatação de fonte do run original
        if paragrafo_modelo.runs:
            run_modelo = paragrafo_modelo.runs[0]
            if hasattr(run_modelo, 'font'):
                if hasattr(run_modelo.font, 'name'):
                    run.font.name = run_modelo.font.name
                if hasattr(run_modelo.font, 'size'):
                    run.font.size = run_modelo.font.size
                if hasattr(run_modelo.font, 'bold'):
                    run.font.bold = run_modelo.font.bold
                if hasattr(run_modelo.font, 'italic'):
                    run.font.italic = run_modelo.font.italic
                if hasattr(run_modelo.font, 'color') and run_modelo.font.color and run_modelo.font.color.rgb:
                    run.font.color.rgb = run_modelo.font.color.rgb
        
        # Inserir o parágrafo de título no documento
        body.insert(posicao_insercao, novo_titulo._p)
        posicao_insercao += 1
        elementos_adicionados.append(novo_titulo._p)
        
        # Adicionar parágrafo "Descrição"
        if paragrafo_descricao_titulo:
            # Criar uma cópia do parágrafo "Descrição"
            novo_descricao_titulo = doc.add_paragraph()
            
            # Copiar estilo e formatação
            if hasattr(paragrafo_descricao_titulo, 'style') and paragrafo_descricao_titulo.style:
                novo_descricao_titulo.style = paragrafo_descricao_titulo.style
            
            # Adicionar o texto "Descrição"
            run = novo_descricao_titulo.add_run("Descrição")
            
            # Copiar formatação de fonte
            if paragrafo_descricao_titulo.runs:
                run_modelo = paragrafo_descricao_titulo.runs[0]
                if hasattr(run_modelo, 'font'):
                    if hasattr(run_modelo.font, 'name'):
                        run.font.name = run_modelo.font.name
                    if hasattr(run_modelo.font, 'size'):
                        run.font.size = run_modelo.font.size
                    if hasattr(run_modelo.font, 'bold'):
                        run.font.bold = run_modelo.font.bold
            
            # Inserir o parágrafo no documento
            body.insert(posicao_insercao, novo_descricao_titulo._p)
            posicao_insercao += 1
            elementos_adicionados.append(novo_descricao_titulo._p)
        
        # Obter valores dos campos do requisito
        local = requisito.get('local', '')
        usuario = requisito.get('usuario', '')
        perfil = requisito.get('perfil', '')
        tipo = requisito.get('tipo', '')
        
        # Obter conteúdo HTML dos campos formatados
        descricao_html = requisito.get('descricao', '')
        if not descricao_html:
            descricao_html = ""
            
        validacoes_html = requisito.get('validacoes', '')
        if not validacoes_html:
            validacoes_html = "Não há validação de campos."
            
        regras_html = requisito.get('regras', '')
        if not regras_html:
            regras_html = "Não há alterações de regra de negócio."
            
        banco_html = requisito.get('banco', '')
        if not banco_html:
            banco_html = "Não há alterações de banco de dados."
        
        # Criar uma cópia da tabela modelo via XML
        nova_tabela_xml = deepcopy(tabela_modelo._tbl)
        
        # Substituir marcadores de texto simples
        substituicoes = {
            "[LOCAL_RF]": local,
            "[USUARIO_RF]": usuario,
            "[PERFIL_RF]": perfil,
            "[TIPO_RF]": tipo,
        }
        
        substituir_marcadores_tabela(nova_tabela_xml, substituicoes)
        
        # Inserir a nova tabela no documento
        body.insert(posicao_insercao, nova_tabela_xml)
        posicao_insercao += 1
        elementos_adicionados.append(nova_tabela_xml)
        
        # Processar o conteúdo formatado (HTML) na tabela inserida
        tabela_inserida = None
        for i, tabela in enumerate(doc.tables):
            contem_marcador = False
            for linha in tabela.rows:
                for celula in linha.cells:
                    for p in celula.paragraphs:
                        if any(marcador in p.text for marcador in ["[DESCRICAO_RF]", "[REGRAS_CAMPOS_RF]", "[REGRAS_NEGOCIO_RF]", "[BANCO_RF]", "[IMAGEM_RF]"]):
                            contem_marcador = True
                            break
                    if contem_marcador:
                        break
                if contem_marcador:
                    break
            
            if contem_marcador:
                tabela_inserida = tabela
                logger.info(f"Encontrada tabela inserida no índice {i}")
                break
        
        if tabela_inserida:
            # Usar processamento HTML para campos formatados
            substituir_html_com_formatacao(tabela_inserida, "[DESCRICAO_RF]", descricao_html)
            substituir_html_com_formatacao(tabela_inserida, "[REGRAS_CAMPOS_RF]", validacoes_html)
            substituir_html_com_formatacao(tabela_inserida, "[REGRAS_NEGOCIO_RF]", regras_html)
            substituir_html_com_formatacao(tabela_inserida, "[BANCO_RF]", banco_html)
            
            # Processar imagens (se houver)
            imagens = requisito.get('imagens', [])
            if imagens:
                logger.info(f"Processando {len(imagens)} imagens para o requisito {idx+1}")
                
                # Localizar célula com o marcador [IMAGEM_RF]
                celula_imagem = None
                for i, linha in enumerate(tabela_inserida.rows):
                    for j, celula in enumerate(linha.cells):
                        for p in celula.paragraphs:
                            if "[IMAGEM_RF]" in p.text:
                                celula_imagem = celula
                                logger.info(f"Encontrado marcador [IMAGEM_RF] na linha {i}, coluna {j}")
                                break
                        if celula_imagem:
                            break
                    if celula_imagem:
                        break
                
                if celula_imagem:
                    # Processar imagens no lugar do marcador
                    for p in celula_imagem.paragraphs:
                        if "[IMAGEM_RF]" in p.text:
                            p.clear()
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            processar_imagens(p, imagens)
                            logger.info("Imagens adicionadas com sucesso")
                            break
                else:
                    logger.warning(f"Marcador [IMAGEM_RF] não encontrado para o requisito {idx+1}")
            else:
                # Remover marcador se não houver imagens
                for linha in tabela_inserida.rows:
                    for celula in linha.cells:
                        for p in celula.paragraphs:
                            if "[IMAGEM_RF]" in p.text:
                                p.text = p.text.replace("[IMAGEM_RF]", "")
        else:
            logger.warning(f"Tabela inserida não encontrada para requisito {idx+1}")
    
    # 8. Remover elementos do modelo original
    try:
        # Remover parágrafo modelo do título do RF
        if paragrafo_modelo._p.getparent() is not None:
            body.remove(paragrafo_modelo._p)
            logger.info("Removido parágrafo modelo do título")
        
        # Remover parágrafo "Descrição" original
        if paragrafo_descricao_titulo and paragrafo_descricao_titulo._p.getparent() is not None:
            body.remove(paragrafo_descricao_titulo._p)
            logger.info("Removido parágrafo 'Descrição' original")
        
        # Remover parágrafo de conteúdo [DESCRICAO] original
        if paragrafo_descricao_conteudo and paragrafo_descricao_conteudo._p.getparent() is not None:
            body.remove(paragrafo_descricao_conteudo._p)
            logger.info("Removido parágrafo [DESCRICAO] original")
        
        # Remover tabela modelo original
        if tabela_modelo._tbl.getparent() is not None:
            body.remove(tabela_modelo._tbl)
            logger.info("Removida tabela modelo original")
    except Exception as e:
        logger.error(f"Erro ao remover elementos do modelo: {str(e)}")
    
    logger.info(f"Processamento de {len(requisitos)} requisitos concluído com sucesso")

def substituir_marcadores_tabela(tabela_xml, substituicoes):
    """
    Substitui marcadores em uma tabela XML.
    
    Args:
        tabela_xml: Elemento XML da tabela
        substituicoes: Dicionário com os valores de substituição para cada marcador
    """
    # Namespace XML para Word
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    
    # Caso especial para [USUARIO_RF] | [PERFIL_RF]
    for elemento_t in tabela_xml.findall(".//w:t", namespace):
        texto = elemento_t.text
        
        if texto:
            # Caso especial para [USUARIO_RF] | [PERFIL_RF]
            if "[USUARIO_RF]" in texto and "[PERFIL_RF]" in texto:
                texto_substituido = texto.replace("[USUARIO_RF]", substituicoes["[USUARIO_RF]"])
                texto_substituido = texto_substituido.replace("[PERFIL_RF]", substituicoes["[PERFIL_RF]"])
                elemento_t.text = texto_substituido
            else:
                # Substituições normais (exceto os que precisam de formatação especial)
                marcadores_especiais = ["[IMAGEM_RF]", "[DESCRICAO_RF]", "[REGRAS_CAMPOS_RF]", "[REGRAS_NEGOCIO_RF]", "[BANCO_RF]"]
                for marcador, valor in substituicoes.items():
                    if marcador in texto and marcador not in marcadores_especiais:
                        elemento_t.text = texto.replace(marcador, valor)

def limpar_texto_html(html_texto):
    """
    Remove tags HTML e retorna o texto puro.
    Usada como uma função de fallback para substituição simples.
    
    Args:
        html_texto: Texto HTML do Quill Editor
        
    Returns:
        Texto limpo sem tags HTML
    """
    if not html_texto:
        return ""
    
    # Usar BeautifulSoup para limpar o HTML
    soup = BeautifulSoup(html_texto, 'html.parser')
    
    # Obter o texto limpo
    return soup.get_text()

def substituir_html_com_formatacao(tabela, marcador, html_texto):
    """
    Substitui um marcador por texto HTML formatado corretamente.
    
    Args:
        tabela: Tabela onde buscar o marcador
        marcador: Marcador a ser substituído (ex: "[DESCRICAO_RF]")
        html_texto: Texto HTML para substituir o marcador
    """
    if not html_texto:
        return
    
    # Localizar a célula que contém o marcador
    celula_destino = None
    for linha in tabela.rows:
        for celula in linha.cells:
            for p in celula.paragraphs:
                if marcador in p.text:
                    celula_destino = celula
                    break
            if celula_destino:
                break
        if celula_destino:
            break
    
    if not celula_destino:
        logger.warning(f"Marcador {marcador} não encontrado na tabela")
        return
    
    # Parsear o HTML
    soup = BeautifulSoup(html_texto, 'html.parser')
    
    # Limpar o conteúdo atual da célula
    for p in list(celula_destino.paragraphs):
        if marcador in p.text:
            # Mantemos o parágrafo, mas limpamos o texto
            p.clear()
            primeiro_paragrafo = p
            break
    else:
        # Se não encontrou o marcador, cria um novo parágrafo
        primeiro_paragrafo = celula_destino.add_paragraph()
    
    # Processar o conteúdo HTML e aplicar ao documento
    paragrafo_atual = primeiro_paragrafo
    
    # Manter controle do primeiro parágrafo
    primeiro_processado = True
    
    # Função auxiliar para processar conteúdo com formatação
    def processar_conteudo_formatado(elemento, paragrafo):
        """Processa elementos mantendo formatação como negrito, itálico, sublinhado e cores"""
        # Criar o run primeiro
        run = paragrafo.add_run(elemento.get_text())
        
        # Negrito
        if elemento.name == 'strong' or elemento.name == 'b':
            run.bold = True
        
        # Itálico
        elif elemento.name == 'em' or elemento.name == 'i':
            run.italic = True
        
        # Sublinhado
        elif elemento.name == 'u':
            run.underline = True
        
        # Código
        elif elemento.name == 'code':
            # Configurar fonte Courier New, tamanho 10
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            
            # Cor do texto azul
            run.font.color.rgb = RGBColor(3, 105, 161)  # #0369A1
            
            # Cor de fundo azul claro
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D2F0FD')  # #D2F0FD
            run._r.get_or_add_rPr().append(shading_elm)
        
        # Span com estilos
        elif elemento.name == 'span':
            style = elemento.get('style', '')
            logger.info(f"Processando span com estilo: {style}")
            
            # Verificar cor do texto
            cor_match = re.search(r'color:\s*rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', style)
            if cor_match:
                r, g, b = map(int, cor_match.groups())
                run.font.color.rgb = RGBColor(r, g, b)
                logger.info(f"Cor do texto definida: RGB({r}, {g}, {b})")
            
            # Verificar classe de alinhamento
            align_class = elemento.get('class', '')
            if align_class:
                if 'ql-align-center' in align_class:
                    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    logger.info("Alinhamento centralizado aplicado ao parágrafo")
                elif 'ql-align-right' in align_class:
                    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    logger.info("Alinhamento à direita aplicado ao parágrafo")
                elif 'ql-align-justify' in align_class:
                    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    logger.info("Alinhamento justificado aplicado ao parágrafo")
                elif 'ql-align-left' in align_class:
                    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    logger.info("Alinhamento à esquerda aplicado ao parágrafo")
        
        return run
    
    # Função auxiliar para processar listas
    def processar_lista(lista_elemento, usar_numeros=True):
        """Processa listas ordenadas ou não ordenadas"""
        nonlocal primeiro_processado, paragrafo_atual
        
        contador = 1
        for li in lista_elemento.find_all('li', recursive=False):
            if not primeiro_processado:
                paragrafo_atual = celula_destino.add_paragraph()
            else:
                primeiro_processado = False
            
            # Verificar se o item da lista tem data-list="bullet"
            if li.get('data-list') == 'bullet' or lista_elemento.get('data-list') == 'bullet':
                # Lista não ordenada
                paragrafo_atual.add_run("• ")
            elif usar_numeros:
                # Lista ordenada
                paragrafo_atual.add_run(f"{contador}. ")
                contador += 1
            else:
                # Lista não ordenada (fallback)
                paragrafo_atual.add_run("• ")
            
            # Processar o conteúdo do item de lista mantendo formatação
            for item in li.contents:
                if item.name in ['strong', 'b', 'em', 'i', 'u', 'code', 'span']:  # Adicionado 'u', 'code' e 'span'
                    processar_conteudo_formatado(item, paragrafo_atual)
                elif item.name in ['ul', 'ol']:
                    # Lista aninhada - processar recursivamente
                    processar_lista(item, item.name == 'ol' and item.get('data-list') != 'bullet')
                elif item.string:
                    paragrafo_atual.add_run(item.string)
            
            # Aplicar recuo
            paragrafo_atual.paragraph_format.left_indent = Pt(20)
    
    # Processar cada elemento do HTML
    for elemento in soup:
        # Ignorar elementos vazios ou apenas com espaço
        if elemento.name is None and (not elemento.string or not elemento.string.strip()):
            continue
            
        # Processar parágrafos
        if elemento.name == 'p':
            # Se não for o primeiro parágrafo, criar um novo
            if not primeiro_processado:
                paragrafo_atual = celula_destino.add_paragraph()
            else:
                primeiro_processado = False
            
            # Verificar classe de alinhamento no parágrafo
            align_class = elemento.get('class', '')
            if align_class:
                if 'ql-align-center' in align_class:
                    paragrafo_atual.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    logger.info("Alinhamento centralizado aplicado ao parágrafo")
                elif 'ql-align-right' in align_class:
                    paragrafo_atual.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    logger.info("Alinhamento à direita aplicado ao parágrafo")
                elif 'ql-align-justify' in align_class:
                    paragrafo_atual.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    logger.info("Alinhamento justificado aplicado ao parágrafo")
                elif 'ql-align-left' in align_class:
                    paragrafo_atual.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    logger.info("Alinhamento à esquerda aplicado ao parágrafo")
            
            # Processar o conteúdo do parágrafo
            for conteudo in elemento.contents:
                if conteudo.name in ['strong', 'b', 'em', 'i', 'u', 'code', 'span']:  # Adicionado 'u', 'code' e 'span'
                    processar_conteudo_formatado(conteudo, paragrafo_atual)
                elif conteudo.name == 'ul' or conteudo.name == 'ol':
                    # Lista dentro de parágrafo
                    usar_numeros = conteudo.name == 'ol' and conteudo.get('data-list') != 'bullet'
                    processar_lista(conteudo, usar_numeros)
                elif conteudo.string:
                    paragrafo_atual.add_run(conteudo.string)
        
        # Se for lista
        elif elemento.name == 'ul' or elemento.name == 'ol':
            # Verificar se é uma lista não ordenada do Quill
            usar_numeros = elemento.name == 'ol' and elemento.get('data-list') != 'bullet'
            processar_lista(elemento, usar_numeros)
        
        # Se for texto plano fora de tags
        elif elemento.string and elemento.string.strip():
            if not primeiro_processado:
                paragrafo_atual = celula_destino.add_paragraph()
            else:
                primeiro_processado = False
            paragrafo_atual.add_run(elemento.string.strip())
        
        # Se for quebra de linha
        elif elemento.name == 'br':
            if not primeiro_processado:
                paragrafo_atual = celula_destino.add_paragraph()
            else:
                primeiro_processado = False

def substituir_texto_com_quebras(tabela, marcador, texto):
    """
    Substitui um marcador por texto preservando quebras de linha.
    
    Args:
        tabela: Tabela onde buscar o marcador
        marcador: Marcador a ser substituído (ex: "[REGRAS_NEGOCIO_RF]")
        texto: Texto com quebras de linha para substituir o marcador
    """
    if not texto:
        return
    
    # Localizar a célula que contém o marcador
    celula_destino = None
    paragrafo_original = None
    
    for linha in tabela.rows:
        for celula in linha.cells:
            for p in celula.paragraphs:
                if marcador in p.text:
                    celula_destino = celula
                    paragrafo_original = p
                    break
            if celula_destino:
                break
        if celula_destino:
            break
    
    if not celula_destino:
        logger.warning(f"Marcador {marcador} não encontrado na tabela")
        return
    
    # Dividir o texto pelas quebras de linha
    linhas = texto.split('\n')
    
    # Limpar o parágrafo original
    paragrafo_original.clear()
    
    # Adicionar a primeira linha ao parágrafo original
    paragrafo_original.add_run(linhas[0])
    
    # Para cada linha adicional, criar um novo parágrafo
    for i in range(1, len(linhas)):
        if linhas[i].strip():  # Ignorar linhas vazias
            p = celula_destino.add_paragraph()
            p.add_run(linhas[i])
            
            # Se a linha começa com um marcador de lista (•, -, *), aplicar recuo
            if linhas[i].strip().startswith(('•', '-', '*')):
                p.paragraph_format.left_indent = Pt(20)

def processar_imagens(paragrafo, imagens):
    """
    Processa imagens em base64 e as adiciona ao parágrafo.
    
    Args:
        paragrafo: Parágrafo do documento Word
        imagens: Lista de imagens em formato base64
    """
    if not imagens or len(imagens) == 0:
        logger.warning("Nenhuma imagem para processar")
        return
    
    logger.info(f"Processando {len(imagens)} imagens")
    
    # Processar todas as imagens da lista
    for i, imagem in enumerate(imagens):
        # Verifica se é uma imagem em base64
        if isinstance(imagem, str) and imagem.startswith('data:image'):
            try:
                # Extrair a parte base64
                base64_data = imagem.split(',')[1]
                imagem_bytes = BytesIO(base64.b64decode(base64_data))
                
                # Adicionar uma quebra de linha entre imagens se não for a primeira
                if i > 0:
                    paragrafo.add_run().add_break()
                
                # Adicionar ao documento
                run = paragrafo.add_run()
                # Define a largura máxima como 6 polegadas (ajuste conforme necessário)
                run.add_picture(imagem_bytes, width=Inches(6))
                
                logger.info(f"Imagem {i+1} processada com sucesso")
                
            except Exception as e:
                logger.error(f"Erro ao processar imagem base64 {i+1}: {str(e)}")
                erro_run = paragrafo.add_run(f"[Erro ao processar imagem {i+1}]")
                # Destacar o erro em vermelho
                erro_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            logger.error(f"Formato de imagem não suportado para imagem {i+1}")
            erro_run = paragrafo.add_run(f"[Formato de imagem {i+1} não suportado]")
            erro_run.font.color.rgb = RGBColor(255, 0, 0)