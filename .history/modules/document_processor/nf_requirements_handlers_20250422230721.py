import logging
import os
import re
from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
from bs4 import BeautifulSoup

# Configuração de logging
logger = logging.getLogger(__name__)

def processar_requisitos_nao_funcionais(doc, requisitos_nao_funcionais):
    """
    Processa requisitos não funcionais no documento de Estratégia de Solução.
    
    Esta função localiza os parágrafos de RNF no documento e os substitui pelos
    requisitos não funcionais fornecidos. Se não houver requisitos, substitui por um
    texto padrão informando que não há requisitos não-funcionais.
    
    Args:
        doc: Documento Word (.docx) aberto
        requisitos_nao_funcionais: Lista de requisitos não funcionais para processar
    """
    if not requisitos_nao_funcionais or len(requisitos_nao_funcionais) == 0:
        logger.info("Nenhum requisito não funcional para processar. Removendo seção de RNF.")
        substituir_por_texto_padrao(doc)
        return
    
    logger.info(f"Processando {len(requisitos_nao_funcionais)} requisitos não funcionais")
    
    # Encontrar TODAS as ocorrências do título da seção de RNF
    indices_titulo_secao = []
    for i, p in enumerate(doc.paragraphs):
        if "3 - Requisitos Não-Funcionais" in p.text:
            indices_titulo_secao.append(i)
            logger.info(f"Encontrado título da seção de RNF no índice {i}: {p.text}")
    
    if not indices_titulo_secao:
        logger.error("Não foi possível encontrar o título da seção de RNF")
        return
        
    # Usamos a ÚLTIMA ocorrência do título, que provavelmente é a seção real de RNF
    indice_titulo_secao = indices_titulo_secao[-1]
    logger.info(f"Usando última ocorrência do título em {indice_titulo_secao}")
    
    # 1. Localizar parágrafos modelos de RNF (após o ÚLTIMO título da seção)
    paragrafo_modelo_rnf_titulo = None
    paragrafo_modelo_rnf_descricao = None
    indice_titulo_rnf = None
    indice_descricao_rnf = None
    
    # Buscar a partir do último título de seção RNF
    for i in range(indice_titulo_secao + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        
        # Localizar parágrafo modelo para título do RNF
        if "RNF-[N_RNF]" in p.text or "RNF-[N_RNF]: [TITULO_RNF]" in p.text:
            paragrafo_modelo_rnf_titulo = p
            indice_titulo_rnf = i
            logger.info(f"Encontrado parágrafo modelo para título de RNF no índice {i}: {p.text}")
            
        # Localizar parágrafo modelo para descrição do RNF
        elif "[DESCRICAO_RNF]" in p.text:
            paragrafo_modelo_rnf_descricao = p
            indice_descricao_rnf = i
            logger.info(f"Encontrado parágrafo modelo para descrição de RNF no índice {i}: {p.text}")
            
        # Se já encontramos um título que não é de RNF após o título de seção,
        # significa que saímos da seção de RNF
        elif p.style and "Heading" in p.style.name and "3 - Requisitos" not in p.text:
            logger.info(f"Encontrado fim da seção de RNF no índice {i}: {p.text}")
            break
    
    # Se não encontrar os modelos, usar padrões básicos
    if not paragrafo_modelo_rnf_titulo or not paragrafo_modelo_rnf_descricao:
        logger.warning("Não foi possível encontrar os parágrafos modelo completos para RNF. Usando estilos padrão.")
        
        # Se não encontrou o modelo de título, usar o primeiro parágrafo após o título da seção como referência
        if not paragrafo_modelo_rnf_titulo and indice_titulo_secao + 1 < len(doc.paragraphs):
            paragrafo_modelo_rnf_titulo = doc.paragraphs[indice_titulo_secao + 1]
            indice_titulo_rnf = indice_titulo_secao + 1
        
        # Se não encontrou o modelo de descrição, usar o segundo parágrafo após o título da seção
        if not paragrafo_modelo_rnf_descricao and indice_titulo_secao + 2 < len(doc.paragraphs):
            paragrafo_modelo_rnf_descricao = doc.paragraphs[indice_titulo_secao + 2]
            indice_descricao_rnf = indice_titulo_secao + 2
    
    # Armazenar o corpo do documento para inserções
    body = doc._body._body
    
    # 2. Remover parágrafos modelo existentes para RNF após o último título da seção
    paragrafos_a_remover = []
    
    # Vamos procurar todos os parágrafos que contêm marcadores de RNF após o título da seção
    for i in range(indice_titulo_secao + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if ("RNF-" in p.text and ("[N_RNF]" in p.text or "[TITULO_RNF]" in p.text)) or "[DESCRICAO_RNF]" in p.text:
            paragrafos_a_remover.append(p._p)
            logger.info(f"Marcando parágrafo para remoção no índice {i}: {p.text[:30]}...")
        
        # Se encontramos outro título de seção, paramos de procurar
        elif p.style and "Heading" in p.style.name and "3 - Requisitos" not in p.text:
            logger.info(f"Encontrado próximo título de seção no índice {i}, parando busca")
            break
    
    # Remover os parágrafos marcados para remoção
    logger.info(f"Removendo {len(paragrafos_a_remover)} parágrafos modelo de RNF")
    for p in paragrafos_a_remover:
        if p.getparent() is not None:
            try:
                p.getparent().remove(p)
            except Exception as e:
                logger.error(f"Erro ao remover parágrafo: {str(e)}")
    
    # 3. Inserir novos parágrafos para cada requisito não funcional
    # Vamos inserir imediatamente após o título da seção
    posicao_insercao = indice_titulo_secao + 1
    paragrafo_titulo = doc.paragraphs[indice_titulo_secao]
    
    for idx, requisito in enumerate(requisitos_nao_funcionais):
        logger.info(f"Processando RNF {idx+1}: {requisito.get('titulo', 'Sem título')}")
        
        # a. Criar novo parágrafo para o título do RNF
        novo_titulo = doc.add_paragraph()
        
        # Copiar estilo ou definir um estilo básico
        if paragrafo_modelo_rnf_titulo and hasattr(paragrafo_modelo_rnf_titulo, 'style') and paragrafo_modelo_rnf_titulo.style:
            novo_titulo.style = paragrafo_modelo_rnf_titulo.style
        else:
            # Estilo padrão se não tiver modelo
            novo_titulo.style = 'Heading 4' if 'Heading 4' in doc.styles else None
            
        # Copiar formatação do parágrafo modelo se existir
        if paragrafo_modelo_rnf_titulo and hasattr(paragrafo_modelo_rnf_titulo, 'paragraph_format'):
            # Copiar alinhamento
            if hasattr(paragrafo_modelo_rnf_titulo.paragraph_format, 'alignment'):
                novo_titulo.paragraph_format.alignment = paragrafo_modelo_rnf_titulo.paragraph_format.alignment
            # Copiar recuo
            if hasattr(paragrafo_modelo_rnf_titulo.paragraph_format, 'left_indent'):
                novo_titulo.paragraph_format.left_indent = paragrafo_modelo_rnf_titulo.paragraph_format.left_indent
            # Copiar espaçamento
            if hasattr(paragrafo_modelo_rnf_titulo.paragraph_format, 'space_before'):
                novo_titulo.paragraph_format.space_before = paragrafo_modelo_rnf_titulo.paragraph_format.space_before
            if hasattr(paragrafo_modelo_rnf_titulo.paragraph_format, 'space_after'):
                novo_titulo.paragraph_format.space_after = paragrafo_modelo_rnf_titulo.paragraph_format.space_after
        
        # Substituir os marcadores no texto do título
        num_rnf = str(idx + 1).zfill(2)  # RNF-01, RNF-02, etc.
        titulo = requisito.get('titulo', '[Sem título]')
        
        # Criar o texto do título formatado
        run = novo_titulo.add_run(f"RNF-{num_rnf}: {titulo}")
        
        # Copiar formatação do run original se existir
        if paragrafo_modelo_rnf_titulo and paragrafo_modelo_rnf_titulo.runs:
            run_modelo = paragrafo_modelo_rnf_titulo.runs[0]
            if hasattr(run_modelo, 'font'):
                # Copiar fonte
                if hasattr(run_modelo.font, 'name'):
                    run.font.name = run_modelo.font.name
                # Copiar tamanho
                if hasattr(run_modelo.font, 'size'):
                    run.font.size = run_modelo.font.size
                # Copiar negrito
                if hasattr(run_modelo.font, 'bold'):
                    run.font.bold = run_modelo.font.bold
                # Copiar itálico
                if hasattr(run_modelo.font, 'italic'):
                    run.font.italic = run_modelo.font.italic
                # Copiar cor
                if hasattr(run_modelo.font, 'color') and run_modelo.font.color and run_modelo.font.color.rgb:
                    run.font.color.rgb = run_modelo.font.color.rgb
        
        # Remove o parágrafo da sua posição atual
        novo_titulo_elemento = novo_titulo._p
        novo_titulo._p.getparent().remove(novo_titulo._p)
        
        # Insere o parágrafo na posição correta
        body.insert(posicao_insercao, novo_titulo_elemento)
        posicao_insercao += 1
        
        # b. Criar novo parágrafo para a descrição do RNF
        nova_descricao = doc.add_paragraph()
        
        # Copiar estilo do parágrafo modelo de descrição
        if paragrafo_modelo_rnf_descricao and hasattr(paragrafo_modelo_rnf_descricao, 'style') and paragrafo_modelo_rnf_descricao.style:
            nova_descricao.style = paragrafo_modelo_rnf_descricao.style
            
        # Copiar formatação do parágrafo modelo
        if paragrafo_modelo_rnf_descricao and hasattr(paragrafo_modelo_rnf_descricao, 'paragraph_format'):
            # Copiar alinhamento
            if hasattr(paragrafo_modelo_rnf_descricao.paragraph_format, 'alignment'):
                nova_descricao.paragraph_format.alignment = paragrafo_modelo_rnf_descricao.paragraph_format.alignment
            # Copiar recuo
            if hasattr(paragrafo_modelo_rnf_descricao.paragraph_format, 'left_indent'):
                nova_descricao.paragraph_format.left_indent = paragrafo_modelo_rnf_descricao.paragraph_format.left_indent
            # Copiar espaçamento
            if hasattr(paragrafo_modelo_rnf_descricao.paragraph_format, 'space_before'):
                nova_descricao.paragraph_format.space_before = paragrafo_modelo_rnf_descricao.paragraph_format.space_before
            if hasattr(paragrafo_modelo_rnf_descricao.paragraph_format, 'space_after'):
                nova_descricao.paragraph_format.space_after = paragrafo_modelo_rnf_descricao.paragraph_format.space_after
        
        # Substituir os marcadores no texto da descrição
        descricao = requisito.get('descricao', '[Sem descrição]')
        
        # Criar o texto da descrição formatado
        run = nova_descricao.add_run(descricao)
        
        # Copiar formatação do run original
        if paragrafo_modelo_rnf_descricao and paragrafo_modelo_rnf_descricao.runs:
            run_modelo = paragrafo_modelo_rnf_descricao.runs[0]
            if hasattr(run_modelo, 'font'):
                # Copiar fonte
                if hasattr(run_modelo.font, 'name'):
                    run.font.name = run_modelo.font.name
                # Copiar tamanho
                if hasattr(run_modelo.font, 'size'):
                    run.font.size = run_modelo.font.size
                # Copiar negrito
                if hasattr(run_modelo.font, 'bold'):
                    run.font.bold = run_modelo.font.bold
                # Copiar itálico
                if hasattr(run_modelo.font, 'italic'):
                    run.font.italic = run_modelo.font.italic
                # Copiar cor
                if hasattr(run_modelo.font, 'color') and run_modelo.font.color and run_modelo.font.color.rgb:
                    run.font.color.rgb = run_modelo.font.color.rgb
        
        # Remove o parágrafo da sua posição atual
        nova_descricao_elemento = nova_descricao._p
        nova_descricao._p.getparent().remove(nova_descricao._p)
        
        # Insere o parágrafo na posição correta
        body.insert(posicao_insercao, nova_descricao_elemento)
        posicao_insercao += 1
    
    logger.info(f"Processamento de {len(requisitos_nao_funcionais)} RNFs concluído com sucesso")

def substituir_por_texto_padrao(doc):
    """
    Substitui a seção de requisitos não funcionais por um texto padrão.
    
    Args:
        doc: Documento Word (.docx) aberto
    """
    # Encontrar TODAS as ocorrências do título da seção de RNF
    indices_titulo_secao = []
    for i, p in enumerate(doc.paragraphs):
        if "3 - Requisitos Não-Funcionais" in p.text:
            indices_titulo_secao.append(i)
            logger.info(f"Encontrado título da seção de RNF no índice {i}: {p.text}")
    
    if not indices_titulo_secao:
        logger.error("Não foi possível encontrar o título da seção de RNF")
        return
        
    # Usamos a ÚLTIMA ocorrência do título, que provavelmente é a seção real de RNF
    indice_titulo_secao = indices_titulo_secao[-1]
    logger.info(f"Usando última ocorrência do título em {indice_titulo_secao}")
    
    # Localizar e remover parágrafos RNF existentes
    paragrafos_a_remover = []
    for i in range(indice_titulo_secao + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if ("RNF-" in p.text and ("[N_RNF]" in p.text or "[TITULO_RNF]" in p.text)) or "[DESCRICAO_RNF]" in p.text:
            paragrafos_a_remover.append(p._p)
            logger.info(f"Marcando parágrafo para remoção no índice {i}: {p.text[:30]}...")
        
        # Se encontramos outro título de seção, paramos de procurar
        elif p.style and "Heading" in p.style.name and "3 - Requisitos" not in p.text:
            logger.info(f"Encontrado próximo título de seção no índice {i}, parando busca")
            break
    
    # Remover parágrafos marcados
    for p in paragrafos_a_remover:
        if p.getparent() is not None:
            p.getparent().remove(p)
            logger.info("Removido parágrafo de RNF")
    
    # Adicionar parágrafo com texto padrão após o título da seção
    paragrafo_padrao = doc.add_paragraph()
    
    # Definir estilo padrão
    run = paragrafo_padrao.add_run("Não há requisitos não-funcionais.")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    
    # Obter o elemento do parágrafo e o corpo do documento
    novo_elemento = paragrafo_padrao._p
    body = doc._body._body
    
    # Remover o parágrafo de sua posição atual (no final do documento)
    if novo_elemento.getparent() is not None:
        novo_elemento.getparent().remove(novo_elemento)
    
    # Inserir após o título da seção
    try:
        # Inserir na posição exatamente após o título da seção
        posicao_insercao = indice_titulo_secao + 1
        
        # Obtém um ponto de referência (o elemento após o título)
        if indice_titulo_secao + 1 < len(doc.paragraphs):
            ponto_referencia = doc.paragraphs[indice_titulo_secao + 1]._p
            # Insere o novo elemento antes do ponto de referência
            ponto_referencia.addprevious(novo_elemento)
            logger.info("Inserido parágrafo com texto padrão após o título de RNF")
        else:
            # Se não tem elemento depois, adiciona no final
            body.append(novo_elemento)
            logger.info("Adicionado parágrafo com texto padrão no final do documento")
    except Exception as e:
        logger.error(f"Erro ao inserir texto padrão: {str(e)}")
        
        # Método alternativo: inserir no corpo do documento
        try:
            body.insert(indice_titulo_secao + 1, novo_elemento)
            logger.info("Inserido parágrafo com texto padrão usando método alternativo")
        except Exception as e:
            logger.error(f"Erro ao inserir texto padrão (método alternativo): {str(e)}")
    
    logger.info("Substituída a seção de RNF por texto padrão")