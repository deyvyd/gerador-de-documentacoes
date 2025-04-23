import logging
import os
import re
from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
from bs4 import BeautifulSoup

# Configuração de logging
logger = logging.getLogger(__name__)

def processar_requisitos_nao_funcionais(doc, requisitos_nao_funcionais):
    """
    Processa requisitos não funcionais no documento de Estratégia de Solução.
    
    Esta função localiza os parágrafos de RNF no documento e os substitui pelos
    requisitos não funcionais fornecidos. Se não houver requisitos, substitui por um
    texto padrão informando que não há requisitos não-funcionais.
    
    Args:
        doc: Documento Word (.docx) aberto
        requisitos_nao_funcionais: Lista de requisitos não funcionais para processar
    """
    logger.info(f"Iniciando processamento de requisitos não funcionais (total: {len(requisitos_nao_funcionais) if requisitos_nao_funcionais else 0})")
    
    # 1. Localizar a última ocorrência do título da seção RNF
    indice_titulo_secao = None
    paragrafo_titulo_secao = None
    
    for i, p in enumerate(doc.paragraphs):
        if "3 - Requisitos Não-Funcionais" in p.text:
            indice_titulo_secao = i
            paragrafo_titulo_secao = p
            logger.info(f"Encontrado título da seção RNF no índice {i}: '{p.text}'")
    
    if indice_titulo_secao is None:
        logger.error("Título da seção de RNF não encontrado no documento")
        return
        
    logger.info(f"Usando título da seção RNF no índice {indice_titulo_secao}")
    
    # 2. Localizar o primeiro título de seção após a seção de RNF (ou o final do documento)
    indice_proxima_secao = len(doc.paragraphs)
    for i in range(indice_titulo_secao + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        
        # Verifica se é um título de seção (por estilo) e diferente do título de RNF
        if p.style and "Heading" in p.style.name and "3 - Requisitos" not in p.text:
            indice_proxima_secao = i
            logger.info(f"Encontrado próximo título de seção no índice {i}: '{p.text}'")
            break
        
        # Ou se tem um texto que indica o início de uma nova seção
        elif p.text.strip().startswith("4 -") or p.text.strip() == "Conclusão":
            indice_proxima_secao = i
            logger.info(f"Encontrado possível próximo título de seção no índice {i}: '{p.text}'")
            break
    
    # 3. Remover todos os parágrafos entre o título da seção RNF e o início da próxima seção
    paragrafos_a_remover = []
    for i in range(indice_titulo_secao + 1, indice_proxima_secao):
        paragrafos_a_remover.append(doc.paragraphs[i]._p)
        logger.info(f"Marcando parágrafo para remoção no índice {i}: '{doc.paragraphs[i].text[:30]}...'")
    
    # Remover os parágrafos marcados de trás para frente para não afetar os índices
    for p in reversed(paragrafos_a_remover):
        if p.getparent() is not None:
            p.getparent().remove(p)
            logger.info("Removido parágrafo")
    
    # 4. Adicionar parágrafos para requisitos não funcionais ou texto padrão
    # Referência para inserção é o corpo do documento
    body = doc._body._body
    
    # Posição de inserção: logo após o título da seção
    posicao_insercao = indice_titulo_secao + 1
    
    if not requisitos_nao_funcionais or len(requisitos_nao_funcionais) == 0:
        # 4.A. Se não houver requisitos, adicionar texto padrão
        logger.info("Nenhum requisito não funcional, adicionando texto padrão")
        
        paragrafo_padrao = doc.add_paragraph()
        run = paragrafo_padrao.add_run("Não há requisitos não-funcionais.")
        
        # Definir estilo padrão
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        
        # Remove o parágrafo de sua posição atual
        elemento_padrao = paragrafo_padrao._p
        if elemento_padrao.getparent() is not None:
            elemento_padrao.getparent().remove(elemento_padrao)
        
        # Inserir após o título da seção
        try:
            body.insert(posicao_insercao, elemento_padrao)
            logger.info("Texto padrão inserido com sucesso")
        except Exception as e:
            logger.error(f"Erro ao inserir texto padrão: {str(e)}")
            
    else:
        # 4.B. Se houver requisitos, adicionar cada um deles
        for idx, requisito in enumerate(requisitos_nao_funcionais):
            logger.info(f"Processando RNF {idx+1}: {requisito.get('titulo', 'Sem título')}")
            
            # Criar parágrafo para o título do RNF
            titulo_rnf = doc.add_paragraph()
            num_rnf = str(idx + 1).zfill(2)  # RNF-01, RNF-02, etc.
            texto_titulo = requisito.get('titulo', 'Sem título')
            
            # Aplicar estilo semelhante ao do título da seção, mas um nível abaixo
            if paragrafo_titulo_secao.style:
                try:
                    # Tentar encontrar um estilo de nível inferior
                    estilo_nivel = int(paragrafo_titulo_secao.style.name.split(' ')[-1]) if paragrafo_titulo_secao.style.name.split(' ')[-1].isdigit() else 1
                    estilo_nivel += 1  # Um nível abaixo
                    estilo_nome = f"Heading {estilo_nivel}"
                    
                    # Verificar se o estilo existe
                    if estilo_nome in doc.styles:
                        titulo_rnf.style = estilo_nome
                    else:
                        titulo_rnf.style = paragrafo_titulo_secao.style
                except:
                    # Usar o mesmo estilo do título da seção
                    titulo_rnf.style = paragrafo_titulo_secao.style
            
            # Copiar formatação relevante do parágrafo título da seção
            if hasattr(paragrafo_titulo_secao, 'paragraph_format'):
                if hasattr(paragrafo_titulo_secao.paragraph_format, 'alignment'):
                    titulo_rnf.paragraph_format.alignment = paragrafo_titulo_secao.paragraph_format.alignment
                if hasattr(paragrafo_titulo_secao.paragraph_format, 'space_before'):
                    titulo_rnf.paragraph_format.space_before = paragrafo_titulo_secao.paragraph_format.space_before
                if hasattr(paragrafo_titulo_secao.paragraph_format, 'space_after'):
                    titulo_rnf.paragraph_format.space_after = paragrafo_titulo_secao.paragraph_format.space_after
            
            # Adicionar o texto do título com fonte similar
            run_titulo = titulo_rnf.add_run(f"RNF-{num_rnf}: {texto_titulo}")
            if paragrafo_titulo_secao.runs:
                for attr in ['bold', 'italic', 'name', 'size', 'underline']:
                    if hasattr(paragrafo_titulo_secao.runs[0].font, attr):
                        setattr(run_titulo.font, attr, getattr(paragrafo_titulo_secao.runs[0].font, attr))
            
            # Remove o elemento de título da posição atual
            elemento_titulo = titulo_rnf._p
            if elemento_titulo.getparent() is not None:
                elemento_titulo.getparent().remove(elemento_titulo)
            
            # Adicionar o elemento de título na posição correta
            body.insert(posicao_insercao, elemento_titulo)
            posicao_insercao += 1
            
            # Criar parágrafo para a descrição do RNF
            descricao_rnf = doc.add_paragraph()
            texto_descricao = requisito.get('descricao', 'Sem descrição.')
            
            # Adicionar o texto da descrição com estilo padrão
            run_descricao = descricao_rnf.add_run(texto_descricao)
            run_descricao.font.name = "Calibri"
            run_descricao.font.size = Pt(12)
            
            # Remove o elemento de descrição da posição atual
            elemento_descricao = descricao_rnf._p
            if elemento_descricao.getparent() is not None:
                elemento_descricao.getparent().remove(elemento_descricao)
            
            # Adicionar o elemento de descrição na posição correta
            body.insert(posicao_insercao, elemento_descricao)
            posicao_insercao += 1
    
    logger.info("Processamento de requisitos não funcionais concluído com sucesso")

def substituir_por_texto_padrao(doc):
    """
    Substitui a seção de requisitos não funcionais por um texto padrão.
    
    Args:
        doc: Documento Word (.docx) aberto
    """
    # Encontrar TODAS as ocorrências do título da seção de RNF
    indices_titulo_secao = []
    for i, p in enumerate(doc.paragraphs):
        if "3 - Requisitos Não-Funcionais" in p.text:
            indices_titulo_secao.append(i)
            logger.info(f"Encontrado título da seção de RNF no índice {i}: {p.text}")
    
    if not indices_titulo_secao:
        logger.error("Não foi possível encontrar o título da seção de RNF")
        return
        
    # Usamos a ÚLTIMA ocorrência do título, que provavelmente é a seção real de RNF
    indice_titulo_secao = indices_titulo_secao[-1]
    logger.info(f"Usando última ocorrência do título em {indice_titulo_secao}")
    
    # Localizar e remover parágrafos RNF existentes
    paragrafos_a_remover = []
    for i in range(indice_titulo_secao + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if ("RNF-" in p.text and ("[N_RNF]" in p.text or "[TITULO_RNF]" in p.text)) or "[DESCRICAO_RNF]" in p.text:
            paragrafos_a_remover.append(p._p)
            logger.info(f"Marcando parágrafo para remoção no índice {i}: {p.text[:30]}...")
        
        # Se encontramos outro título de seção, paramos de procurar
        elif p.style and "Heading" in p.style.name and "3 - Requisitos" not in p.text:
            logger.info(f"Encontrado próximo título de seção no índice {i}, parando busca")
            break
    
    # Remover parágrafos marcados
    for p in paragrafos_a_remover:
        if p.getparent() is not None:
            p.getparent().remove(p)
            logger.info("Removido parágrafo de RNF")
    
    # Adicionar parágrafo com texto padrão após o título da seção
    paragrafo_padrao = doc.add_paragraph()
    
    # Definir estilo padrão
    run = paragrafo_padrao.add_run("Não há requisitos não-funcionais.")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    
    # Obter o elemento do parágrafo e o corpo do documento
    novo_elemento = paragrafo_padrao._p
    body = doc._body._body
    
    # Remover o parágrafo de sua posição atual (no final do documento)
    if novo_elemento.getparent() is not None:
        novo_elemento.getparent().remove(novo_elemento)
    
    # Inserir após o título da seção
    try:
        # Inserir na posição exatamente após o título da seção
        posicao_insercao = indice_titulo_secao + 1
        
        # Obtém um ponto de referência (o elemento após o título)
        if indice_titulo_secao + 1 < len(doc.paragraphs):
            ponto_referencia = doc.paragraphs[indice_titulo_secao + 1]._p
            # Insere o novo elemento antes do ponto de referência
            ponto_referencia.addprevious(novo_elemento)
            logger.info("Inserido parágrafo com texto padrão após o título de RNF")
        else:
            # Se não tem elemento depois, adiciona no final
            body.append(novo_elemento)
            logger.info("Adicionado parágrafo com texto padrão no final do documento")
    except Exception as e:
        logger.error(f"Erro ao inserir texto padrão: {str(e)}")
        
        # Método alternativo: inserir no corpo do documento
        try:
            body.insert(indice_titulo_secao + 1, novo_elemento)
            logger.info("Inserido parágrafo com texto padrão usando método alternativo")
        except Exception as e:
            logger.error(f"Erro ao inserir texto padrão (método alternativo): {str(e)}")
    
    logger.info("Substituída a seção de RNF por texto padrão")