import logging
import os
import re
from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
from bs4 import BeautifulSoup

# Configuração de logging
logger = logging.getLogger(__name__)

def processar_requisitos_nao_funcionais(doc, requisitos_nao_funcionais, posicao_insercao_inicial):
    """
    Processa requisitos não funcionais no documento de Estratégia de Solução.
    
    Esta função localiza os parágrafos de RNF no documento e os substitui pelos
    requisitos não funcionais fornecidos. Se não houver requisitos, substitui por um
    texto padrão informando que não há requisitos não-funcionais.
    
    Args:
        doc: Documento Word (.docx) aberto
        requisitos_nao_funcionais: Lista de requisitos não funcionais para processar
    """
    if not requisitos_nao_funcionais or len(requisitos_nao_funcionais) == 0:
        logger.info("Nenhum requisito não funcional para processar. Removendo seção de RNF.")
        substituir_por_texto_padrao(doc)
        return
    
    logger.info(f"Processando {len(requisitos_nao_funcionais)} requisitos não funcionais")
    
    # 1. Localizar parágrafos modelos de RNF
    paragrafo_titulo_rnf = None
    paragrafo_modelo_rnf_titulo = None
    paragrafo_modelo_rnf_descricao = None
    
    indice_titulo_secao = None
    indice_titulo_rnf = None
    indice_descricao_rnf = None
    
    for i, p in enumerate(doc.paragraphs):
        # Localizar título da seção
        if "3 - Requisitos Não-Funcionais" in p.text:
            paragrafo_titulo_rnf = p
            indice_titulo_secao = i
            logger.info(f"Encontrado título da seção de RNF no índice {i}: {p.text}")
        
        # Localizar parágrafo modelo para título do RNF
        elif "RNF-[N_RNF]" in p.text or "RNF-[N_RNF]: [TITULO_RNF]" in p.text:
            paragrafo_modelo_rnf_titulo = p
            indice_titulo_rnf = i
            logger.info(f"Encontrado parágrafo modelo para título de RNF no índice {i}: {p.text}")
        
        # Localizar parágrafo modelo para descrição do RNF
        elif "[DESCRICAO_RNF]" in p.text:
            paragrafo_modelo_rnf_descricao = p
            indice_descricao_rnf = i
            logger.info(f"Encontrado parágrafo modelo para descrição de RNF no índice {i}: {p.text}")
    
    # Verificar se encontrou ambos os parágrafos modelo necessários
    if not (paragrafo_modelo_rnf_titulo and paragrafo_modelo_rnf_descricao):
        logger.error("Não foi possível encontrar os parágrafos modelo para RNF")
        return
    
    # Armazenar o corpo do documento para inserções
    body = doc._body._body
    
    # 2. Remover parágrafos modelo existentes para RNF, preservando o título da seção
    paragrafos_a_remover = []
    
    if indice_titulo_rnf is not None and indice_titulo_rnf > indice_titulo_secao:
        # Adicionar todos os parágrafos a partir do parágrafo modelo até o final do documento
        for i in range(indice_titulo_rnf, len(doc.paragraphs)):
            paragrafos_a_remover.append(doc.paragraphs[i]._p)
    
    # Remover os parágrafos marcados para remoção
    logger.info(f"Removendo {len(paragrafos_a_remover)} parágrafos modelo de RNF")
    for p in paragrafos_a_remover:
        if p.getparent() is not None:
            try:
                p.getparent().remove(p)
            except Exception as e:
                logger.error(f"Erro ao remover parágrafo: {str(e)}")
    
    # 3. Inserir novos parágrafos para cada requisito não funcional
    #posicao_insercao = indice_titulo_secao + 1  # Após o título da seção
    posicao_insercao = posicao_insercao_inicial  # Posição inicial fornecida
    logger.info(f"Inserindo novo parágrafo para requisito não funcional no índice {posicao_insercao}")
    
    for idx, requisito in enumerate(requisitos_nao_funcionais):
        logger.info(f"Processando RNF {idx+1}: {requisito.get('titulo', 'Sem título')}")
        
        # a. Criar novo parágrafo para o título do RNF
        novo_titulo = doc.add_paragraph()
        
        # Copiar estilo do parágrafo modelo de título
        if hasattr(paragrafo_modelo_rnf_titulo, 'style') and paragrafo_modelo_rnf_titulo.style:
            novo_titulo.style = paragrafo_modelo_rnf_titulo.style
            
        # Copiar formatação do parágrafo modelo
        if hasattr(paragrafo_modelo_rnf_titulo, 'paragraph_format'):
            # Copiar alinhamento
            if hasattr(paragrafo_modelo_rnf_titulo.paragraph_format, 'alignment'):
                novo_titulo.paragraph_format.alignment = paragrafo_modelo_rnf_titulo.paragraph_format.alignment
            # Copiar recuo
            if hasattr(paragrafo_modelo_rnf_titulo.paragraph_format, 'left_indent'):
                novo_titulo.paragraph_format.left_indent = paragrafo_modelo_rnf_titulo.paragraph_format.left_indent
            # Copiar espaçamento
            if hasattr(paragrafo_modelo_rnf_titulo.paragraph_format, 'space_before'):
                novo_titulo.paragraph_format.space_before = paragrafo_modelo_rnf_titulo.paragraph_format.space_before
            if hasattr(paragrafo_modelo_rnf_titulo.paragraph_format, 'space_after'):
                novo_titulo.paragraph_format.space_after = paragrafo_modelo_rnf_titulo.paragraph_format.space_after
        
        # Substituir os marcadores no texto do título
        num_rnf = str(idx + 1).zfill(2)  # RNF-01, RNF-02, etc.
        titulo = requisito.get('titulo', '[Sem título]')
        
        # Criar o texto do título formatado
        run = novo_titulo.add_run(f"RNF-{num_rnf}: {titulo}")
        
        # Copiar formatação do run original
        if paragrafo_modelo_rnf_titulo.runs:
            run_modelo = paragrafo_modelo_rnf_titulo.runs[0]
            if hasattr(run_modelo, 'font'):
                # Copiar fonte
                if hasattr(run_modelo.font, 'name'):
                    run.font.name = run_modelo.font.name
                # Copiar tamanho
                if hasattr(run_modelo.font, 'size'):
                    run.font.size = run_modelo.font.size
                # Copiar negrito
                if hasattr(run_modelo.font, 'bold'):
                    run.font.bold = run_modelo.font.bold
                # Copiar itálico
                if hasattr(run_modelo.font, 'italic'):
                    run.font.italic = run_modelo.font.italic
                # Copiar cor
                if hasattr(run_modelo.font, 'color') and run_modelo.font.color and run_modelo.font.color.rgb:
                    run.font.color.rgb = run_modelo.font.color.rgb
        
        # Remove o parágrafo da sua posição atual
        novo_titulo_elemento = novo_titulo._p
        novo_titulo._p.getparent().remove(novo_titulo._p)
        
        # Insere o parágrafo na posição correta
        body.insert(posicao_insercao, novo_titulo_elemento)
        posicao_insercao += 1
        
        # b. Criar novo parágrafo para a descrição do RNF
        nova_descricao = doc.add_paragraph()
        
        # Copiar estilo do parágrafo modelo de descrição
        if hasattr(paragrafo_modelo_rnf_descricao, 'style') and paragrafo_modelo_rnf_descricao.style:
            nova_descricao.style = paragrafo_modelo_rnf_descricao.style
            
        # Copiar formatação do parágrafo modelo
        if hasattr(paragrafo_modelo_rnf_descricao, 'paragraph_format'):
            # Copiar alinhamento
            if hasattr(paragrafo_modelo_rnf_descricao.paragraph_format, 'alignment'):
                nova_descricao.paragraph_format.alignment = paragrafo_modelo_rnf_descricao.paragraph_format.alignment
            # Copiar recuo
            if hasattr(paragrafo_modelo_rnf_descricao.paragraph_format, 'left_indent'):
                nova_descricao.paragraph_format.left_indent = paragrafo_modelo_rnf_descricao.paragraph_format.left_indent
            # Copiar espaçamento
            if hasattr(paragrafo_modelo_rnf_descricao.paragraph_format, 'space_before'):
                nova_descricao.paragraph_format.space_before = paragrafo_modelo_rnf_descricao.paragraph_format.space_before
            if hasattr(paragrafo_modelo_rnf_descricao.paragraph_format, 'space_after'):
                nova_descricao.paragraph_format.space_after = paragrafo_modelo_rnf_descricao.paragraph_format.space_after
        
        # Substituir os marcadores no texto da descrição
        descricao = requisito.get('descricao', '[Sem descrição]')
        
        # Criar o texto da descrição formatado
        run = nova_descricao.add_run(descricao)
        
        # Copiar formatação do run original
        if paragrafo_modelo_rnf_descricao.runs:
            run_modelo = paragrafo_modelo_rnf_descricao.runs[0]
            if hasattr(run_modelo, 'font'):
                # Copiar fonte
                if hasattr(run_modelo.font, 'name'):
                    run.font.name = run_modelo.font.name
                # Copiar tamanho
                if hasattr(run_modelo.font, 'size'):
                    run.font.size = run_modelo.font.size
                # Copiar negrito
                if hasattr(run_modelo.font, 'bold'):
                    run.font.bold = run_modelo.font.bold
                # Copiar itálico
                if hasattr(run_modelo.font, 'italic'):
                    run.font.italic = run_modelo.font.italic
                # Copiar cor
                if hasattr(run_modelo.font, 'color') and run_modelo.font.color and run_modelo.font.color.rgb:
                    run.font.color.rgb = run_modelo.font.color.rgb
        
        # Remove o parágrafo da sua posição atual
        nova_descricao_elemento = nova_descricao._p
        nova_descricao._p.getparent().remove(nova_descricao._p)
        
        # Insere o parágrafo na posição correta
        body.insert(posicao_insercao, nova_descricao_elemento)
        posicao_insercao += 1
    
    logger.info(f"Processamento de {len(requisitos_nao_funcionais)} RNFs concluído com sucesso")

def substituir_por_texto_padrao(doc):
    """
    Substitui a seção de requisitos não funcionais por um texto padrão.
    
    Args:
        doc: Documento Word (.docx) aberto
    """
    # Localizar o título da seção e os parágrafos de RNF
    indice_titulo_secao = None
    paragrafos_rnf = []
    
    for i, p in enumerate(doc.paragraphs):
        if "3 - Requisitos Não-Funcionais" in p.text:
            indice_titulo_secao = i
            logger.info(f"Encontrado título da seção de RNF no índice {i}")
        elif indice_titulo_secao is not None and i > indice_titulo_secao:
            if "RNF-" in p.text or "[DESCRICAO_RNF]" in p.text or "[TITULO_RNF]" in p.text or "[N_RNF]" in p.text:
                paragrafos_rnf.append(p)
                logger.info(f"Marcando parágrafo de RNF para remoção: índice {i}, texto: '{p.text[:20]}...'")
    
    if indice_titulo_secao is None:
        logger.error("Não foi possível encontrar o título da seção de RNF")
        return
    
    # Remover parágrafos marcados
    for p in paragrafos_rnf:
        if p._p.getparent() is not None:
            p._p.getparent().remove(p._p)
            logger.info(f"Removido parágrafo de RNF")
    
    # Adicionar parágrafo com texto padrão após o título da seção
    novo_paragrafo = doc.add_paragraph()
    
    # Definir estilo padrão
    run = novo_paragrafo.add_run("Não há requisitos não-funcionais.")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    
    # Obter o elemento do parágrafo e o corpo do documento
    novo_elemento = novo_paragrafo._p
    body = doc._body._body
    
    # Remover o parágrafo de sua posição atual (no final do documento)
    if novo_elemento.getparent() is not None:
        novo_elemento.getparent().remove(novo_elemento)
    
    # Inserir após o título da seção
    try:
        # Obter o elemento do título da seção
        elemento_titulo = doc.paragraphs[indice_titulo_secao]._p
        
        # Inserir o novo parágrafo após o título
        elemento_titulo.addnext(novo_elemento)
        logger.info("Inserido parágrafo com texto padrão após o título de RNF")
    except Exception as e:
        logger.error(f"Erro ao inserir texto padrão: {str(e)}")
        
        # Método alternativo: inserir no índice após o título
        try:
            body.insert(indice_titulo_secao + 1, novo_elemento)
            logger.info("Inserido parágrafo com texto padrão usando método alternativo")
        except Exception as e:
            logger.error(f"Erro ao inserir texto padrão (método alternativo): {str(e)}")
    
    logger.info("Substituída a seção de RNF por texto padrão")