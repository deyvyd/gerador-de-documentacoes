# modules/document_processor/summary_handlers.py
import logging
import os
import re
from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from .utils import identificar_quebras_pagina
from .hyperlink_handlers import criar_hyperlink_marcador

# Configuração de logging
logger = logging.getLogger(__name__)

def obter_titulos_sumario(caminho_arquivo, titulo_documento=None):
    """
    Obtém os títulos do sumário de um documento.
    """
    doc = None

    doc = Document(caminho_arquivo)
    logger.debug(f"Documento carregado com sucesso: {caminho_arquivo}")
    
    if titulo_documento and hasattr(doc, 'core_properties'):
        doc.core_properties.title = titulo_documento
        logger.debug(f"Título do documento atualizado para: {titulo_documento}")
    
    # 1. Identificar quebras de página no documento
    quebras_pagina = identificar_quebras_pagina(doc)
    logger.info(f"Identificadas {len(quebras_pagina)} quebras de página explícitas no documento")
    
    # 2. Encontrar início do sumário (parágrafo com o título "Sumário")
    sumario_inicio = None
    for i, paragrafo in enumerate(doc.paragraphs):
        if paragrafo.text.strip().lower() == "sumário":
            sumario_inicio = i
            logger.info(f"Título do sumário encontrado no parágrafo {i}")
            break
    
    if sumario_inicio is None:
        # Busca alternativa por parágrafo com estilo TOC_Heading
        for i, paragrafo in enumerate(doc.paragraphs):
            if hasattr(paragrafo, 'style') and paragrafo.style and "TOC" in str(paragrafo.style.name):
                sumario_inicio = i
                logger.info(f"Título do sumário encontrado por estilo no parágrafo {i}")
                break
    
    if sumario_inicio is None:
        # Última tentativa: localizar pelo contexto do documento
        for i, paragrafo in enumerate(doc.paragraphs):
            if paragrafo.text.strip().lower() in ["índice", "conteúdo", "table of contents"]:
                sumario_inicio = i
                logger.info(f"Possível título do sumário encontrado no parágrafo {i}: '{paragrafo.text}'")
                break

    if sumario_inicio is None:
        logger.error("Título do sumário não encontrado no documento!")
        return False
    
    # 3. Encontrar fim do sumário atual
    sumario_fim = None
    
    for i in range(sumario_inicio + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        texto = p.text.strip()
        estilo = p.style.name if hasattr(p, 'style') and p.style and hasattr(p.style, 'name') else ""
        
        if (estilo == 'Title' and texto.startswith("1")) or \
           re.match(r'^1\s*[-–]', texto) or \
           texto.lower().startswith("introdução"):
            sumario_fim = i
            logger.info(f"Introdução encontrada no parágrafo {i}: '{texto}'")
            break
            
    if sumario_fim is None:
        sumario_fim = sumario_inicio + 3
        logger.warning(f"Introdução não encontrada após o sumário. Usando valor relativo {sumario_fim}.")
    
    # 4. Criar mapa de páginas baseado nas quebras de página
    pagina_atual = 1
    mapa_paragrafos_paginas = {}
    
    # Definimos o primeiro parágrafo como página 1
    for i in range(len(doc.paragraphs)):
        # Verifica se há uma quebra de página neste parágrafo
        if i in quebras_pagina:
            pagina_atual += 1
        mapa_paragrafos_paginas[i] = pagina_atual
    
    # 5. Coletar todos os títulos, suas páginas corretas e criar marcadores para links
    titulos = []
    marcadores_titulos = {}
    
    # Primeiro, vamos adicionar marcadores (bookmarks) a todos os títulos
    for i, p in enumerate(doc.paragraphs):
        # Ignoramos parágrafos dentro do sumário
        if sumario_inicio <= i < sumario_fim:
            continue
            
        texto = p.text.strip()
        estilo = p.style.name if hasattr(p, 'style') and p.style and hasattr(p.style, 'name') else ""
        
        # Determinamos a página do título baseada no mapa de quebras de página
        pagina_do_titulo = mapa_paragrafos_paginas.get(i, 1)
        
        # Verifica se é um título
        eh_titulo = False
        nivel_titulo = 0
        
        if (estilo == 'Title' or estilo == 'Heading 1') and texto:
            eh_titulo = True
            nivel_titulo = 1
        elif estilo == 'Heading 2' and texto:
            eh_titulo = True
            nivel_titulo = 2
        elif estilo == 'Heading 3' and texto:
            eh_titulo = True
            nivel_titulo = 3
        
        # Se é um título, cria um marcador (bookmark)
        if eh_titulo:
            # Gera um ID único para o título
            import hashlib
            marcador_id = f"titulo_{hashlib.md5(texto.encode()).hexdigest()[:8]}"
            
            # Adiciona um marcador (bookmark) para o título
            from docx.oxml.shared import OxmlElement, qn

            # Cria marcador de início
            start = OxmlElement('w:bookmarkStart')
            start.set(qn('w:id'), str(i))  # Usando o índice como ID único
            start.set(qn('w:name'), marcador_id)
            
            # Cria marcador de fim
            end = OxmlElement('w:bookmarkEnd')
            end.set(qn('w:id'), str(i))
            
            # Insere o marcador no início do parágrafo
            p._p.insert(0, start)
            p._p.append(end)
            
            # Guarda o ID do marcador para usá-lo nos links
            marcadores_titulos[i] = marcador_id
            
            # Processamento específico para títulos de nível 1
            if nivel_titulo == 1:
                numero_match = re.match(r'(\d+(?:\.\d+)*)\s*[-–—]?\s*(.*)', texto)
                
                if numero_match:
                    numero, titulo_texto = numero_match.groups()
                    titulos.append({
                        'nivel': 1,
                        'numero': numero,
                        'texto': titulo_texto.strip(),
                        'pagina': pagina_do_titulo,
                        'marcador': marcador_id,
                        'indice': i
                    })
                else:
                    titulos.append({
                        'nivel': 1,
                        'numero': None,
                        'texto': texto,
                        'pagina': pagina_do_titulo,
                        'marcador': marcador_id,
                        'indice': i
                    })
            elif nivel_titulo > 1:
                titulos.append({
                    'nivel': nivel_titulo,
                    'numero': None,
                    'texto': texto,
                    'pagina': pagina_do_titulo,
                    'marcador': marcador_id,
                    'indice': i
                })

    return titulos

def atualizar_sumario_com_python_docx(caminho_arquivo, titulo_documento=None, titulos=None):
    """
    Atualiza o sumário de um documento usando python-docx e detecção de quebras de página.
    """
    logger.info(f"Iniciando atualização do sumário: {caminho_arquivo}")
    
    if not os.path.exists(caminho_arquivo):
        logger.error(f"Arquivo não encontrado: {caminho_arquivo}")
        return False
    
    doc = None
    
    try:
        doc = Document(caminho_arquivo)
        logger.debug(f"Documento carregado com sucesso: {caminho_arquivo}")
        
        if titulo_documento and hasattr(doc, 'core_properties'):
            doc.core_properties.title = titulo_documento
            logger.debug(f"Título do documento atualizado para: {titulo_documento}")
        
        # 1. Identificar quebras de página no documento
        quebras_pagina = identificar_quebras_pagina(doc)
        logger.info(f"Identificadas {len(quebras_pagina)} quebras de página explícitas no documento")
        
        # 2. Encontrar início do sumário (parágrafo com o título "Sumário")
        sumario_inicio = None
        for i, paragrafo in enumerate(doc.paragraphs):
            if paragrafo.text.strip().lower() == "sumário":
                sumario_inicio = i
                logger.info(f"Título do sumário encontrado no parágrafo {i}")
                break
        
        if sumario_inicio is None:
            # Busca alternativa por parágrafo com estilo TOC_Heading
            for i, paragrafo in enumerate(doc.paragraphs):
                if hasattr(paragrafo, 'style') and paragrafo.style and "TOC" in str(paragrafo.style.name):
                    sumario_inicio = i
                    logger.info(f"Título do sumário encontrado por estilo no parágrafo {i}")
                    break
        
        if sumario_inicio is None:
            # Última tentativa: localizar pelo contexto do documento
            for i, paragrafo in enumerate(doc.paragraphs):
                if paragrafo.text.strip().lower() in ["índice", "conteúdo", "table of contents"]:
                    sumario_inicio = i
                    logger.info(f"Possível título do sumário encontrado no parágrafo {i}: '{paragrafo.text}'")
                    break

        if sumario_inicio is None:
            logger.error("Título do sumário não encontrado no documento!")
            return False
        
        # 3. Encontrar fim do sumário atual
        sumario_fim = None
        
        for i in range(sumario_inicio + 1, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            texto = p.text.strip()
            estilo = p.style.name if hasattr(p, 'style') and p.style and hasattr(p.style, 'name') else ""
            
            if (estilo == 'Title' and texto.startswith("1")) or \
               re.match(r'^1\s*[-–]', texto) or \
               texto.lower().startswith("introdução"):
                sumario_fim = i
                logger.info(f"Introdução encontrada no parágrafo {i}: '{texto}'")
                break
                
        if sumario_fim is None:
            sumario_fim = sumario_inicio + 3
            logger.warning(f"Introdução não encontrada após o sumário. Usando valor relativo {sumario_fim}.")
        
        # 4. Criar mapa de páginas baseado nas quebras de página
        pagina_atual = 1
        mapa_paragrafos_paginas = {}
        
        # Definimos o primeiro parágrafo como página 1
        for i in range(len(doc.paragraphs)):
            # Verifica se há uma quebra de página neste parágrafo
            if i in quebras_pagina:
                pagina_atual += 1
            mapa_paragrafos_paginas[i] = pagina_atual
        
        # 5. Coletar todos os títulos, suas páginas corretas e criar marcadores para links
        marcadores_titulos = {}
        
        # Primeiro, vamos adicionar marcadores (bookmarks) a todos os títulos
        for i, p in enumerate(doc.paragraphs):
            # Ignoramos parágrafos dentro do sumário
            if sumario_inicio <= i < sumario_fim:
                continue
                
            texto = p.text.strip()
            estilo = p.style.name if hasattr(p, 'style') and p.style and hasattr(p.style, 'name') else ""
            
            # Verifica se é um título
            eh_titulo = False
            
            if (estilo == 'Title' or estilo == 'Heading 1') and texto:
                eh_titulo = True
            elif estilo == 'Heading 2' and texto:
                eh_titulo = True
            elif estilo == 'Heading 3' and texto:
                eh_titulo = True
            
            # Se é um título, cria um marcador (bookmark)
            if eh_titulo:
                # Gera um ID único para o título
                import hashlib
                marcador_id = f"titulo_{hashlib.md5(texto.encode()).hexdigest()[:8]}"
                
                # Adiciona um marcador (bookmark) para o título
                from docx.oxml.shared import OxmlElement, qn

                # Cria marcador de início
                start = OxmlElement('w:bookmarkStart')
                start.set(qn('w:id'), str(i))  # Usando o índice como ID único
                start.set(qn('w:name'), marcador_id)
                
                # Cria marcador de fim
                end = OxmlElement('w:bookmarkEnd')
                end.set(qn('w:id'), str(i))
                
                # Insere o marcador no início do parágrafo
                p._p.insert(0, start)
                p._p.append(end)
                
                # Guarda o ID do marcador para usá-lo nos links
                marcadores_titulos[i] = marcador_id
                for tit in titulos:
                    if(tit["texto"] == texto):
                        tit["marcador"] = marcador_id 
                        break
        
        # 6. Remover os parágrafos do sumário antigo EXCETO o título TOC
        logger.info(f"Removendo sumário antigo do parágrafo {sumario_inicio+1} ao {sumario_fim-1}")
        for i in range(sumario_fim - 1, sumario_inicio, -1):
            try:
                p_element = doc.paragraphs[i]._element
                if p_element.getparent() is not None:
                    p_element.getparent().remove(p_element)
                    logger.debug(f"Removido parágrafo {i}")
            except Exception as e:
                logger.warning(f"Erro ao remover parágrafo {i}: {str(e)}")
        
        # 7. Criar itens do sumário com hyperlinks
        logger.info(f"Criando novo sumário com {len(titulos)} títulos")
        
        # Obter o parágrafo de referência (título do sumário)
        paragrafo_referencia = doc.paragraphs[sumario_inicio]
        elemento_referencia = paragrafo_referencia._element
        pai_elemento = elemento_referencia.getparent()
        
        # Obter o índice do elemento de referência entre seus irmãos
        indice_no_pai = pai_elemento.index(elemento_referencia)
        
        # Inserir os novos itens do sumário após o título
        posicao_atual = indice_no_pai
        
        for titulo in titulos:
            try:
                p = doc.add_paragraph()
                
                # Remove o parágrafo de sua posição atual
                novo_elemento = p._element
                novo_elemento.getparent().remove(novo_elemento)
                
                # Insere após a posição atual (que começou com o título do sumário)
                posicao_atual += 1
                pai_elemento.insert(posicao_atual, novo_elemento)
                
                # Formatação do parágrafo
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Pt(12 * (titulo['nivel'] - 1))
                
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Cm(16.99), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
                
                texto_item = ""
                if titulo['nivel'] > 1:
                    texto_item += "  " * (titulo['nivel'] - 1)
                
                if titulo['numero']:
                    texto_item += f"{titulo['numero']} - {titulo['texto']}"
                else:
                    texto_item += titulo['texto']
                
                # Definir o tamanho da fonte com base no nível do título
                # Títulos nível 1: 12pt, negrito
                # Títulos nível 2+: 11pt, normal
                tamanho_fonte = Pt(12) if titulo['nivel'] == 1 else Pt(11)
                negrito = (titulo['nivel'] == 1)
                
                # Cria o hyperlink para o texto do item usando o marcador
                cor_rgb = (51, 70, 130)
                criar_hyperlink_marcador(
                    p, 
                    texto_item, 
                    titulo['marcador'], 
                    cor_rgb=cor_rgb, 
                    tamanho_fonte=tamanho_fonte, 
                    negrito=negrito
                )
                
                # Adiciona o tab
                tab_run = p.add_run("\t")
                
                # Adiciona o número da página como hyperlink também
                criar_hyperlink_marcador(
                    p,
                    str(titulo['pagina']),
                    titulo['marcador'],
                    cor_rgb=cor_rgb,
                    tamanho_fonte=tamanho_fonte,
                    negrito=negrito
                )
                    
                logger.debug(f"Adicionado item ao sumário: {texto_item} (Página {titulo['pagina']})")
            except Exception as e:
                logger.warning(f"Erro ao adicionar item ao sumário: {str(e)}")
                logger.exception("Detalhes do erro:")
        
        # 8. Adicionar quebra de página após o último item do sumário
        logger.info("Adicionando quebra de página após o sumário")
        
        # Criar parágrafo com quebra de página
        p_quebra = doc.add_paragraph()
        
        # Remove o parágrafo da posição atual
        p_quebra_element = p_quebra._element
        p_quebra_element.getparent().remove(p_quebra_element)
        
        # Insere após o último item do sumário
        pai_elemento.insert(posicao_atual + 1, p_quebra_element)
        
        # Adiciona a quebra de página
        run = p_quebra.add_run()
        run._element.append(OxmlElement('w:br'))
        run._element.get_or_add_rPr().append(OxmlElement('w:lastRenderedPageBreak'))
        
        # Configuração adicional da quebra de página para garantir que seja uma quebra de página real
        from docx.enum.text import WD_BREAK
        run.add_break(WD_BREAK.PAGE)
        
        logger.info("Quebra de página adicionada após o sumário")
        
        # 9. Adicionar parágrafo vazio após a quebra de página para manter o padrão de espaçamento
        logger.info("Adicionando parágrafo vazio após a quebra de página")
        
        # Criar parágrafo vazio
        p_espaco = doc.add_paragraph()
        
        # Remove o parágrafo da posição atual
        p_espaco_element = p_espaco._element
        p_espaco_element.getparent().remove(p_espaco_element)
        
        # Insere após o parágrafo de quebra de página
        pai_elemento.insert(posicao_atual + 2, p_espaco_element)
        
        # Configura o parágrafo vazio SEM espaçamento
        p_espaco.paragraph_format.space_before = Pt(0)
        p_espaco.paragraph_format.space_after = Pt(0)
        
        logger.info("Parágrafo vazio adicionado após a quebra de página")
        
        # Salvar o documento
        doc.save(caminho_arquivo)
        logger.info("Documento salvo com sucesso após atualização do sumário")
        
        return True
        
    except Exception as e:
        logger.error(f"Erro ao atualizar sumário: {str(e)}")
        logger.exception("Detalhes do erro:")
        
        if doc:
            try:
                doc.save(caminho_arquivo)
                logger.info("Documento salvo, apesar do erro na atualização do sumário")
            except Exception as save_err:
                logger.error(f"Não foi possível salvar o documento após erro: {str(save_err)}")
            
        return False