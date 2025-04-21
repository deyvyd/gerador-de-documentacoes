# modules/document_processor/requisitos_handlers.py
import logging
import os
import base64
from io import BytesIO
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .formatters import copiar_formatacao_run, copiar_formatacao_bullets_completa
from .table_handlers import (
    tab_mapear_formatacao_paragrafo,
    tab_processar_formatacao_paragrafo,
    tab_copiar_estilo_celula
)

# Configuração de logging
logger = logging.getLogger(__name__)

def processar_requisitos_funcionais(doc, requisitos):
    """
    Processa requisitos funcionais no formato específico do documento Estratégia de Solução.
    
    Este formato tem o cabeçalho fora da tabela e conteúdo detalhado dentro de tabelas.
    Adiciona quebras de página entre requisitos.
    """
    if not requisitos:
        logger.info("Nenhum requisito funcional para processar")
        return
    
    logger.info(f"Processando {len(requisitos)} requisitos funcionais para Estratégia de Solução")
    
    # Identificar o parágrafo de modelo para o cabeçalho do RF e a tabela modelo
    paragrafo_modelo = None
    tabela_modelo = None
    paragrafo_indice = None
    tabela_indice = None
    
    # Buscar parágrafo modelo que contém "RF-[N_RF]: [TITULO_RF]"
    for i, paragrafo in enumerate(doc.paragraphs):
        if "RF-[N_RF]" in paragrafo.text:
            paragrafo_modelo = paragrafo
            paragrafo_indice = i
            logger.info(f"Parágrafo modelo RF encontrado no índice {i}: {paragrafo.text}")
            break
    
    if not paragrafo_modelo:
        logger.error("Parágrafo modelo para cabeçalho de RF não encontrado")
        return
    
    # Buscar tabela modelo logo após o parágrafo modelo
    for i, tabela in enumerate(doc.tables):
        # Verificar se a tabela contém marcadores de RF
        contem_marcadores = False
        for row in tabela.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "[LOCAL_RF]" in p.text or "[USUARIO_RF]" in p.text:
                        contem_marcadores = True
                        break
            if contem_marcadores:
                break
        
        if contem_marcadores:
            tabela_modelo = tabela
            tabela_indice = i
            logger.info(f"Tabela modelo RF encontrada no índice {i}")
            break
    
    if not tabela_modelo:
        logger.error("Tabela modelo para detalhes de RF não encontrada")
        return
    
    # Remover o parágrafo e tabela modelo originais
    paragrafos_a_remover = []
    tabelas_a_remover = []
    
    # Remover apenas depois de processar todos os requisitos para evitar problemas de índices
    paragrafos_a_remover.append(paragrafo_indice)
    tabelas_a_remover.append(tabela_indice)
    
    # Posição de inserção - logo após o parágrafo modelo original
    pos_insercao = paragrafo_indice
    
    # Processar cada requisito e criar elementos correspondentes
    for idx, requisito in enumerate(requisitos):
        # Adicionar quebra de página antes de cada RF, exceto o primeiro
        if idx > 0:
            paragrafo_quebra = doc.add_paragraph()
            run = paragrafo_quebra.add_run()
            run.add_break(WD_BREAK.PAGE)
            
            # Inserir parágrafo de quebra
            doc._body._body.insert(pos_insercao, paragrafo_quebra._p)
            pos_insercao += 1
        
        # 1. Criar parágrafo de cabeçalho para este RF
        novo_paragrafo = doc.add_paragraph()
        
        # Copiar formatação do parágrafo modelo
        copiar_formatacao_bullets_completa(novo_paragrafo, paragrafo_modelo)
        
        # Substituir o texto do marcador com os dados reais
        numero_rf = f"{idx+1:02d}"  # Formatar como 01, 02, etc.
        titulo_rf = requisito.get('tituloRF', '')
        texto_cabecalho = f"RF-{numero_rf}: {titulo_rf}"
        
        run = novo_paragrafo.add_run(texto_cabecalho)
        if paragrafo_modelo.runs:
            copiar_formatacao_run(run, paragrafo_modelo.runs[0])
        
        # Inserir parágrafo de cabeçalho após a quebra de página
        doc._body._body.insert(pos_insercao, novo_paragrafo._p)
        pos_insercao += 1
        
        # 2. Criar tabela para este RF
        nova_tabela = doc.add_table(rows=tabela_modelo.rows.__len__(), cols=tabela_modelo.columns.__len__())
        
        # Copiar formatação e estilos da tabela modelo
        # CORREÇÃO: Verificar atributos antes de copiar
        try:
            if hasattr(tabela_modelo._tbl, 'tblPr'):
                if not hasattr(nova_tabela._tbl, 'tblPr'):
                    nova_tabela._tbl.append(OxmlElement('w:tblPr'))
                for attr in tabela_modelo._tbl.tblPr:
                    nova_tabela._tbl.tblPr.append(attr.copy())
        except Exception as e:
            logger.warning(f"Não foi possível copiar propriedades da tabela: {e}")
        
        # Copiar estilo/formatação para cada célula
        for i, row in enumerate(tabela_modelo.rows):
            for j, cell in enumerate(row.cells):
                if i < len(nova_tabela.rows) and j < len(nova_tabela.rows[i].cells):
                    nova_celula = nova_tabela.rows[i].cells[j]
                    try:
                        # Usar método seguro de cópia que não depende de tcPr
                        tab_copiar_estilo_celula(nova_celula, cell)
                    except Exception as e:
                        logger.warning(f"Erro ao copiar estilo de célula: {e}")
                        
                        # Fallback: Pelo menos copiar o texto
                        for p_idx, p in enumerate(cell.paragraphs):
                            if p_idx < len(nova_celula.paragraphs):
                                nova_celula.paragraphs[p_idx].text = p.text
                            else:
                                nova_celula.add_paragraph(p.text)
        
        # Processar cada célula da tabela para substituir marcadores
        for i, row in enumerate(nova_tabela.rows):
            for j, cell in enumerate(row.cells):
                # Obter a célula correspondente na tabela modelo
                if i < len(tabela_modelo.rows) and j < len(tabela_modelo.rows[i].cells):
                    celula_modelo = tabela_modelo.rows[i].cells[j]
                    
                    # Garantir que cada célula tenha pelo menos um parágrafo
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    
                    # Substituir os marcadores nas células
                    for p_idx, p in enumerate(cell.paragraphs):
                        texto_original = ""
                        
                        # Obter texto do parágrafo modelo correspondente, se existir
                        if p_idx < len(celula_modelo.paragraphs):
                            texto_original = celula_modelo.paragraphs[p_idx].text
                        
                        # Dicionário com as substituições específicas para cada requisito
                        substituicoes_rf = {
                            '[LOCAL_RF]': requisito.get('local', ''),
                            '[USUARIO_RF]': requisito.get('usuario', ''),
                            '[PERFIL_RF]': requisito.get('perfil', ''),
                            '[TIPO_RF]': requisito.get('tipo', ''),
                            '[DESCRICAO_RF]': requisito.get('descricao', ''),
                            '[REGRAS_CAMPOS_RF]': requisito.get('validacoes', ''),
                            '[REGRAS_NEGOCIO_RF]': requisito.get('regras', ''),
                            '[BANCO_RF]': requisito.get('banco', '')
                        }
                        
                        # Substituir marcadores
                        texto_substituido = texto_original
                        for marcador, valor in substituicoes_rf.items():
                            if marcador in texto_substituido:
                                texto_substituido = texto_substituido.replace(marcador, valor)
                        
                        # Simples substitução de texto para evitar erros
                        p.text = texto_substituido
                        
                        # Se tiver o marcador de imagem, processar separadamente
                        if '[IMAGEM_RF]' in texto_original and requisito.get('imagens'):
                            # Limpar o texto para evitar mostrar o marcador
                            p.text = ""
                            processar_imagens_requisito(p, requisito['imagens'])
        
        # Inserir a nova tabela
        doc._body._body.insert(pos_insercao, nova_tabela._tbl)
        pos_insercao += 1
    
    # Remover os elementos modelo do documento
    # Fazemos isso em ordem reversa para evitar problemas com índices
    for tabela_idx in sorted(tabelas_a_remover, reverse=True):
        if tabela_idx < len(doc.tables):  # Verificar se o índice ainda é válido
            table_element = doc.tables[tabela_idx]._tbl
            if table_element.getparent() is not None:
                doc._body._body.remove(table_element)
    
    for paragrafo_idx in sorted(paragrafos_a_remover, reverse=True):
        if paragrafo_idx < len(doc.paragraphs):  # Verificar se o índice ainda é válido
            paragraph_element = doc.paragraphs[paragrafo_idx]._p
            if paragraph_element.getparent() is not None:
                doc._body._body.remove(paragraph_element)

def processar_imagens_requisito(paragrafo, imagens):
    """
    Processa as imagens do requisito e as adiciona no parágrafo especificado.
    """
    if not imagens:
        return
    
    try:
        # Adicionar apenas a primeira imagem por enquanto
        img_data = imagens[0]
        
        # Se a imagem está em formato base64, decodificá-la
        if isinstance(img_data, str) and img_data.startswith('data:image'):
            try:
                # Extrair a parte base64 da string (remover o prefixo "data:image/...")
                base64_data = img_data.split(',')[1]
                image_stream = BytesIO(base64.b64decode(base64_data))
                
                # Adicionar a imagem ao parágrafo
                run = paragrafo.add_run()
                run.add_picture(image_stream, width=Inches(6))  # Ajustar largura conforme necessário
                
            except Exception as e:
                logger.error(f"Erro ao processar imagem base64: {e}")
                # Adicionar um texto de placeholder em caso de erro
                paragrafo.text = "[Imagem não pôde ser processada]"
        else:
            logger.warning("Formato de imagem não suportado. Esperado: base64")
            paragrafo.text = "[Formato de imagem não suportado]"
            
    except Exception as e:
        logger.error(f"Erro geral ao processar imagem: {e}")
        paragrafo.text = "[Erro ao processar imagem]"