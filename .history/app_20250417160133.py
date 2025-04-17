# Imports
## Imports da Flask
import json
import logging
import os
import re
import shutil
import tempfile
import zipfile
from copy import deepcopy
## Imports da biblioteca padrão Python
from datetime import date, datetime
from urllib.parse import quote

import docx.opc.constants
import docx.oxml.shared
## Imports relacionados ao python-docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from flask import (Flask, render_template, request, send_file,
                   send_from_directory, jsonify)
## Imports de terceiros
from unidecode import unidecode

# Configuração inicial
## Configuração do logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

## Inicialização da aplicação Flask
app = Flask(__name__)

# Classes
class RelatorioAcompanhamentoProjeto:
    def __init__(self, numero_ss, ano_ss, iniciais_autor, titulo_ss, descricao,
                 data_inicio, data_fim, total_horas, link_board, atividades=None):
        self.numero_ss = numero_ss
        self.ano_ss = ano_ss
        self.iniciais_autor = iniciais_autor
        self.titulo_ss = titulo_ss
        self.descricao = descricao
        self.data_inicio = datetime.strptime(data_inicio, '%Y-%m-%d').date()
        self.data_fim = datetime.strptime(data_fim, '%Y-%m-%d').date()
        self.total_horas = int(total_horas)
        self.link_board = link_board 
        self.atividades = atividades or []
        
        # Calcular valores derivados
        self.dias_uteis = self._calcular_dias_uteis()
        self.n_pf = self._calcular_pontos_funcao()
        
    def _calcular_dias_uteis(self):
        dias = (self.data_fim - self.data_inicio).days + 1
        dias_uteis = sum(1 for day in range(dias) 
                        if (self.data_inicio + date.resolution * day).weekday() < 5)
        return dias_uteis
    
    def _calcular_pontos_funcao(self):
        """
        Calcula os pontos de função baseado no total de horas
        Fórmula: (total_horas / 10) * (250 / 100)
        """
        try:
            return round((self.total_horas / 10) * (250 / 100), 2)
        except Exception as e:
            logger.error(f"Erro ao calcular pontos de função: {e}")
            return 0
    
    def get_substituicoes(self):
        """Retorna o dicionário de substituições para o documento"""
        data_atual = datetime.now().strftime('%d/%m/%Y')
        
        return {
            '[NNN]': str(self.numero_ss).zfill(3),
            '[AAAA]': str(self.ano_ss),
            '[INICIAIS_AUTOR]': str(self.iniciais_autor),
            '[TITULO]': str(self.titulo_ss),
            '[DESCRICAO]': str(self.descricao),
            '[DATA_ATUAL]': data_atual,
            '[DATA_INICIO]': self.data_inicio.strftime('%d/%m/%Y'),
            '[DATA_FIM]': self.data_fim.strftime('%d/%m/%Y'),
            '[TOTAL_HORAS]': str(self.total_horas),
            '[DIAS_UTEIS]': str(self.dias_uteis),
            '[N_PF]': str(self.n_pf),
            '[LINK_BOARD]': str(self.link_board)
        }

# Funções Auxiliares

## Funções de formatação
def configurar_fonte_padrao(run):
    """Configura a fonte como Calibri 12"""
    from docx.shared import Pt
    if hasattr(run, 'font'):
        run.font.name = 'Calibri'
        run.font.size = Pt(12)

def copiar_formatacao_run(run_destino, run_origem):
    """Copia a formatação de um run para outro"""
    try:
        if not hasattr(run_origem, '_element') or not hasattr(run_destino, '_element'):
            return
            
        # Copiar propriedades da fonte
        if hasattr(run_origem, 'font') and hasattr(run_destino, 'font'):
            fonte_origem = run_origem.font
            fonte_destino = run_destino.font
            
            # Copiar propriedades específicas
            if hasattr(fonte_origem, 'name'):
                fonte_destino.name = fonte_origem.name
            if hasattr(fonte_origem, 'size'):
                fonte_destino.size = fonte_origem.size
            if hasattr(fonte_origem, 'bold'):
                fonte_destino.bold = fonte_origem.bold
            if hasattr(fonte_origem, 'italic'):
                fonte_destino.italic = fonte_origem.italic
            if hasattr(fonte_origem, 'underline'):
                fonte_destino.underline = fonte_origem.underline
            if hasattr(fonte_origem, 'color') and fonte_origem.color:
                fonte_destino.color.rgb = fonte_origem.color.rgb

        # Copiar XML direto
        if hasattr(run_origem._element, 'rPr') and run_origem._element.rPr is not None:
            if not hasattr(run_destino._element, 'rPr'):
                run_destino._element.get_or_add_rPr()
            
            # Limpar propriedades existentes
            for child in list(run_destino._element.rPr):
                run_destino._element.rPr.remove(child)
                
            # Copiar todas as propriedades do original
            for prop in run_origem._element.rPr.xpath('*'):
                new_prop = deepcopy(prop)
                run_destino._element.rPr.append(new_prop)
    except Exception as e:
        # Tratamento básico de erro sem logging
        print(f"Erro ao copiar formatação: {e}")

## Funções de processamento de documentos
def identificar_tipo_documento(nome_arquivo):
    """Identifica o tipo de documento baseado no nome"""
    from docx.document import Document as DocxDocument

    if isinstance(nome_arquivo, DocxDocument):
        try:
            nome_arquivo = nome_arquivo.core_properties.title or ""
        except:
            nome_arquivo = ""
    
    nome_normalizado = unidecode(str(nome_arquivo).lower())
    logger.debug(f"Nome original: {nome_arquivo}")
    logger.debug(f"Nome normalizado: {nome_normalizado}")
    
    tipo = None
    if "estimativa" in nome_normalizado:
        tipo = 'estimativa'
        logger.info(f"Documento identificado como Estimativa de Esforço e Cronograma: {nome_arquivo}")
    elif "estrategia" in nome_normalizado or "solucao" in nome_normalizado:
        tipo = 'estrategia'
        logger.info(f"Documento identificado como Estratégia da Solução: {nome_arquivo}")
    elif "relatorio" in nome_normalizado:
        tipo = 'relatorio'
        logger.info(f"Documento identificado como Relatório de Acompanhamento de Projeto: {nome_arquivo}")
    else:
        logger.warning(f"Tipo de documento não identificado para o arquivo: {nome_arquivo}")
    
    return tipo


def copiar_formatacao_bullets_completa(paragrafo_destino, paragrafo_origem):
    """
    Copia formatação completa de um parágrafo para outro, com verificações de segurança.
    """
    try:
        # Verificação inicial de segurança
        if not hasattr(paragrafo_destino, '_p') or not hasattr(paragrafo_origem, '_p'):
            logger.warning("Parágrafo de destino ou origem inválido")
            return

        # Copia o estilo do parágrafo, se existir
        try:
            if hasattr(paragrafo_origem, 'style'):
                paragrafo_destino.style = paragrafo_origem.style
        except Exception as style_error:
            logger.warning(f"Erro ao copiar estilo: {style_error}")

        # Copia as propriedades do parágrafo (pPr) com verificações extras
        try:
            # Obtém as propriedades do parágrafo original com verificação
            pPr_origem = paragrafo_origem._p.get_or_add_pPr()
            
            if pPr_origem is not None:
                # Remove propriedades existentes no parágrafo de destino
                pPr_destino = paragrafo_destino._p.get_or_add_pPr()
                
                # Limpa propriedades existentes
                for child in list(pPr_destino):
                    pPr_destino.remove(child)
                
                # Copia todas as propriedades XML do parágrafo original
                for prop in pPr_origem:
                    new_prop = deepcopy(prop)
                    pPr_destino.append(new_prop)

        except Exception as pPr_error:
            logger.warning(f"Erro ao copiar propriedades do parágrafo: {pPr_error}")

    except Exception as e:
        logger.error(f"Erro geral ao copiar formatação do parágrafo: {e}")

def processar_atividades_bullet_points(doc, atividades):
    """
    Localiza o marcador [ITEM] no documento e o substitui pelas atividades fornecidas,
    mantendo a formatação original e adicionando apenas uma quebra de página sem texto.
    """
    # Validação inicial dos dados de entrada
    if not atividades:
        logger.warning("Nenhuma atividade fornecida para substituição. O documento não será modificado.")
        return
    
    # Preparação para localizar o marcador [ITEM]
    paragrafo_referencia = None
    indice_elemento = None
    body = doc._body._body
    elementos = list(body)
    
    # Localizar o parágrafo com [ITEM] e o parágrafo de quebra de página
    paragrafo_quebra = None
    
    for i, paragrafo in enumerate(doc.paragraphs):
        if '[ITEM]' in paragrafo.text:
            paragrafo_referencia = paragrafo
            indice_elemento = elementos.index(paragrafo._element)
            logger.info(f"Marcador [ITEM] encontrado no documento, índice {i}")
            
            # Verifica se o próximo parágrafo contém quebra de página
            if i+1 < len(doc.paragraphs):
                # Verificamos se o próximo parágrafo contém uma quebra de página
                for run in doc.paragraphs[i+1].runs:
                    if hasattr(run, '_element') and 'br' in str(run._element.xml):
                        paragrafo_quebra = doc.paragraphs[i+1]
                        logger.info(f"Parágrafo com quebra de página encontrado no índice {i+1}")
                        break
            break
    
    # Verificação de segurança - se não encontrou o marcador
    if paragrafo_referencia is None:
        logger.error("Marcador [ITEM] não encontrado no documento. Nenhuma modificação será realizada.")
        return
    
    # Capturar informações do parágrafo de quebra de página (se existir)
    quebra_elemento = None
    if paragrafo_quebra:
        quebra_elemento = paragrafo_quebra._element
        indice_quebra = elementos.index(quebra_elemento)
        logger.info(f"Elemento de quebra encontrado no índice {indice_quebra}")
    
    try:
        # Lista para armazenar os parágrafos criados
        paragrafos_criados = []
        
        # Processa cada atividade
        for idx, atividade in enumerate(atividades):
            # Criação de um novo parágrafo
            novo_paragrafo = doc.add_paragraph()
            
            # Remove o novo parágrafo do documento (para manipulação)
            body.remove(novo_paragrafo._element)
            
            # Preserva a formatação do parágrafo original
            copiar_formatacao_bullets_completa(novo_paragrafo, paragrafo_referencia)
            
            # Adiciona o texto da atividade
            novo_run = novo_paragrafo.add_run(atividade['nome'])
            if paragrafo_referencia.runs:
                copiar_formatacao_run(novo_run, paragrafo_referencia.runs[0])
            
            # Armazena o novo parágrafo
            paragrafos_criados.append(novo_paragrafo._element)
            logger.info(f"Criado parágrafo para atividade: {atividade['nome']}")
            
            # Se for o último item, adiciona a quebra de página diretamente a ele
            if idx == len(atividades) - 1:
                from docx.enum.text import WD_BREAK
                novo_run.add_break(WD_BREAK.PAGE)
                logger.info("Adicionada quebra de página ao último item da lista")
        
        # Remove o parágrafo de referência ([ITEM]) e quebra de página (se existir)
        body.remove(paragrafo_referencia._element)
        if quebra_elemento is not None and quebra_elemento.getparent() is not None:
            body.remove(quebra_elemento)
        
        # Insere os novos parágrafos de atividades
        for i, paragrafo in enumerate(paragrafos_criados):
            body.insert(indice_elemento + i, paragrafo)
        
        logger.info(f"Adicionados {len(paragrafos_criados)} itens com quebra de página no último item")
        
    except Exception as e:
        logger.error(f"Erro durante a substituição das atividades: {e}")
        logger.exception("Detalhes do erro:")
        raise

def tab_copiar_estilo_celula(celula_destino, celula_origem):
    """
    Copia a formatação completa de uma célula para outra, mantendo a formatação
    específica de cada parte do texto exatamente na posição correta.
    """
    try:
        # Primeiro capturamos o alinhamento original
        alinhamento_origem = None
        if celula_origem.paragraphs:
            alinhamento_origem = celula_origem.paragraphs[0].alignment
            logger.debug(f"Alinhamento original capturado: {alinhamento_origem}")

        # Copia propriedades da célula
        if hasattr(celula_origem._tc, 'tcPr'):
            tcPr_origem = celula_origem._tc.tcPr
            tcPr_destino = celula_destino._tc.get_or_add_tcPr()
            
            # Limpa propriedades existentes
            for child in list(tcPr_destino):
                tcPr_destino.remove(child)
            
            # Copia propriedades da célula
            for prop in tcPr_origem.xpath('*'):
                new_prop = deepcopy(prop)
                tcPr_destino.append(new_prop)

        # Trata os parágrafos
        if celula_origem.paragraphs and celula_destino.paragraphs:
            p_origem = celula_origem.paragraphs[0]
            p_destino = celula_destino.paragraphs[0]
            
            # Aplica o alinhamento capturado anteriormente
            if alinhamento_origem is not None:
                p_destino.alignment = alinhamento_origem
                logger.debug(f"Alinhamento aplicado: {alinhamento_origem}")
            
            # Copia propriedades específicas do parágrafo
            if hasattr(p_origem._p, 'pPr') and p_origem._p.pPr is not None:
                pPr_origem = p_origem._p.pPr
                pPr_destino = p_destino._p.get_or_add_pPr()
                
                # Limpa propriedades existentes mantendo o alinhamento
                alinhamento_temp = None
                if hasattr(pPr_destino, 'jc'):
                    alinhamento_temp = deepcopy(pPr_destino.jc)
                
                for child in list(pPr_destino):
                    pPr_destino.remove(child)
                
                # Restaura o alinhamento se existia
                if alinhamento_temp is not None:
                    pPr_destino.append(alinhamento_temp)
                
                # Copia outras propriedades do parágrafo
                for prop in pPr_origem.xpath('*'):
                    if prop.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc':
                        new_prop = deepcopy(prop)
                        pPr_destino.append(new_prop)

            # Obtém o texto completo da célula destino
            texto_destino = celula_destino.text.strip()
            palavras_destino = texto_destino.split()
            
            # Remove runs existentes mantendo o alinhamento
            alinhamento_atual = p_destino.alignment
            for run in p_destino.runs:
                run._element.getparent().remove(run._element)
            
            # Restaura o alinhamento
            p_destino.alignment = alinhamento_atual
            
            # Processa o texto preservando a formatação
            if texto_destino:
                run_destino = p_destino.add_run(texto_destino)
                if p_origem.runs:
                    copiar_formatacao_run(run_destino, p_origem.runs[0])

            logger.info("Formatação da célula copiada com sucesso, mantendo posicionamento correto")
            logger.debug(f"Alinhamento final: {p_destino.alignment}")

    except Exception as e:
        logger.error(f"Erro ao copiar estilo da célula: {e}")
        raise

def tab_mapear_formatacao_paragrafo(paragraph):
    """
    Cria um mapa da formatação de cada parte do texto em um parágrafo.
    
    Esta função analisa cada 'run' (segmento de texto) no parágrafo e registra
    sua posição e formatação. Isso nos permite replicar a mesma formatação
    ao substituir o texto posteriormente.
    
    Args:
        paragraph: O parágrafo do documento Word a ser analisado
        
    Returns:
        list: Lista de dicionários contendo a informação de formatação de cada parte do texto
    """
    mapa_formatacao = []
    posicao_atual = 0
    
    # Analisamos cada run (segmento de texto) no parágrafo
    for run in paragraph.runs:
        # Só mapeamos se o run contiver texto
        comprimento = len(run.text)
        if comprimento > 0:
            mapa_formatacao.append({
                'inicio': posicao_atual,
                'fim': posicao_atual + comprimento,
                'run': run  # Guardamos o run original para copiar sua formatação
            })
            posicao_atual += comprimento
            
    return mapa_formatacao

def tab_processar_formatacao_paragrafo(paragraph, texto_novo, mapa_formatacao):
    """
    Processa e aplica formatação completa a um parágrafo, incluindo alinhamento e estilos.
    
    Esta função unifica a obtenção do alinhamento e aplicação de formatação, cuidando de:
    1. Alinhamento do parágrafo (direto, XML e herança de células)
    2. Propriedades do parágrafo (espaçamento, recuo)
    3. Formatação do texto (fonte, tamanho, cor, estilos)
    
    Args:
        paragraph: O parágrafo a ser processado
        texto_novo: O novo texto a ser inserido
        mapa_formatacao: Mapa de formatação do texto original
        
    Returns:
        None
    """
    try:
        # Parte 1: Obtenção do alinhamento
        alinhamento_real = None
        
        # 1.1 Verifica alinhamento direto
        if paragraph.alignment is not None:
            logger.debug(f"Alinhamento encontrado diretamente: {paragraph.alignment}")
            alinhamento_real = paragraph.alignment
            
        # 1.2 Verifica alinhamento nas propriedades XML
        if alinhamento_real is None and hasattr(paragraph._p, 'pPr'):
            pPr = paragraph._p.pPr
            if pPr is not None:
                jc_element = pPr.xpath('.//w:jc')
                if jc_element:
                    val = jc_element[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    logger.debug(f"Alinhamento encontrado no XML: {val}")
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    alignment_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    alinhamento_real = alignment_map.get(val)
        
        # 1.3 Verifica alinhamento na célula da tabela
        if alinhamento_real is None:
            cell = paragraph._p.getparent()
            while cell is not None:
                if cell.tag.endswith('tc'):
                    tcPr = cell.xpath('.//w:tcPr')
                    if tcPr:
                        align_elements = tcPr[0].xpath('.//w:vAlign')
                        if align_elements:
                            val = align_elements[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            logger.debug(f"Alinhamento encontrado na célula: {val}")
                            from docx.enum.text import WD_ALIGN_PARAGRAPH
                            cell_alignment_map = {
                                'center': WD_ALIGN_PARAGRAPH.CENTER,
                                'left': WD_ALIGN_PARAGRAPH.LEFT,
                                'right': WD_ALIGN_PARAGRAPH.RIGHT
                            }
                            alinhamento_real = cell_alignment_map.get(val)
                    break
                cell = cell.getparent()

        # Parte 2: Preservação e aplicação da formatação
        # 2.1 Salva propriedades do parágrafo
        estilo_paragrafo = None
        if hasattr(paragraph._p, 'pPr') and paragraph._p.pPr is not None:
            estilo_paragrafo = deepcopy(paragraph._p.pPr)
        
        # 2.2 Remove runs existentes
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
        
        # 2.3 Aplica o novo texto com formatação
        if not mapa_formatacao:
            run = paragraph.add_run(texto_novo)
        else:
            texto_original_len = mapa_formatacao[-1]['fim']
            proporcao = len(texto_novo) / texto_original_len if texto_original_len > 0 else 1
            
            for formato in mapa_formatacao:
                novo_inicio = int(formato['inicio'] * proporcao)
                novo_fim = int(formato['fim'] * proporcao)
                novo_fim = min(novo_fim, len(texto_novo))
                
                if novo_inicio < novo_fim:
                    texto_parte = texto_novo[novo_inicio:novo_fim]
                    run_novo = paragraph.add_run(texto_parte)
                    try:
                        copiar_formatacao_run(run_novo, formato['run'])
                    except Exception as e:
                        logger.warning(f"Erro ao copiar formatação do run: {e}")
                        configurar_fonte_padrao(run_novo)
        
        # 2.4 Restaura alinhamento e propriedades
        if alinhamento_real is not None:
            paragraph.alignment = alinhamento_real
            logger.debug(f"Alinhamento restaurado: {alinhamento_real}")
        
        if estilo_paragrafo is not None:
            if hasattr(paragraph._p, 'pPr'):
                for child in list(paragraph._p.pPr):
                    paragraph._p.pPr.remove(child)
                    
            for prop in estilo_paragrafo.xpath('*'):
                paragraph._p.get_or_add_pPr().append(deepcopy(prop))

    except Exception as e:
        logger.error(f"Erro ao processar formatação: {e}")
        # Tratamento de erro: mantém pelo menos o texto e alinhamento básico
        run = paragraph.add_run(texto_novo)
        if alinhamento_real is not None:
            paragraph.alignment = alinhamento_real

def processar_atividades_tabela(doc, atividades):
    """
    Processa uma tabela do Word para inserir atividades mantendo a formatação original.
    
    Esta função realiza as seguintes operações:
    1. Localiza a tabela que contém os marcadores [ITEM] e [HORAS_ITEM]
    2. Preserva o cabeçalho da tabela
    3. Substitui os marcadores pelos dados da primeira atividade
    4. Adiciona linhas adicionais para as demais atividades
    5. Mantém a linha de totais no final com sua formatação original
    
    Args:
        doc: Objeto Document do python-docx contendo a tabela
        atividades: Lista de dicionários, cada um com 'nome' e 'horas' da atividade
    """
    try:
        # Vamos localizar a tabela específica que contém os marcadores
        tabela_alvo = None
        linha_item = None
        linha_total = None

        # Iteramos por todas as tabelas do documento
        for table in doc.tables:
            for idx_row, row in enumerate(table.rows):
                texto_linha = ' '.join(cell.text for cell in row.cells)
                if '[ITEM]' in texto_linha:
                    tabela_alvo = table
                    linha_item = idx_row
                elif 'Total' in texto_linha and tabela_alvo == table:
                    linha_total = idx_row
                    break
            if tabela_alvo and linha_total is not None:
                break

        if not tabela_alvo:
            raise ValueError("Não foi possível encontrar a tabela com os marcadores necessários")

        logger.info(f"Tabela encontrada: linha do item={linha_item}, linha do total={linha_total}")

        # Preservamos a linha de totais removendo-a temporariamente
        linha_total_row = tabela_alvo.rows[linha_total]
        linha_total_elemento = linha_total_row._tr
        tabela_alvo._tbl.remove(linha_total_elemento)

        # Processamos a primeira atividade (substituindo os marcadores)
        primeira_linha = tabela_alvo.rows[linha_item]
        
        # Substituímos os marcadores na primeira linha
        for cell in primeira_linha.cells:
            for paragraph in cell.paragraphs:
                texto = paragraph.text
                if '[ITEM]' in texto:
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(paragraph)
                    tab_processar_formatacao_paragrafo(paragraph, atividades[0]['nome'], mapa_formatacao)
                elif '[HORAS_ITEM]' in texto:
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(paragraph)
                    tab_processar_formatacao_paragrafo(paragraph, str(atividades[0]['horas']), mapa_formatacao)

        # Para cada atividade adicional, criamos uma nova linha
        for atividade in atividades[1:]:
            nova_linha = tabela_alvo.add_row()
            
            # Copiamos a formatação da primeira linha para a nova
            for idx, (cell_destino, cell_origem) in enumerate(zip(nova_linha.cells, primeira_linha.cells)):
                # Copiamos o estilo da célula usando a função existente
                tab_copiar_estilo_celula(cell_destino, cell_origem)
                
                # Substituímos o conteúdo mantendo a formatação
                if idx == 0:  # Coluna do nome da atividade 
                    p_destino = cell_destino.paragraphs[0]
                    p_origem = cell_origem.paragraphs[0]
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(p_origem)
                    tab_processar_formatacao_paragrafo(p_destino, atividade['nome'], mapa_formatacao)
                elif idx == 1:  # Coluna das horas
                    p_destino = cell_destino.paragraphs[0]
                    p_origem = cell_origem.paragraphs[0]
                    mapa_formatacao = tab_mapear_formatacao_paragrafo(p_origem)
                    tab_processar_formatacao_paragrafo(p_destino, str(atividade['horas']), mapa_formatacao)

        # Recolocamos a linha de totais no final
        tabela_alvo._tbl.append(linha_total_elemento)

        logger.info("Processamento das atividades na tabela concluído com sucesso")

    except Exception as e:
        logger.error(f"Erro ao processar atividades na tabela: {e}")
        logger.exception("Detalhes do erro:")
        raise

def criar_hyperlink(paragraph, url, text, run_modelo=None):
    """
    Cria um hyperlink em um parágrafo do documento Word mantendo a formatação original.
    
    Args:
        paragraph: O parágrafo onde o link será criado
        url: A URL do link
        text: O texto que será exibido para o link
        run_modelo: Run de referência para copiar a formatação
    """
    # Cria o relacionamento da URL
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Cria o elemento hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Cria um novo run
    new_run = OxmlElement('w:r')
    
    # Adiciona as propriedades do run
    rPr = OxmlElement('w:rPr')
    
    # Se temos um run modelo, copiamos suas propriedades de fonte
    if run_modelo and hasattr(run_modelo, '_element') and hasattr(run_modelo._element, 'rPr'):
        modelo_rPr = run_modelo._element.rPr
        
        # Copia o nome da fonte
        if modelo_rPr.xpath('.//w:rFonts'):
            for rFonts in modelo_rPr.xpath('.//w:rFonts'):
                new_rFonts = deepcopy(rFonts)
                rPr.append(new_rFonts)
        
        # Copia o tamanho da fonte
        if modelo_rPr.xpath('.//w:sz'):
            for sz in modelo_rPr.xpath('.//w:sz'):
                new_sz = deepcopy(sz)
                rPr.append(new_sz)
    
    # Adiciona o estilo de hyperlink
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    # Adiciona a cor azul
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    # Adiciona o sublinhado
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    # Adiciona o texto
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)
    return hyperlink

def substituir_texto_preservando_formatacao(paragraph, substituicoes):
    """
    Substitui texto em um parágrafo preservando formatações e tratando links especialmente.
    Mantém quebras de linha simples sem adicionar linhas vazias extras.
    """
    try:
        texto_original = ''.join(run.text for run in paragraph.runs)
        
        # Tratamento especial para o link do board
        if '[LINK_BOARD]' in texto_original:
            url = substituicoes.get('[LINK_BOARD]', '')
            
            # Se o URL estiver vazio, retornamos sem fazer nada, pois o parágrafo
            # já terá sido removido pela função remover_paragrafo_link_board
            if not url or url.strip() == '':
                return False
                
            run_modelo = paragraph.runs[0] if paragraph.runs else None
            
            # Textos padrão para o link
            texto_antes = "Para maior detalhamento do fluxo das atividades basta acessar o "
            texto_depois = " com o seu login e senha da plataforma do GitLab. Caso seu usuário não tenha acesso a esse board, por favor solicitar acesso a LogAp."
            
            # Limpa o parágrafo mantendo alinhamento
            alinhamento_atual = paragraph.alignment
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
            paragraph.alignment = alinhamento_atual
            
            # Adiciona texto antes do link
            run_antes = paragraph.add_run(texto_antes)
            if run_modelo:
                copiar_formatacao_run(run_antes, run_modelo)
            
            # Cria e adiciona o hyperlink
            hyperlink = criar_hyperlink(paragraph, url, url, run_modelo)
            paragraph._p.append(hyperlink)
            
            # Adiciona texto depois do link
            run_depois = paragraph.add_run(texto_depois)
            if run_modelo:
                copiar_formatacao_run(run_depois, run_modelo)
            
            return True
            
        # Se não há marcadores para substituir, retorna
        if not any(marcador in texto_original for marcador in substituicoes.keys()):
            return False
            
        # Detecta se estamos processando o marcador de descrição
        processando_descricao = '[DESCRICAO]' in texto_original
        
        # Mapeia cada run com sua posição exata e formatação
        mapa_formatacao = []
        posicao = 0
        
        for run in paragraph.runs:
            if run.text:
                mapa_formatacao.append({
                    'inicio': posicao,
                    'fim': posicao + len(run.text),
                    'texto': run.text,
                    'run': run
                })
                posicao += len(run.text)
        
        # Faz todas as substituições necessárias
        novo_texto = texto_original
        substituicoes_realizadas = []
        
        for marcador, substituicao in substituicoes.items():
            if marcador in novo_texto and marcador != '[LINK_BOARD]':
                # Para a descrição, garantimos que as quebras de linha sejam simples
                if marcador == '[DESCRICAO]' and isinstance(substituicao, str):
                    # Normaliza todos os tipos de quebras de linha para \n
                    substituicao = substituicao.replace('\r\n', '\n').replace('\r', '\n')
                    # Evita quebras de linha duplas que causariam linhas vazias
                    substituicao = substituicao.replace('\n\n', '\n')
                
                # Guarda as posições antes da substituição
                pos_inicio = novo_texto.find(marcador)
                while pos_inicio != -1:
                    substituicoes_realizadas.append({
                        'inicio_original': pos_inicio,
                        'fim_original': pos_inicio + len(marcador),
                        'texto_novo': substituicao,
                        'tamanho_original': len(marcador),
                        'tamanho_novo': len(substituicao)
                    })
                    pos_inicio = novo_texto.find(marcador, pos_inicio + 1)
                
                novo_texto = novo_texto.replace(marcador, substituicao)
        
        # Remove runs existentes mantendo alinhamento
        alinhamento_atual = paragraph.alignment
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
        paragraph.alignment = alinhamento_atual
        
        # Para o texto da descrição, dividimos por quebras de linha e adicionamos cada parte como um run separado
        if processando_descricao and '\n' in novo_texto:
            linhas = novo_texto.split('\n')
            for i, linha in enumerate(linhas):
                run = paragraph.add_run(linha)
                # Aplicamos a formatação do primeiro run do parágrafo original
                if mapa_formatacao and mapa_formatacao[0]['run']:
                    copiar_formatacao_run(run, mapa_formatacao[0]['run'])
                
                # Adicionamos quebra de linha após cada linha, exceto a última
                if i < len(linhas) - 1:
                    run.add_break()  # Adiciona quebra de linha simples
        else:
            # Reconstrói o texto mantendo as formatações exatas
            posicao_atual = 0
            while posicao_atual < len(novo_texto):
                # Encontra qual run original corresponde a esta posição
                run_correspondente = None
                ajuste_posicao = 0
                
                # Ajusta a posição considerando as substituições realizadas
                posicao_ajustada = posicao_atual
                for subst in substituicoes_realizadas:
                    if posicao_atual >= subst['inicio_original']:
                        ajuste_posicao += subst['tamanho_novo'] - subst['tamanho_original']
                        posicao_ajustada = posicao_atual - ajuste_posicao
                
                # Encontra o run original correspondente à posição ajustada
                for formato in mapa_formatacao:
                    if formato['inicio'] <= posicao_ajustada < formato['fim']:
                        run_correspondente = formato['run']
                        break
                
                # Define quanto texto processar neste run
                if run_correspondente:
                    # Encontra o próximo ponto de quebra
                    proximo_ponto = len(novo_texto)
                    for formato in mapa_formatacao:
                        if formato['inicio'] > posicao_ajustada:
                            proximo_ponto = min(proximo_ponto, 
                                            formato['inicio'] + ajuste_posicao)
                    
                    texto_parte = novo_texto[posicao_atual:proximo_ponto]
                    
                    if texto_parte:
                        novo_run = paragraph.add_run(texto_parte)
                        copiar_formatacao_run(novo_run, run_correspondente)
                    
                    posicao_atual = proximo_ponto
                else:
                    # Se não encontrou formatação, usa a do último run
                    texto_parte = novo_texto[posicao_atual:]
                    novo_run = paragraph.add_run(texto_parte)
                    if mapa_formatacao:
                        copiar_formatacao_run(novo_run, mapa_formatacao[-1]['run'])
                    break
        
        return True
        
    except Exception as e:
        logger.error(f"Erro ao substituir texto preservando formatação: {e}")
        try:
            paragraph.text = texto_original
        except:
            pass
        return False
    
def remover_paragrafo_link_board(doc):
    """
    Remove parágrafos que contêm o marcador [LINK_BOARD] quando o link não foi fornecido
    e adiciona uma única quebra de página explícita no lugar.
    
    Args:
        doc: Objeto Document do python-docx
    """
    from docx.enum.text import WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Lista de parágrafos que contêm o marcador e devem ser removidos
    paragrafos_para_remover = []
    indices_paragrafos = []
    
    # Primeiro, identificamos os parágrafos que contêm o marcador
    for i, paragrafo in enumerate(doc.paragraphs):
        if '[LINK_BOARD]' in paragrafo.text:
            paragrafos_para_remover.append(paragrafo._element)
            indices_paragrafos.append(i)
            logger.info(f"Marcado parágrafo com [LINK_BOARD] para remoção: índice {i}")
    
    # Remove os parágrafos identificados e adiciona quebra de página
    for elemento, indice in zip(paragrafos_para_remover, indices_paragrafos):
        if elemento.getparent() is not None:
            # Obtém o elemento pai e o índice do elemento no pai
            pai = elemento.getparent()
            indice_no_pai = pai.index(elemento)
            
            # Remove o elemento
            pai.remove(elemento)
            logger.info("Parágrafo com [LINK_BOARD] removido com sucesso")
            
            try:
                # Cria um novo parágrafo com quebra de página
                novo_paragrafo = doc.add_paragraph()
                novo_elemento = novo_paragrafo._element
                
                # Remove o parágrafo da posição atual
                if novo_elemento.getparent() is not None:
                    novo_elemento.getparent().remove(novo_elemento)
                
                # Adiciona o novo parágrafo na posição do removido
                pai.insert(indice_no_pai, novo_elemento)
                
                # Adiciona APENAS UMA quebra de página ao novo parágrafo
                # Método 1: Usando o método add_break (mais compatível)
                run = novo_paragrafo.add_run()
                run.add_break(WD_BREAK.PAGE)
                
                # Não adicionar os elementos XML adicionais que estavam causando 
                # múltiplas quebras de página
                
                logger.info("Uma única quebra de página adicionada no lugar do parágrafo removido")
            except Exception as e:
                logger.error(f"Erro ao adicionar quebra de página: {str(e)}")
    
    # Também verifica nas tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                paragrafos_para_remover = []
                for i, paragrafo in enumerate(celula.paragraphs):
                    if '[LINK_BOARD]' in paragrafo.text:
                        paragrafos_para_remover.append(paragrafo._element)
                        logger.info(f"Marcado parágrafo com [LINK_BOARD] em tabela para remoção")
                
                for elemento in paragrafos_para_remover:
                    if elemento.getparent() is not None:
                        # Remove o elemento
                        pai = elemento.getparent()
                        indice_no_pai = pai.index(elemento)
                        pai.remove(elemento)
                        logger.info("Parágrafo com [LINK_BOARD] em tabela removido com sucesso")
                        
                        try:
                            # Para tabelas, não adicionamos quebra de página, apenas um parágrafo vazio
                            # para manter o espaçamento
                            novo_paragrafo = celula.add_paragraph()
                            novo_elemento = novo_paragrafo._element
                            
                            # Remove o parágrafo da posição atual
                            if novo_elemento.getparent() is not None:
                                novo_elemento.getparent().remove(novo_elemento)
                            
                            # Adiciona o novo parágrafo na posição do removido
                            pai.insert(indice_no_pai, novo_elemento)
                            logger.info("Parágrafo vazio adicionado na tabela para substituir o removido")
                        except Exception as e:
                            logger.error(f"Erro ao adicionar parágrafo vazio na tabela: {str(e)}")

def processar_documento(doc, tipo_documento, substituicoes, atividades=None):
    """Função principal para processar o documento"""
    try:
        # Verifica se o link do board está vazio e, se estiver, remove os parágrafos relacionados
        if not substituicoes.get('[LINK_BOARD]') or substituicoes.get('[LINK_BOARD]').strip() == '':
            logger.info("Link do board não fornecido. Removendo parágrafos relacionados.")
            remover_paragrafo_link_board(doc)
        
         # Lista de marcadores para verificação rápida
        marcadores = set(substituicoes.keys())

        # Processa parágrafos no corpo do documento
        for paragraph in doc.paragraphs:
            try:
                # Verificação prévia se o parágrafo contém algum marcador
                texto_paragrafo = ''.join(run.text for run in paragraph.runs)
                
                # Só processa se contiver pelo menos um marcador ou o link do board
                if any(marcador in texto_paragrafo for marcador in marcadores):
                    substituir_texto_preservando_formatacao(paragraph, substituicoes)
            except Exception as e:
                logger.error(f"Erro ao processar parágrafo: {str(e)}")
        
        # Processa parágrafos nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        try:
                            # Verificação prévia se o parágrafo contém algum marcador
                            texto_paragrafo = ''.join(run.text for run in paragraph.runs)
                            
                            # Só processa se contiver pelo menos um marcador
                            if any(marcador in texto_paragrafo for marcador in marcadores):
                                substituir_texto_preservando_formatacao(paragraph, substituicoes)
                        except Exception as e:
                            logger.error(f"Erro ao processar parágrafo na tabela: {str(e)}")
        
        # Processa atividades se existirem
        if atividades:            
            if tipo_documento == 'estimativa':
                processar_atividades_tabela(doc, atividades)
            elif tipo_documento == 'estrategia':
                processar_atividades_bullet_points(doc, atividades)
            else:
                logger.info("O documento não tem atividades a serem processadas.")

    except Exception as e:
        logger.error(f"Erro ao processar documento: {str(e)}")
        raise

def gerar_nome_arquivo(tipo_documento, numero_ss, ano_ss):
    """Gera o nome do arquivo baseado no tipo de documento"""
    base_codigo = "CTNI-0534-2021.00"
    
    if tipo_documento == 'relatorio':
        prefixo = "AP"
        sufixo = "Relatório de Acompanhamento de Projeto"
    elif tipo_documento == 'estimativa':
        prefixo = "ES"
        sufixo = "Estimativa de Esforço e Cronograma"
    elif tipo_documento == 'estrategia':
        prefixo = "ES"
        sufixo = "Estratégia de solução"
    else:
        prefixo = "DOC"
        sufixo = "Documento"

    nome_final_arquivo = f"{prefixo}-{base_codigo}-{numero_ss}-{ano_ss} - {sufixo}.docx"
    logger.info(f"Nome final arquivo: {nome_final_arquivo}")
    
    return nome_final_arquivo

def obter_titulos_sumario(caminho_arquivo, titulo_documento=None):
    doc = None

    doc = Document(caminho_arquivo)
    logger.debug(f"Documento carregado com sucesso: {caminho_arquivo}")
    
    if titulo_documento and hasattr(doc, 'core_properties'):
        doc.core_properties.title = titulo_documento
        logger.debug(f"Título do documento atualizado para: {titulo_documento}")
    
    # 1. Identificar quebras de página no documento
    quebras_pagina = identificar_quebras_pagina(doc)
    logger.info(f"Identificadas {len(quebras_pagina)} quebras de página explícitas no documento")
    
    # 2. Encontrar início do sumário (parágrafo com o título "Sumário")
    sumario_inicio = None
    for i, paragrafo in enumerate(doc.paragraphs):
        if paragrafo.text.strip().lower() == "sumário":
            sumario_inicio = i
            logger.info(f"Título do sumário encontrado no parágrafo {i}")
            break
    
    if sumario_inicio is None:
        # Busca alternativa por parágrafo com estilo TOC_Heading
        for i, paragrafo in enumerate(doc.paragraphs):
            if hasattr(paragrafo, 'style') and paragrafo.style and "TOC" in str(paragrafo.style.name):
                sumario_inicio = i
                logger.info(f"Título do sumário encontrado por estilo no parágrafo {i}")
                break
    
    if sumario_inicio is None:
        # Última tentativa: localizar pelo contexto do documento
        for i, paragrafo in enumerate(doc.paragraphs):
            if paragrafo.text.strip().lower() in ["índice", "conteúdo", "table of contents"]:
                sumario_inicio = i
                logger.info(f"Possível título do sumário encontrado no parágrafo {i}: '{paragrafo.text}'")
                break

    if sumario_inicio is None:
        logger.error("Título do sumário não encontrado no documento!")
        return False
    
    # 3. Encontrar fim do sumário atual
    sumario_fim = None
    
    for i in range(sumario_inicio + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        texto = p.text.strip()
        estilo = p.style.name if hasattr(p, 'style') and p.style and hasattr(p.style, 'name') else ""
        
        if (estilo == 'Title' and texto.startswith("1")) or \
            re.match(r'^1\s*[-–]', texto) or \
            texto.lower().startswith("introdução"):
            sumario_fim = i
            logger.info(f"Introdução encontrada no parágrafo {i}: '{texto}'")
            break
            
    if sumario_fim is None:
        sumario_fim = sumario_inicio + 3
        logger.warning(f"Introdução não encontrada após o sumário. Usando valor relativo {sumario_fim}.")
    
    # 4. Criar mapa de páginas baseado nas quebras de página
    pagina_atual = 1
    mapa_paragrafos_paginas = {}
    
    # Definimos o primeiro parágrafo como página 1
    for i in range(len(doc.paragraphs)):
        # Verifica se há uma quebra de página neste parágrafo
        if i in quebras_pagina:
            pagina_atual += 1
        mapa_paragrafos_paginas[i] = pagina_atual
    
    # 5. Coletar todos os títulos, suas páginas corretas e criar marcadores para links
    titulos = []
    marcadores_titulos = {}
    
    # Primeiro, vamos adicionar marcadores (bookmarks) a todos os títulos
    for i, p in enumerate(doc.paragraphs):
        # Ignoramos parágrafos dentro do sumário
        if sumario_inicio <= i < sumario_fim:
            continue
            
        texto = p.text.strip()
        estilo = p.style.name if hasattr(p, 'style') and p.style and hasattr(p.style, 'name') else ""
        
        # Determinamos a página do título baseada no mapa de quebras de página
        pagina_do_titulo = mapa_paragrafos_paginas.get(i, 1)
        
        # Verifica se é um título
        eh_titulo = False
        nivel_titulo = 0
        
        if (estilo == 'Title' or estilo == 'Heading 1') and texto:
            eh_titulo = True
            nivel_titulo = 1
        elif estilo == 'Heading 2' and texto:
            eh_titulo = True
            nivel_titulo = 2
        elif estilo == 'Heading 3' and texto:
            eh_titulo = True
            nivel_titulo = 3
        
        # Se é um título, cria um marcador (bookmark)
        if eh_titulo:
            # Gera um ID único para o título
            import hashlib
            marcador_id = f"titulo_{hashlib.md5(texto.encode()).hexdigest()[:8]}"
            
            # Adiciona um marcador (bookmark) para o título
            from docx.oxml.shared import OxmlElement, qn

            # Cria marcador de início
            start = OxmlElement('w:bookmarkStart')
            start.set(qn('w:id'), str(i))  # Usando o índice como ID único
            start.set(qn('w:name'), marcador_id)
            
            # Cria marcador de fim
            end = OxmlElement('w:bookmarkEnd')
            end.set(qn('w:id'), str(i))
            
            # Insere o marcador no início do parágrafo
            p._p.insert(0, start)
            p._p.append(end)
            
            # Guarda o ID do marcador para usá-lo nos links
            marcadores_titulos[i] = marcador_id
            
            # Processamento específico para títulos de nível 1
            if nivel_titulo == 1:
                numero_match = re.match(r'(\d+(?:\.\d+)*)\s*[-–—]?\s*(.*)', texto)
                
                if numero_match:
                    numero, titulo_texto = numero_match.groups()
                    titulos.append({
                        'nivel': 1,
                        'numero': numero,
                        'texto': titulo_texto.strip(),
                        'pagina': pagina_do_titulo,
                        'marcador': marcador_id,
                        'indice': i
                    })
                else:
                    titulos.append({
                        'nivel': 1,
                        'numero': None,
                        'texto': texto,
                        'pagina': pagina_do_titulo,
                        'marcador': marcador_id,
                        'indice': i
                    })
            elif nivel_titulo > 1:
                titulos.append({
                    'nivel': nivel_titulo,
                    'numero': None,
                    'texto': texto,
                    'pagina': pagina_do_titulo,
                    'marcador': marcador_id,
                    'indice': i
                })

    return titulos

def atualizar_sumario_com_python_docx(caminho_arquivo, titulo_documento=None, titulos=None):
    """
    Atualiza o sumário de um documento usando python-docx e detecção de quebras de página.
    
    Args:
        caminho_arquivo (str): Caminho do arquivo DOCX
        titulo_documento (str): Título que será definido nas propriedades do documento
        
    Returns:
        bool: True se o sumário foi atualizado com sucesso, False caso contrário
    """
    logger.info(f"Iniciando atualização do sumário: {caminho_arquivo}")
    
    if not os.path.exists(caminho_arquivo):
        logger.error(f"Arquivo não encontrado: {caminho_arquivo}")
        return False
    
    doc = None
    
    try:
        doc = Document(caminho_arquivo)
        logger.debug(f"Documento carregado com sucesso: {caminho_arquivo}")
        
        if titulo_documento and hasattr(doc, 'core_properties'):
            doc.core_properties.title = titulo_documento
            logger.debug(f"Título do documento atualizado para: {titulo_documento}")
        
        # 1. Identificar quebras de página no documento
        quebras_pagina = identificar_quebras_pagina(doc)
        logger.info(f"Identificadas {len(quebras_pagina)} quebras de página explícitas no documento")
        
        # 2. Encontrar início do sumário (parágrafo com o título "Sumário")
        sumario_inicio = None
        for i, paragrafo in enumerate(doc.paragraphs):
            if paragrafo.text.strip().lower() == "sumário":
                sumario_inicio = i
                logger.info(f"Título do sumário encontrado no parágrafo {i}")
                break
        
        if sumario_inicio is None:
            # Busca alternativa por parágrafo com estilo TOC_Heading
            for i, paragrafo in enumerate(doc.paragraphs):
                if hasattr(paragrafo, 'style') and paragrafo.style and "TOC" in str(paragrafo.style.name):
                    sumario_inicio = i
                    logger.info(f"Título do sumário encontrado por estilo no parágrafo {i}")
                    break
        
        if sumario_inicio is None:
            # Última tentativa: localizar pelo contexto do documento
            for i, paragrafo in enumerate(doc.paragraphs):
                if paragrafo.text.strip().lower() in ["índice", "conteúdo", "table of contents"]:
                    sumario_inicio = i
                    logger.info(f"Possível título do sumário encontrado no parágrafo {i}: '{paragrafo.text}'")
                    break

        if sumario_inicio is None:
            logger.error("Título do sumário não encontrado no documento!")
            return False
        
        # 3. Encontrar fim do sumário atual
        sumario_fim = None
        
        for i in range(sumario_inicio + 1, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            texto = p.text.strip()
            estilo = p.style.name if hasattr(p, 'style') and p.style and hasattr(p.style, 'name') else ""
            
            if (estilo == 'Title' and texto.startswith("1")) or \
               re.match(r'^1\s*[-–]', texto) or \
               texto.lower().startswith("introdução"):
                sumario_fim = i
                logger.info(f"Introdução encontrada no parágrafo {i}: '{texto}'")
                break
                
        if sumario_fim is None:
            sumario_fim = sumario_inicio + 3
            logger.warning(f"Introdução não encontrada após o sumário. Usando valor relativo {sumario_fim}.")
        
        # 4. Criar mapa de páginas baseado nas quebras de página
        pagina_atual = 1
        mapa_paragrafos_paginas = {}
        
        # 5. Coletar todos os títulos, suas páginas corretas e criar marcadores para links
        marcadores_titulos = {}
        
        # Primeiro, vamos adicionar marcadores (bookmarks) a todos os títulos
        for i, p in enumerate(doc.paragraphs):
            # Ignoramos parágrafos dentro do sumário
            if sumario_inicio <= i < sumario_fim:
                continue
                
            texto = p.text.strip()
            estilo = p.style.name if hasattr(p, 'style') and p.style and hasattr(p.style, 'name') else ""
            
            # Determinamos a página do título baseada no mapa de quebras de página
            
            # Verifica se é um título
            eh_titulo = False
            
            if (estilo == 'Title' or estilo == 'Heading 1') and texto:
                eh_titulo = True
            elif estilo == 'Heading 2' and texto:
                eh_titulo = True
            elif estilo == 'Heading 3' and texto:
                eh_titulo = True
            
            # Se é um título, cria um marcador (bookmark)
            if eh_titulo:
                # Gera um ID único para o título
                import hashlib
                marcador_id = f"titulo_{hashlib.md5(texto.encode()).hexdigest()[:8]}"
                
                # Adiciona um marcador (bookmark) para o título
                from docx.oxml.shared import OxmlElement, qn

                # Cria marcador de início
                start = OxmlElement('w:bookmarkStart')
                start.set(qn('w:id'), str(i))  # Usando o índice como ID único
                start.set(qn('w:name'), marcador_id)
                
                # Cria marcador de fim
                end = OxmlElement('w:bookmarkEnd')
                end.set(qn('w:id'), str(i))
                
                # Insere o marcador no início do parágrafo
                p._p.insert(0, start)
                p._p.append(end)
                
                # Guarda o ID do marcador para usá-lo nos links
                marcadores_titulos[i] = marcador_id
                for tit in titulos:
                    if(tit["texto"] == texto):
                        tit["marcador"] = marcador_id 
                        break

        # Definimos o primeiro parágrafo como página 1
        for i in range(len(doc.paragraphs)):
            # Verifica se há uma quebra de página neste parágrafo
            if i in quebras_pagina:
                pagina_atual += 1
            mapa_paragrafos_paginas[i] = pagina_atual
        
        # 6. Remover os parágrafos do sumário antigo EXCETO o título TOC
        logger.info(f"Removendo sumário antigo do parágrafo {sumario_inicio+1} ao {sumario_fim-1}")
        for i in range(sumario_fim - 1, sumario_inicio, -1):
            try:
                p_element = doc.paragraphs[i]._element
                if p_element.getparent() is not None:
                    p_element.getparent().remove(p_element)
                    logger.debug(f"Removido parágrafo {i}")
            except Exception as e:
                logger.warning(f"Erro ao remover parágrafo {i}: {str(e)}")
        
        # 7. Criar itens do sumário com hyperlinks
        import docx.opc.constants
        from docx.enum.text import (WD_LINE_SPACING, WD_TAB_ALIGNMENT,
                                    WD_TAB_LEADER)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
        
        logger.info(f"Criando novo sumário com {len(titulos)} títulos")
        
        # Obter o parágrafo de referência (título do sumário)
        paragrafo_referencia = doc.paragraphs[sumario_inicio]
        elemento_referencia = paragrafo_referencia._element
        pai_elemento = elemento_referencia.getparent()
        
        # Obter o índice do elemento de referência entre seus irmãos
        indice_no_pai = pai_elemento.index(elemento_referencia)
        
        # Inserir os novos itens do sumário após o título
        posicao_atual = indice_no_pai
        
        # Função auxiliar para criar hyperlinks para marcadores internos sem sublinhado
        def criar_hyperlink_marcador(paragrafo, texto, marcador, cor_rgb=(51, 70, 130), tamanho_fonte=Pt(12), negrito=False):
            """
            Cria um hyperlink no parágrafo que aponta para um marcador interno.
            
            Args:
                paragrafo: O parágrafo onde o link será criado
                texto: O texto que será exibido para o link
                marcador: O nome do marcador (bookmark) para onde o link apontará
                cor_rgb: Cor RGB do texto (tupla de 3 valores)
                tamanho_fonte: Tamanho da fonte
                negrito: Se o texto deve estar em negrito
            """
            # Cria o elemento hyperlink
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), marcador)
            
            # Cria um novo run
            new_run = OxmlElement('w:r')
            
            # Adiciona as propriedades do run
            rPr = OxmlElement('w:rPr')
            
            # Adiciona o nome da fonte
            font = OxmlElement('w:rFonts')
            font.set(qn('w:ascii'), "Calibri")
            font.set(qn('w:hAnsi'), "Calibri")
            rPr.append(font)
            
            # Adiciona o tamanho da fonte
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(tamanho_fonte.pt * 2)))  # Word armazena o tamanho em 1/2 pts
            rPr.append(sz)
            
            # Adiciona o negrito, se necessário
            if negrito:
                b = OxmlElement('w:b')
                b.set(qn('w:val'), '1')
                rPr.append(b)
            
            # Adiciona a cor
            color = OxmlElement('w:color')
            r, g, b = cor_rgb
            color.set(qn('w:val'), f"{r:02X}{g:02X}{b:02X}")
            rPr.append(color)
            
            # Importante: NÃO adiciona o estilo de hyperlink nem o sublinhado
            # para manter o estilo visual igual aos itens sem link
            
            new_run.append(rPr)
            
            # Adiciona o texto
            t = OxmlElement('w:t')
            t.text = texto
            new_run.append(t)
            
            hyperlink.append(new_run)
            paragrafo._p.append(hyperlink)
            
            return hyperlink
        
        for titulo in titulos:
            try:
                p = doc.add_paragraph()
                
                # Remove o parágrafo de sua posição atual
                novo_elemento = p._element
                novo_elemento.getparent().remove(novo_elemento)
                
                # Insere após a posição atual (que começou com o título do sumário)
                posicao_atual += 1
                pai_elemento.insert(posicao_atual, novo_elemento)
                
                # Formatação do parágrafo
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Pt(12 * (titulo['nivel'] - 1))
                
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Cm(16.99), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
                
                texto_item = ""
                if titulo['nivel'] > 1:
                    texto_item += "  " * (titulo['nivel'] - 1)
                
                if titulo['numero']:
                    texto_item += f"{titulo['numero']} - {titulo['texto']}"
                else:
                    texto_item += titulo['texto']
                
                # Definir o tamanho da fonte com base no nível do título
                # Títulos nível 1: 12pt, negrito
                # Títulos nível 2+: 11pt, normal
                tamanho_fonte = Pt(12) if titulo['nivel'] == 1 else Pt(11)
                negrito = (titulo['nivel'] == 1)
                
                # Cria o hyperlink para o texto do item usando o marcador
                cor_rgb = (51, 70, 130)
                criar_hyperlink_marcador(
                    p, 
                    texto_item, 
                    titulo['marcador'], 
                    cor_rgb=cor_rgb, 
                    tamanho_fonte=tamanho_fonte, 
                    negrito=negrito
                )
                
                # Adiciona o tab
                tab_run = p.add_run("\t")
                
                # Adiciona o número da página como hyperlink também
                criar_hyperlink_marcador(
                    p,
                    str(titulo['pagina']),
                    titulo['marcador'],
                    cor_rgb=cor_rgb,
                    tamanho_fonte=tamanho_fonte,
                    negrito=negrito
                )
                    
                logger.debug(f"Adicionado item ao sumário: {texto_item} (Página {titulo['pagina']})")
            except Exception as e:
                logger.warning(f"Erro ao adicionar item ao sumário: {str(e)}")
                logger.exception("Detalhes do erro:")
        
        # 8. Adicionar quebra de página após o último item do sumário
        logger.info("Adicionando quebra de página após o sumário")
        
        # Criar parágrafo com quebra de página
        p_quebra = doc.add_paragraph()
        
        # Remove o parágrafo da posição atual
        p_quebra_element = p_quebra._element
        p_quebra_element.getparent().remove(p_quebra_element)
        
        # Insere após o último item do sumário
        pai_elemento.insert(posicao_atual + 1, p_quebra_element)
        
        # Adiciona a quebra de página
        run = p_quebra.add_run()
        run._element.append(OxmlElement('w:br'))
        run._element.get_or_add_rPr().append(OxmlElement('w:lastRenderedPageBreak'))
        
        # Configuração adicional da quebra de página para garantir que seja uma quebra de página real
        from docx.enum.text import WD_BREAK
        run.add_break(WD_BREAK.PAGE)
        
        logger.info("Quebra de página adicionada após o sumário")
        
        # 9. Adicionar parágrafo vazio após a quebra de página para manter o padrão de espaçamento
        logger.info("Adicionando parágrafo vazio após a quebra de página")
        
        # Criar parágrafo vazio
        p_espaco = doc.add_paragraph()
        
        # Remove o parágrafo da posição atual
        p_espaco_element = p_espaco._element
        p_espaco_element.getparent().remove(p_espaco_element)
        
        # Insere após o parágrafo de quebra de página
        pai_elemento.insert(posicao_atual + 2, p_espaco_element)
        
        # Configura o parágrafo vazio SEM espaçamento
        p_espaco.paragraph_format.space_before = Pt(0)
        p_espaco.paragraph_format.space_after = Pt(0)
        
        logger.info("Parágrafo vazio adicionado após a quebra de página")
        
        # Salvar o documento
        doc.save(caminho_arquivo)
        logger.info("Documento salvo com sucesso após atualização do sumário")
        
        return True
        
    except Exception as e:
        logger.error(f"Erro ao atualizar sumário: {str(e)}")
        logger.exception("Detalhes do erro:")
        
        if doc:
            try:
                doc.save(caminho_arquivo)
                logger.info("Documento salvo, apesar do erro na atualização do sumário")
            except Exception as save_err:
                logger.error(f"Não foi possível salvar o documento após erro: {str(save_err)}")
            
        return False

# Função auxiliar para identificar quebras de página (se você ainda não tiver uma)
def identificar_quebras_pagina(doc):
    """
    Identifica parágrafos que contêm quebras de página no documento.
    
    Args:
        doc: Objeto Document do python-docx
        
    Returns:
        list: Lista de índices de parágrafos que contêm quebras de página
    """
    quebras_pagina = []
    
    for i, p in enumerate(doc.paragraphs):
        # Verifica se há quebras de página explícitas
        for run in p.runs:
            if "lastRenderedPageBreak" in str(run._element.xml) or "w:br w:type=\"page\"" in str(run._element.xml):
                quebras_pagina.append(i)
                break
        
        # Verifica configurações do parágrafo que forçam quebras de página
        if hasattr(p._p, 'pPr') and p._p.pPr is not None:
            if p._p.pPr.xpath('.//w:pageBreakBefore'):
                quebras_pagina.append(i)
    
    return quebras_pagina

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    from docx.oxml import ns
    element.set(ns.qn(name), value)

def gerar_pdf_do_docx(caminho_arquivo):
    """
    Gera um arquivo PDF a partir de um documento DOCX usando LibreOffice.
    Detecta automaticamente o caminho de instalação do LibreOffice no Windows.
    
    Args:
        caminho_arquivo (str): Caminho do arquivo DOCX
        
    Returns:
        str: Caminho do arquivo PDF gerado ou None em caso de erro
    """
    logger.info(f"Iniciando geração de PDF: {caminho_arquivo}")
    
    import glob
    import os
    import platform
    import subprocess

    # Define o caminho do PDF a ser gerado
    pdf_path = os.path.splitext(caminho_arquivo)[0] + '.pdf'
    diretorio = os.path.dirname(os.path.abspath(caminho_arquivo))
    
    try:
        soffice_path = None
        
        # No Windows, precisamos encontrar o caminho do soffice.exe
        if platform.system() == 'Windows':
            # Locais comuns de instalação do LibreOffice
            possible_locations = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
            
            # Procura em locais comuns
            for location in possible_locations:
                if os.path.exists(location):
                    soffice_path = location
                    break
            
            # Busca por instalações em locais não padrão
            if not soffice_path:
                # Procura em Program Files
                program_files_dirs = [
                    os.environ.get('ProgramFiles', r'C:\Program Files'),
                    os.environ.get('ProgramFiles(x86)', r'C:\Program Files (x86)')
                ]
                
                for program_dir in program_files_dirs:
                    if os.path.exists(program_dir):
                        # Procura por pastas que começam com LibreOffice
                        libreoffice_dirs = glob.glob(os.path.join(program_dir, 'LibreOffice*'))
                        for libre_dir in libreoffice_dirs:
                            soffice_path_candidate = os.path.join(libre_dir, 'program', 'soffice.exe')
                            if os.path.exists(soffice_path_candidate):
                                soffice_path = soffice_path_candidate
                                break
            
            # Se ainda não encontrou, tenta através do registro do Windows
            if not soffice_path:
                try:
                    import winreg
                    with winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r'SOFTWARE\LibreOffice\Layers') as key:
                        install_path, _ = winreg.QueryValueEx(key, 'BRAND')
                        soffice_path = os.path.join(install_path, 'program', 'soffice.exe')
                        if not os.path.exists(soffice_path):
                            soffice_path = None
                except:
                    # Se falhar a busca no registro, continua com outras alternativas
                    pass
                    
            # Se não encontrou o LibreOffice, loga o erro e retorna
            if not soffice_path:
                logger.error("LibreOffice (soffice.exe) não encontrado. Verifique se está instalado.")
                return None
            
            logger.info(f"Usando LibreOffice em: {soffice_path}")
            
            # Comando para converter usando o LibreOffice no Windows
            cmd = [
                soffice_path, 
                '--headless', 
                '--convert-to', 'pdf',
                '--outdir', diretorio, 
                caminho_arquivo
            ]
        else:
            # Linux/Docker - mais simples
            cmd = [
                'libreoffice', 
                '--headless', 
                '--convert-to', 'pdf',
                '--outdir', diretorio, 
                caminho_arquivo
            ]
        
        logger.debug(f"Executando comando: {' '.join(cmd)}")
        
        # Executa o comando com shell=True para Windows
        process = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=60,  # Timeout de 60 segundos
            check=False,  # Não levanta exceção se o comando falhar
            shell=(platform.system() == 'Windows')  # Usa shell no Windows
        )
        
        # Verifica se o processo foi bem-sucedido
        if process.returncode != 0:
            logger.error(f"Erro ao gerar PDF: {process.stderr}")
            return None
        
        # Verifica se o arquivo foi realmente criado
        if os.path.exists(pdf_path):
            logger.info(f"PDF gerado com sucesso: {pdf_path}")
            return pdf_path
        else:
            logger.error(f"Arquivo PDF não foi criado: {pdf_path}")
            return None
        
    except subprocess.TimeoutExpired:
        logger.error("Timeout ao converter para PDF")
        return None
        
    except Exception as e:
        logger.error(f"Erro durante geração do PDF: {str(e)}")
        logger.exception("Detalhes do erro:")
        return None

# Rotas Flask
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/dev')
def dev_index():
    return render_template('index-dev.html')

@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static'),
                             'favicon.ico', mimetype='image/vnd.microsoft.icon')


def encontrar_pagina_pdf(file, text, pag_min):
    import re

    import PyPDF2

    pdf = PyPDF2.PdfReader(file)

    num_pages = len(pdf.pages)
    for i in range(0, num_pages):
        page = pdf.pages[i]
        page_text = page.extract_text()
        if re.search(text, page_text) and i >= pag_min - 1:
            return i
    return None

def obter_paginas_pdf(caminho_pdf, titulos):
    titulos_aux = []
    for titulo in titulos:
        pagina = encontrar_pagina_pdf(caminho_pdf, titulo["texto"], titulo["pagina"])
        if pagina is not None:
            titulo["pagina"] = pagina + 1
        titulos_aux.append(titulo)
    return titulos_aux

@app.route('/gerar_relatorio', methods=['POST'])
def gerar_relatorio():
    """
    Rota principal para geração de relatórios em DOCX e/ou PDF.
    Processa os modelos de documento, substitui os marcadores e gera os arquivos nos formatos solicitados.
    """
    logger.info("Iniciando geração de relatórios")
    
    # Listas para controle de arquivos temporários e arquivos de saída
    temp_files = []    # Arquivos temporários que serão deletados ao final
    output_files = []  # Arquivos que serão incluídos no ZIP final
    zip_path = None    # Caminho do arquivo ZIP que será gerado
    
    try:
        # Obtém os dados do formulário
        dados = request.form
        
        # Obtém o tipo de formulário (tecnico ou desenvolvimento)
        tipo_formulario = dados.get('tipo', 'tecnico')
        
        # Processa as opções de formato selecionadas pelo usuário
        gerar_json = request.form.get('gerar_json', 'true').lower() == 'true'  # Por padrão, sempre gera JSON
        gerar_docx = request.form.get('gerar_docx', 'false').lower() == 'true'
        gerar_pdf = request.form.get('gerar_pdf', 'false').lower() == 'true'
        apenas_json = request.form.get('apenas_json', 'false').lower() == 'true'

        # Dados específicos de cada tipo de formulário
        if tipo_formulario == 'tecnico':
            atividades = json.loads(dados.get('atividades', '[]'))
            total_horas = dados.get('totalHoras', 0)
            # Outros dados específicos do formulário técnico
            dados_especificos = {
                'atividades': atividades,
                'totalHoras': total_horas
            }
        else:  # desenvolvimento
            requisitos = json.loads(dados.get('requisitos', '[]'))
            requisitos_nao_funcionais = json.loads(dados.get('requisitosNaoFuncionais', '[]'))
            total_pontos_funcao = dados.get('totalPontosFuncao', 0)
            # Outros dados específicos do formulário de desenvolvimento
            dados_especificos = {
                'requisitos': requisitos,
                'requisitosNaoFuncionais': requisitos_nao_funcionais,
                'totalPontosFuncao': total_pontos_funcao
            }

        if apenas_json:
            try:
                # Criamos um dicionário com os dados disponíveis
                dados_json = {
                    'tipo': tipo_formulario,  # Adiciona o tipo
                    'info': {
                        'numeroSS': dados.get('numeroSS', ''),
                        'anoSS': dados.get('anoSS', ''), 
                        'tituloSS': dados.get('tituloSS', ''),
                        'descricao': dados.get('descricao', ''),
                        'dataInicio': dados.get('dataInicio', ''),
                        'dataFim': dados.get('dataFim', ''),
                        'linkBoard': dados.get('linkBoard', ''),
                        'iniciaisAutor': dados.get('iniciaisAutor', '')
                    }
                }
                
                # Adiciona os dados específicos ao JSON com base no tipo
                dados_json.update(dados_especificos)
                
                # Caminho e nome do arquivo JSON
                json_filename = f"SS {dados.get('numeroSS', '').zfill(3)}-{dados.get('anoSS', '')}.json"
                json_path = os.path.join(tempfile.gettempdir(), json_filename)
                
                # Salvar arquivo JSON
                with open(json_path, 'w', encoding='utf-8') as json_file:
                    json.dump(dados_json, json_file, ensure_ascii=False, indent=2)
                
                # Retornar resposta JSON com confirmação
                return send_file(
                    json_path,
                    as_attachment=True,
                    download_name=json_filename,
                    mimetype='application/json'
                )
                
            except Exception as e:
                logger.error(f"Erro ao gerar arquivo JSON: {str(e)}")
                return jsonify({"error": f"Erro ao gerar JSON: {str(e)}"}), 500

        json_path = None
        try:
            # Criamos um dicionário com todos os dados
            dados_json = {
                'info': {
                    'numeroSS': dados['numeroSS'],
                    'anoSS': dados['anoSS'], 
                    'tituloSS': dados['tituloSS'],
                    'descricao': dados['descricao'],
                    'dataInicio': dados['dataInicio'],
                    'dataFim': dados['dataFim'],
                    'totalHoras': dados.get('totalHoras', 0),
                    'linkBoard': dados['linkBoard'],
                    'iniciaisAutor': dados['iniciaisAutor']
                },
                'atividades': atividades
            }
            
            # Caminho do arquivo JSON
            json_path = os.path.join(tempfile.gettempdir(), f"SS {dados['numeroSS'].zfill(3)}-{dados['anoSS']}.json")
            
            # Salvar arquivo JSON
            with open(json_path, 'w', encoding='utf-8') as json_file:
                json.dump(dados_json, json_file, ensure_ascii=False, indent=2)
            
            # Adicionar à lista de arquivos temporários e de saída
            temp_files.append(json_path)
            output_files.append({
                'path': json_path,
                'filename': os.path.basename(json_path)
            })
            
        except Exception as e:
            logger.error(f"Erro ao gerar arquivo JSON: {str(e)}")

        # Cria instância do relatório com os dados fornecidos
        relatorio = RelatorioAcompanhamentoProjeto(
            numero_ss=dados['numeroSS'],
            ano_ss=dados['anoSS'],
            iniciais_autor=dados['iniciaisAutor'],
            titulo_ss=dados['tituloSS'],
            descricao=dados['descricao'],
            data_inicio=dados['dataInicio'],
            data_fim=dados['dataFim'],
            total_horas=dados.get('totalHoras', 0),
            link_board=dados['linkBoard'],
            atividades=atividades
        )

        # Define o caminho dos modelos de documento
        modelos_path = os.path.join(app.root_path, 'modelos', 'tecnica')
        modelos = [
            os.path.join(modelos_path, "ModeloTec - Estimativa de Esforço e Cronograma.docx"),
            os.path.join(modelos_path, "ModeloTec - Estratégia de solução.docx"),
            os.path.join(modelos_path, "ModeloTec - Relatório de Acompanhamento de Projeto.docx")
        ]
        
        # Verifica se todos os modelos existem
        for modelo in modelos:
            if not os.path.exists(modelo):
                return {"error": f"Modelo não encontrado: {os.path.basename(modelo)}"}, 404

        # Obtém as substituições que serão feitas nos documentos
        substituicoes = relatorio.get_substituicoes()
        
        # Processa cada modelo de documento
        for modelo_path in modelos:
            # Cria uma cópia temporária do modelo
            temp_output = None
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_output = temp_file.name
                temp_files.append(temp_output)
                shutil.copy2(modelo_path, temp_output)
            
            try:
                # Processa o documento em um contexto controlado
                with open(temp_output, 'rb') as doc_file:
                    doc = Document(doc_file)
                    doc.core_properties.title = os.path.basename(modelo_path)
                    
                    # Identifica o tipo de documento e gera o nome do arquivo
                    tipo_documento = identificar_tipo_documento(doc)
                    nome_arquivo = gerar_nome_arquivo(
                        tipo_documento,
                        dados['numeroSS'].zfill(3),
                        dados['anoSS']
                    )

                    # Define o título do documento sem a extensão
                    titulo_documento = os.path.splitext(nome_arquivo)[0]
                    logger.debug(f"Título definido no DOCX: {titulo_documento}")
                    
                    # Define as propriedades do documento DOCX
                    doc.core_properties.title = titulo_documento

                    # Processa o documento substituindo os marcadores
                    processar_documento(doc, tipo_documento, substituicoes, atividades)
                    
                    # Define o caminho do arquivo temporário final
                    temp_final = os.path.join(tempfile.gettempdir(), nome_arquivo)

                    # Salvamos o documento com as substituições
                    doc.save(temp_final)
                    pdf_pre_sumario = gerar_pdf_do_docx(temp_final)
                    titulos_sumario = obter_titulos_sumario(temp_final)
                    titulos_atualizados = obter_paginas_pdf(pdf_pre_sumario, titulos_sumario)

                    temp_files.append(temp_final)
                    
                    # Atualizamos o sumário - SEMPRE, independente se vai gerar PDF ou não
                    atualizar_sumario_com_python_docx(temp_final, titulo_documento, titulos_atualizados)
                        
                    # Adiciona o DOCX aos arquivos de saída se solicitado
                    if gerar_docx:
                        output_files.append({
                            'path': temp_final,
                            'filename': nome_arquivo
                        })
                    
                    # Gera PDF se solicitado
                    if gerar_pdf:
                        pdf_path = gerar_pdf_do_docx(temp_final)
                        if pdf_path:
                            temp_files.append(pdf_path)
                            output_files.append({
                                'path': pdf_path,
                                'filename': os.path.splitext(nome_arquivo)[0] + '.pdf'
                            })
                        else:
                            # Se falhou ao gerar PDF mas o usuário quer só PDF, adiciona mensagem de erro
                            if not gerar_docx:
                                return {"error": "Falha ao gerar PDF"}, 500
                
            except Exception as e:
                logger.error(f"Erro ao processar documento {modelo_path}: {str(e)}")
                raise
        
        # Verifica se temos arquivos para incluir no ZIP
        if not output_files:
            return {"error": "Nenhum arquivo foi gerado com sucesso"}, 500
        
        # Cria nome do arquivo ZIP
        zip_filename = f"SS {dados['numeroSS'].zfill(3)}-{dados['anoSS']}.zip"
        zip_path = os.path.join(tempfile.gettempdir(), zip_filename)
        
        # Cria o arquivo ZIP com os documentos gerados
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for output in output_files:
                if os.path.exists(output['path']):
                    zipf.write(output['path'], output['filename'])

        # Configura e retorna a resposta com o arquivo ZIP
        response = send_file(
            zip_path,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )

        # Garante que o Content-Disposition esteja correto para caracteres especiais
        encoded_filename = quote(zip_filename)
        response.headers["Content-Disposition"] = (
            f"attachment; filename=\"{zip_filename}\"; "
            f"filename*=UTF-8''{encoded_filename}"
        )

        # Função para limpeza dos arquivos temporários após envio da resposta
        @response.call_on_close
        def cleanup():
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.unlink(temp_file)
                except Exception as e:
                    logger.error(f"Erro ao limpar arquivo temporário: {e}")
            
            try:
                if os.path.exists(zip_path):
                    os.unlink(zip_path)
            except Exception as e:
                logger.error(f"Erro ao limpar arquivo ZIP: {e}")

        return response

    except Exception as e:
        logger.exception("Erro durante a geração dos relatórios")
        return {"error": str(e)}, 500
        
    finally:
        # Limpeza final de segurança dos arquivos temporários
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except:
                pass
            
        try:
            if zip_path and os.path.exists(zip_path):
                os.unlink(zip_path)
        except:
            pass

if __name__ == '__main__':
    app.run(debug=False)